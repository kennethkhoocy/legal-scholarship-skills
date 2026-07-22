"""Build a new .docx document from a structured JSON specification.

Supports two modes:

1. **Template mode** -- When a ``template_path`` is supplied, the spec is
   rendered into a ``.docx`` template via *docxtpl* (Jinja2 templating for Word).
2. **Programmatic mode** -- Without a template, a new document is assembled
   from scratch using *python-docx*, with headings, sections, and an optional
   "Response to Comments" block.

After writing, the output is validated as a well-formed ZIP archive containing
``[Content_Types].xml``, and a sidecar manifest JSON is written alongside the
output file.

Typical usage::

    from pathlib import Path
    from build_docx import build_from_spec, write_build_manifest
    from models import BuildSpec

    spec = BuildSpec.model_validate_json(Path("spec.json").read_text())
    diagnostics = build_from_spec(spec, Path("output.docx"))
    write_build_manifest(Path("output.docx"), spec, diagnostics)
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

_DIR = Path(__file__).resolve().parent
if str(_DIR) not in sys.path:
    sys.path.insert(0, str(_DIR))

from models import BuildSpec, DiagnosticEntry

logger = logging.getLogger(__name__)


# ── Routing helpers ────────────────────────────────────────────────────────

class BuildSpecRequiresJSError(RuntimeError):
    """Raised when runtime='python' is set on a spec that uses js-only features."""


def needs_js_runtime(spec) -> bool:
    """Return True iff the spec uses any feature only the JS path can deliver."""
    if spec.toc is not None:
        return True
    if spec.columns > 1:
        return True
    for ps in spec.page_sections:
        if any(el.type == "page_number" for el in (ps.header + ps.footer)):
            return True
    if spec.internal_links:
        return True
    if spec.native_footnotes:
        return True
    return False


def decide_runtime(spec) -> str:
    """Resolve spec.runtime ('auto' | 'python' | 'js') to a concrete backend."""
    if spec.runtime == "js":
        return "js"
    if spec.runtime == "python":
        if needs_js_runtime(spec):
            raise BuildSpecRequiresJSError(
                "Spec sets runtime='python' but uses js-only features "
                "(toc / columns>1 / page_number / internal_links / native_footnotes)."
            )
        return "python"
    # auto
    return "js" if needs_js_runtime(spec) else "python"


# ── Validation ─────────────────────────────────────────────────────────────

def _validate_docx(path: Path) -> list[DiagnosticEntry]:
    """Verify that *path* is a valid ZIP containing ``[Content_Types].xml``."""
    diagnostics: list[DiagnosticEntry] = []
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            names = zf.namelist()
            if "[Content_Types].xml" not in names:
                diagnostics.append(DiagnosticEntry(
                    level="error", source="build_docx",
                    message=(
                        f"Output file {path.name} is a ZIP but lacks "
                        "[Content_Types].xml -- may be corrupt"
                    ),
                ))
            else:
                diagnostics.append(DiagnosticEntry(
                    level="info", source="build_docx",
                    message=f"Validated {path.name} as a well-formed .docx",
                ))
    except zipfile.BadZipFile:
        diagnostics.append(DiagnosticEntry(
            level="error", source="build_docx",
            message=f"Output file {path.name} is not a valid ZIP archive",
        ))
    return diagnostics


# ── Programmatic build ─────────────────────────────────────────────────────

def _apply_page_config(doc, spec: BuildSpec) -> None:
    """Apply page size, orientation, margins, and text-only headers/footers."""
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENT

    section = doc.sections[0]

    # Page size
    if spec.page.size.lower() == "a4":
        section.page_width = Inches(11906 / 1440)
        section.page_height = Inches(16838 / 1440)
    elif spec.page.size.lower() == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Orientation
    if spec.page.orientation.lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # Margins
    section.top_margin = Inches(spec.page.margins.top)
    section.bottom_margin = Inches(spec.page.margins.bottom)
    section.left_margin = Inches(spec.page.margins.left)
    section.right_margin = Inches(spec.page.margins.right)

    # Simple text-only headers/footers (page_number elements route to JS path)
    for ps in spec.page_sections:
        if any(el.type == "text" for el in ps.header):
            text = " ".join(el.value or "" for el in ps.header if el.type == "text")
            if text:
                section.header.paragraphs[0].text = text
        if any(el.type == "text" for el in ps.footer):
            text = " ".join(el.value or "" for el in ps.footer if el.type == "text")
            if text:
                section.footer.paragraphs[0].text = text


def _build_programmatic(spec: BuildSpec, output_path: Path) -> list[DiagnosticEntry]:
    """Create a .docx from *spec* using python-docx."""
    from docx import Document
    from docx.shared import Pt  # noqa: F401 -- available for future formatting

    diagnostics: list[DiagnosticEntry] = []

    try:
        doc = Document()

        _apply_page_config(doc, spec)

        doc.add_heading(spec.title, level=0)
        if spec.subtitle:
            doc.add_paragraph(spec.subtitle)

        for section in spec.sections:
            doc.add_heading(section.heading, level=1)
            for para_text in section.paragraphs:
                doc.add_paragraph(para_text)

        if spec.items:
            doc.add_heading("Response to Comments", level=1)
            for item in spec.items:
                doc.add_heading(item.comment_id, level=2)
                # Comment in italic.
                p = doc.add_paragraph()
                run = p.add_run(f"Comment: {item.comment}")
                run.italic = True
                doc.add_paragraph(f"Response: {item.response}")
                if item.revision_made:
                    doc.add_paragraph(f"Revision: {item.revision_made}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        diagnostics.append(DiagnosticEntry(
            level="info", source="build_docx",
            message=f"Built {output_path.name} (programmatic mode)",
        ))
    except Exception as exc:
        diagnostics.append(DiagnosticEntry(
            level="error", source="build_docx",
            message=f"Programmatic build failed: {exc}",
        ))

    return diagnostics


# ── Template build ─────────────────────────────────────────────────────────

def _build_from_template(
    spec: BuildSpec,
    output_path: Path,
    template_path: Path,
) -> list[DiagnosticEntry]:
    """Render *spec* into a docxtpl template."""
    from docxtpl import DocxTemplate

    diagnostics: list[DiagnosticEntry] = []

    if not template_path.exists():
        diagnostics.append(DiagnosticEntry(
            level="error", source="build_docx",
            message=f"Template not found: {template_path}",
        ))
        return diagnostics

    try:
        doc = DocxTemplate(str(template_path))
        doc.render(spec.model_dump())
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        diagnostics.append(DiagnosticEntry(
            level="info", source="build_docx",
            message=(
                f"Built {output_path.name} from template "
                f"{template_path.name}"
            ),
        ))
    except Exception as exc:
        diagnostics.append(DiagnosticEntry(
            level="error", source="build_docx",
            message=f"Template build failed: {exc}",
        ))

    return diagnostics


# ── Footnote injection ────────────────────────────────────────────────────

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_FOOTNOTE_MARKER_RE = re.compile(r"\[\^(\d+)\]")


def _inject_footnotes(
    docx_path: Path,
    footnotes: dict[str, str],
) -> list[DiagnosticEntry]:
    """Post-process a saved .docx to inject OOXML footnotes.

    Scans ``word/document.xml`` for ``[^N]`` markers, replaces each with a
    ``w:footnoteReference`` run, creates ``word/footnotes.xml``, and updates
    the relationship and content-type entries.  Uses an atomic write via
    ``tempfile.mkstemp`` + ``os.replace``.
    """
    from lxml import etree

    diagnostics: list[DiagnosticEntry] = []

    # Read all ZIP entries into memory.
    all_entries: dict[str, bytes] = {}
    with zipfile.ZipFile(str(docx_path), "r") as zf:
        for name in zf.namelist():
            all_entries[name] = zf.read(name)

    # ── Parse document.xml and replace markers ───────────────────────────
    doc_xml = etree.fromstring(all_entries["word/document.xml"])
    ns = {"w": W_NS}
    used_ids: set[str] = set()

    # Pre-processing: merge adjacent runs with identical formatting
    # within each paragraph so that [^N] markers split across runs
    # are reassembled before scanning.
    for para in doc_xml.iter(f"{{{W_NS}}}p"):
        runs = list(para.findall(f"{{{W_NS}}}r"))
        i = 0
        while i < len(runs) - 1:
            r1 = runs[i]
            r2 = runs[i + 1]
            t1_list = r1.findall(f"{{{W_NS}}}t")
            t2_list = r2.findall(f"{{{W_NS}}}t")
            # Only merge if both runs are pure text (one w:t, optional w:rPr, nothing else)
            rpr_tag = f"{{{W_NS}}}rPr"
            t_tag = f"{{{W_NS}}}t"
            r1_children = [c.tag for c in r1]
            r2_children = [c.tag for c in r2]
            r1_pure = all(t in (rpr_tag, t_tag) for t in r1_children)
            r2_pure = all(t in (rpr_tag, t_tag) for t in r2_children)
            if len(t1_list) == 1 and len(t2_list) == 1 and r1_pure and r2_pure:
                t1, t2 = t1_list[0], t2_list[0]
                if t1.text and t2.text:
                    rpr1 = etree.tostring(r1.find(f"{{{W_NS}}}rPr") or etree.Element("x"))
                    rpr2 = etree.tostring(r2.find(f"{{{W_NS}}}rPr") or etree.Element("x"))
                    if rpr1 == rpr2:
                        t1.text = (t1.text or "") + (t2.text or "")
                        para.remove(r2)
                        runs.pop(i + 1)
                        continue
            i += 1

    for t_elem in doc_xml.iter(f"{{{W_NS}}}t"):
        text = t_elem.text or ""
        markers = list(_FOOTNOTE_MARKER_RE.finditer(text))
        if not markers:
            continue

        # The <w:t> lives inside a <w:r>; we will replace that run with
        # a sequence of runs: text-before, footnoteRef, text-between, ...
        run = t_elem.getparent()
        para = run.getparent()
        run_index = list(para).index(run)

        # Gather the run properties (rPr) from the original run so that
        # plain-text segments keep their formatting.
        orig_rpr = run.find(f"{{{W_NS}}}rPr")

        new_elements: list[etree._Element] = []
        prev_end = 0

        for m in markers:
            fn_id = m.group(1)
            if fn_id not in footnotes:
                diagnostics.append(DiagnosticEntry(
                    level="warning", source="build_docx",
                    message=f"Footnote marker [^{fn_id}] has no entry in footnotes dict; left as literal text.",
                ))
                # Emit the literal marker text as a plain segment
                literal = text[prev_end:m.end()]
                if literal:
                    r = etree.SubElement(etree.Element("dummy"), f"{{{W_NS}}}r")
                    if orig_rpr is not None:
                        from copy import deepcopy
                        r.append(deepcopy(orig_rpr))
                    t = etree.SubElement(r, f"{{{W_NS}}}t")
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    t.text = literal
                    new_elements.append(r)
                prev_end = m.end()
                continue
            used_ids.add(fn_id)

            # Text segment before this marker.
            before = text[prev_end:m.start()]
            if before:
                r = etree.SubElement(
                    etree.Element("dummy"), f"{{{W_NS}}}r",
                )
                if orig_rpr is not None:
                    from copy import deepcopy
                    r.append(deepcopy(orig_rpr))
                t = etree.SubElement(r, f"{{{W_NS}}}t")
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t.text = before
                new_elements.append(r)

            # Footnote reference run.
            fn_run = etree.SubElement(
                etree.Element("dummy"), f"{{{W_NS}}}r",
            )
            fn_rpr = etree.SubElement(fn_run, f"{{{W_NS}}}rPr")
            etree.SubElement(
                fn_rpr, f"{{{W_NS}}}rStyle",
            ).set(f"{{{W_NS}}}val", "FootnoteReference")
            etree.SubElement(
                fn_run, f"{{{W_NS}}}footnoteReference",
            ).set(f"{{{W_NS}}}id", fn_id)
            new_elements.append(fn_run)

            prev_end = m.end()

        # Trailing text after the last marker.
        after = text[prev_end:]
        if after:
            r = etree.SubElement(
                etree.Element("dummy"), f"{{{W_NS}}}r",
            )
            if orig_rpr is not None:
                from copy import deepcopy
                r.append(deepcopy(orig_rpr))
            t = etree.SubElement(r, f"{{{W_NS}}}t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = after
            new_elements.append(r)

        # Replace the original run in the paragraph.
        for i, elem in enumerate(new_elements):
            para.insert(run_index + i, elem)
        para.remove(run)

    if not used_ids:
        has_marker_warnings = any(
            "[^" in d.message and "no entry" in d.message
            for d in diagnostics
        )
        if not has_marker_warnings:
            diagnostics.append(DiagnosticEntry(
                level="warning", source="build_docx",
                message="Footnotes dict provided but no [^N] markers found in document text.",
            ))
        return diagnostics

    # ── Build or merge word/footnotes.xml ─────────────────────────────────
    if "word/footnotes.xml" in all_entries:
        fn_root = etree.fromstring(all_entries["word/footnotes.xml"])
        diagnostics.append(DiagnosticEntry(
            level="info", source="build_docx",
            message="Merging new footnotes with existing word/footnotes.xml.",
        ))
        # Find existing footnote IDs to avoid collisions
        existing_fn_ids = set()
        for fn_el in fn_root.findall(f"{{{W_NS}}}footnote"):
            fid = fn_el.get(f"{{{W_NS}}}id")
            if fid:
                try:
                    existing_fn_ids.add(int(fid))
                except ValueError:
                    pass
        # Remap new IDs if they collide, tracking all allocated IDs
        id_remap: dict[str, str] = {}
        all_taken = set(existing_fn_ids)
        # Also reserve IDs requested by non-colliding markers
        for fn_id in used_ids:
            all_taken.add(int(fn_id))
        next_id = max(all_taken | {0}) + 1
        for fn_id in sorted(used_ids, key=int):
            if int(fn_id) in existing_fn_ids:
                new_id = str(next_id)
                id_remap[fn_id] = new_id
                all_taken.add(next_id)
                next_id += 1
                while next_id in all_taken:
                    next_id += 1
                diagnostics.append(DiagnosticEntry(
                    level="info", source="build_docx",
                    message=f"Remapped footnote ID {fn_id} -> {new_id} to avoid collision.",
                ))
        # Update footnoteReference elements in document.xml with remapped IDs
        if id_remap:
            for fn_ref in doc_xml.iter(f"{{{W_NS}}}footnoteReference"):
                old_id = fn_ref.get(f"{{{W_NS}}}id")
                if old_id in id_remap:
                    fn_ref.set(f"{{{W_NS}}}id", id_remap[old_id])
    else:
        fn_root = etree.Element(
            f"{{{W_NS}}}footnotes",
            nsmap={"w": W_NS, "r": R_NS},
        )

        # Separator (id=-1).
        sep = etree.SubElement(fn_root, f"{{{W_NS}}}footnote")
        sep.set(f"{{{W_NS}}}type", "separator")
        sep.set(f"{{{W_NS}}}id", "-1")
        sep_p = etree.SubElement(sep, f"{{{W_NS}}}p")
        sep_r = etree.SubElement(sep_p, f"{{{W_NS}}}r")
        etree.SubElement(sep_r, f"{{{W_NS}}}separator")

        # Continuation separator (id=0).
        cont = etree.SubElement(fn_root, f"{{{W_NS}}}footnote")
        cont.set(f"{{{W_NS}}}type", "continuationSeparator")
        cont.set(f"{{{W_NS}}}id", "0")
        cont_p = etree.SubElement(cont, f"{{{W_NS}}}p")
        cont_r = etree.SubElement(cont_p, f"{{{W_NS}}}r")
        etree.SubElement(cont_r, f"{{{W_NS}}}continuationSeparator")

    # User-defined footnotes (use remapped IDs if available).
    id_remap = id_remap if "id_remap" in dir() else {}
    for fn_id in sorted(used_ids, key=int):
        actual_id = id_remap.get(fn_id, fn_id)
        fn_elem = etree.SubElement(fn_root, f"{{{W_NS}}}footnote")
        fn_elem.set(f"{{{W_NS}}}id", actual_id)
        fn_p = etree.SubElement(fn_elem, f"{{{W_NS}}}p")
        fn_ppr = etree.SubElement(fn_p, f"{{{W_NS}}}pPr")
        etree.SubElement(fn_ppr, f"{{{W_NS}}}pStyle").set(
            f"{{{W_NS}}}val", "FootnoteText",
        )
        # Self-reference run.
        ref_r = etree.SubElement(fn_p, f"{{{W_NS}}}r")
        ref_rpr = etree.SubElement(ref_r, f"{{{W_NS}}}rPr")
        etree.SubElement(ref_rpr, f"{{{W_NS}}}rStyle").set(
            f"{{{W_NS}}}val", "FootnoteReference",
        )
        etree.SubElement(ref_r, f"{{{W_NS}}}footnoteRef")
        # Footnote text run.
        txt_r = etree.SubElement(fn_p, f"{{{W_NS}}}r")
        txt_t = etree.SubElement(txt_r, f"{{{W_NS}}}t")
        txt_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        txt_t.text = " " + footnotes[fn_id]

    fn_xml_bytes = etree.tostring(
        fn_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # ── Update document.xml ──────────────────────────────────────────────
    all_entries["word/document.xml"] = etree.tostring(
        doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True,
    )
    all_entries["word/footnotes.xml"] = fn_xml_bytes

    # ── Add relationship for footnotes.xml ───────────────────────────────
    rels_path = "word/_rels/document.xml.rels"
    if rels_path in all_entries:
        rels_tree = etree.fromstring(all_entries[rels_path])
    else:
        rels_tree = etree.Element(
            "Relationships",
            xmlns="http://schemas.openxmlformats.org/package/2006/relationships",
        )

    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    # Check if a footnotes relationship already exists.
    fn_rel_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    )
    existing = [
        r for r in rels_tree
        if r.get("Type") == fn_rel_type
    ]
    if not existing:
        # Pick a unique rId.
        existing_ids = {r.get("Id") for r in rels_tree}
        rid_num = 1
        while f"rId{rid_num}" in existing_ids:
            rid_num += 1
        new_rel = etree.SubElement(rels_tree, f"{{{rels_ns}}}Relationship")
        new_rel.set("Id", f"rId{rid_num}")
        new_rel.set("Type", fn_rel_type)
        new_rel.set("Target", "footnotes.xml")

    all_entries[rels_path] = etree.tostring(
        rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # ── Add content type for footnotes.xml ───────────────────────────────
    ct_xml = etree.fromstring(all_entries["[Content_Types].xml"])
    ct_ns = ct_xml.nsmap.get(None, CT_NS)
    fn_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
    existing_ct = [
        o for o in ct_xml
        if o.get("PartName") == "/word/footnotes.xml"
    ]
    if not existing_ct:
        override = etree.SubElement(ct_xml, f"{{{ct_ns}}}Override")
        override.set("PartName", "/word/footnotes.xml")
        override.set("ContentType", fn_ct)

    all_entries["[Content_Types].xml"] = etree.tostring(
        ct_xml, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # ── Atomic write ─────────────────────────────────────────────────────
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=docx_path.parent)
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in all_entries.items():
                zf.writestr(name, data)
        Path(tmp_path).replace(docx_path)
    except Exception as exc:
        Path(tmp_path).unlink(missing_ok=True)
        diagnostics.append(DiagnosticEntry(
            level="error", source="build_docx",
            message=f"Footnote injection failed during write: {exc}",
        ))
        return diagnostics

    diagnostics.append(DiagnosticEntry(
        level="info", source="build_docx",
        message=f"Injected {len(used_ids)} footnote(s) into {docx_path.name}.",
    ))
    return diagnostics


# ── Public API ─────────────────────────────────────────────────────────────

def build_from_spec(
    spec: BuildSpec,
    output_path: Path,
    template_path: Path | None = None,
) -> list[DiagnosticEntry]:
    """Build a ``.docx`` from *spec*, routing to python or JS backend.

    Parameters
    ----------
    spec:
        The structured build specification (title, sections, items).
    output_path:
        Where to write the resulting ``.docx``.
    template_path:
        If provided, a ``.docx`` template rendered via *docxtpl* (python path
        only).

    Returns
    -------
    list[DiagnosticEntry]
        Diagnostics from the build and subsequent ZIP validation.

    Raises
    ------
    BuildSpecRequiresJSError
        When the spec sets ``runtime='python'`` but uses JS-only features
        (toc / columns>1 / page_number / internal_links / native_footnotes).
    """
    runtime = decide_runtime(spec)
    if runtime == "python":
        diagnostics = _build_python(spec, output_path, template_path)
    else:
        diagnostics = _build_js(spec, Path(output_path))

    # Post-build validation (best-effort; warn but do not block on validator failure)
    try:
        from commands.validate import run as validate_run

        code = validate_run(output_path)
        diagnostics.append(DiagnosticEntry(
            level="info" if code == 0 else "warning",
            source="post_build_validate",
            message=f"validate exited {code}",
        ))
    except Exception:
        # Validator unavailable (plugin absent) — silent skip
        pass

    return diagnostics


def _build_js(spec: BuildSpec, out: Path) -> list[DiagnosticEntry]:
    """Generate a .mjs builder script and invoke node to produce the output .docx."""
    from build_js_codegen import emit_builder
    from node_bridge import ensure_docx_package, ensure_node, run_node_script

    ensure_node()
    ensure_docx_package()

    out = Path(out).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        mjs_path = tmp_path / "build.mjs"
        spec_path = tmp_path / "spec.json"
        emit_builder(spec, mjs_path, spec_path, out)

        result = run_node_script(mjs_path)
        if result.returncode != 0:
            return [DiagnosticEntry(
                level="error",
                source="build_js",
                message=f"node exited {result.returncode}: {result.stderr or result.stdout}",
            )]

    # Validate the output file produced by node.
    if out.exists():
        return _validate_docx(out)
    return [DiagnosticEntry(
        level="error",
        source="build_js",
        message=f"node exited 0 but output file was not created: {out}",
    )]


def _fix_zoom_percent(docx_path: Path) -> bool:
    """Add ``w:percent="100"`` to ``<w:zoom>`` in word/settings.xml if missing.

    python-docx writes ``<w:zoom w:val="bestFit"/>`` without the required
    ``w:percent`` attribute, so Anthropic's schema validator (and Word's
    own validator) reject every python-routed build. This helper repairs
    settings.xml in place after the document has been written.

    Returns True when the file was modified, False otherwise.
    """
    from lxml import etree

    with zipfile.ZipFile(str(docx_path), "r") as zin:
        members = {n: zin.read(n) for n in zin.namelist()}
    settings_xml = members.get("word/settings.xml")
    if settings_xml is None:
        return False
    try:
        tree = etree.fromstring(settings_xml)
    except etree.XMLSyntaxError:
        return False
    changed = False
    for zoom in tree.iter(f"{{{W_NS}}}zoom"):
        if zoom.get(f"{{{W_NS}}}percent") is None:
            zoom.set(f"{{{W_NS}}}percent", "100")
            changed = True
    if not changed:
        return False
    members["word/settings.xml"] = etree.tostring(
        tree, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=docx_path.parent)
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in members.items():
                zout.writestr(name, data)
        Path(tmp_path).replace(docx_path)
    except Exception:
        Path(tmp_path).unlink(missing_ok=True)
        raise
    return True


def _build_python(
    spec: BuildSpec,
    output_path: Path,
    template_path: Path | None = None,
) -> list[DiagnosticEntry]:
    """Execute the python-docx / docxtpl build path."""
    output_path = Path(output_path).resolve()

    if template_path is not None:
        template_path = Path(template_path).resolve()
        diagnostics = _build_from_template(spec, output_path, template_path)
    else:
        diagnostics = _build_programmatic(spec, output_path)

    # Patch the python-docx zoom-percent omission so the file passes
    # Anthropic's schema validator (codex audit issue 5). Run BEFORE
    # footnote injection so subsequent steps see the corrected zoom.
    if output_path.exists():
        try:
            _fix_zoom_percent(output_path)
        except Exception as exc:
            diagnostics.append(DiagnosticEntry(
                level="warning", source="build_docx",
                message=f"Failed to patch zoom percent in settings.xml: {exc}",
            ))

    # Inject footnotes if the spec defines any and the file was created.
    if output_path.exists() and spec.footnotes:
        diagnostics.extend(_inject_footnotes(output_path, spec.footnotes))

    # Validate output if it was created.
    if output_path.exists():
        diagnostics.extend(_validate_docx(output_path))

    return diagnostics


def write_build_manifest(
    output_path: Path,
    spec: BuildSpec,
    diagnostics: list[DiagnosticEntry],
    template_path: Path | None = None,
) -> Path:
    """Write a sidecar manifest JSON next to *output_path*.

    The manifest file is named ``<stem>.manifest.json`` and records metadata
    about the build: output file, template used, timestamp, spec title,
    section/item counts, and diagnostics.

    Returns
    -------
    Path
        The path to the written manifest file.
    """
    output_path = Path(output_path).resolve()
    manifest_path = output_path.parent / f"{output_path.stem}.manifest.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)

    manifest = {
        "output_file": str(output_path),
        "template": str(template_path) if template_path else None,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "spec_title": spec.title,
        "sections_count": len(spec.sections),
        "items_count": len(spec.items),
        "diagnostics": [d.model_dump() for d in diagnostics],
    }

    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    logger.info("Wrote manifest to %s", manifest_path)
    return manifest_path


# ── CLI entry point ────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Build a .docx from a JSON spec")
    parser.add_argument("spec", type=Path, help="Path to the JSON spec file")
    parser.add_argument("output", type=Path, help="Output .docx path")
    parser.add_argument(
        "-t", "--template", type=Path, default=None,
        help="Optional .docx template for docxtpl rendering",
    )
    args = parser.parse_args()

    spec_path = args.spec.resolve()
    try:
        spec = BuildSpec.model_validate_json(spec_path.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[ERROR] Invalid spec: {exc}")
        sys.exit(1)

    diags = build_from_spec(spec, args.output.resolve(), args.template)
    manifest = write_build_manifest(
        args.output.resolve(), spec, diags, args.template,
    )

    has_errors = False
    for d in diags:
        print(f"[{d.level.upper()}] {d.source}: {d.message}")
        if d.level == "error":
            has_errors = True
    print(f"Manifest: {manifest}")
    if has_errors:
        sys.exit(1)
