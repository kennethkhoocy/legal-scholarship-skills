"""Test OOXML audit functionality with minimal fake .docx ZIP archives."""

import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))

import pytest

from audit_ooxml import audit_ooxml, validate_docx

# ---------------------------------------------------------------------------
# XML content for fake .docx files
# ---------------------------------------------------------------------------

CONTENT_TYPES_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

DOCUMENT_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Hello World</w:t></w:r></w:p>
    <w:p><w:r><w:t>Second paragraph</w:t></w:r></w:p>
  </w:body>
</w:document>"""

COMMENTS_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="1" w:author="Alice" w:date="2026-01-15T10:00:00Z">
    <w:p><w:r><w:t>First comment</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="2" w:author="Bob" w:date="2026-01-16T12:30:00Z">
    <w:p><w:r><w:t>Second comment</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_fake_docx(
    directory: Path,
    filename: str = "test.docx",
    include_comments: bool = False,
) -> Path:
    """Create a minimal .docx ZIP archive in *directory*."""
    path = directory / filename
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", DOCUMENT_XML)
        if include_comments:
            zf.writestr("word/comments.xml", COMMENTS_XML)
    return path


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.fixture()
def tmp_dir():
    """Yield a temporary directory, cleaned up after the test."""
    with tempfile.TemporaryDirectory() as td:
        yield Path(td)


def test_validate_valid_docx(tmp_dir):
    """validate_docx should return no errors for a well-formed fake .docx."""
    docx_path = _make_fake_docx(tmp_dir)
    errors = validate_docx(docx_path)
    assert errors == [], f"Expected no validation errors, got: {errors}"


def test_validate_invalid_file(tmp_dir):
    """validate_docx should return errors for a plain-text file."""
    bad_path = tmp_dir / "not_a_docx.docx"
    bad_path.write_text("this is not a zip file", encoding="utf-8")
    errors = validate_docx(bad_path)
    assert len(errors) > 0, "Expected at least one validation error for a non-ZIP file"


def test_audit_counts_paragraphs(tmp_dir):
    """audit_ooxml should count 2 w:p elements in the fake document."""
    docx_path = _make_fake_docx(tmp_dir)
    result = audit_ooxml(docx_path)
    assert result["valid"] is True
    doc_counts = result["counts"].get("word/document.xml", {})
    assert doc_counts.get("w:p") == 2


def test_audit_with_comments(tmp_dir):
    """audit_ooxml should count comment elements when word/comments.xml is present."""
    docx_path = _make_fake_docx(tmp_dir, include_comments=True)
    result = audit_ooxml(docx_path)
    assert result["valid"] is True
    comment_counts = result["counts"].get("word/comments.xml", {})
    assert comment_counts.get("w:comment") == 2
    doc_counts = result["counts"].get("word/document.xml", {})
    assert doc_counts.get("w:p") == 2


def test_audit_malformed_xml(tmp_dir):
    """audit_ooxml should report parse errors for malformed XML."""
    path = tmp_dir / "bad.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", "<invalid>not closed")
    result = audit_ooxml(path)
    assert result["valid"] is False
    doc_entry = result["counts"].get("word/document.xml", {})
    assert "_parse_error" in doc_entry


TRACKED_CHANGES_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:del w:author="Alice" w:date="2026-01-15T10:00:00Z">
        <w:r><w:delText>old text</w:delText></w:r>
      </w:del>
      <w:ins w:author="Alice" w:date="2026-01-15T10:00:00Z">
        <w:r><w:t>new text</w:t></w:r>
      </w:ins>
    </w:p>
  </w:body>
</w:document>"""


def test_revision_extraction_ooxml_order(tmp_dir):
    """OOXML fallback should extract revisions in document order (del before ins)."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "tracked.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", TRACKED_CHANGES_XML)

    revisions, diags = _extract_with_ooxml(path)
    assert len(revisions) == 2
    assert revisions[0].type == "deletion"
    assert revisions[0].text == "old text"
    assert revisions[1].type == "insertion"
    assert revisions[1].text == "new text"


def test_extract_text_includes_tracked_changes(tmp_dir):
    """extract_text should include text from w:ins and w:del elements."""
    from extract_text import extract_text

    path = tmp_dir / "tracked.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Normal paragraph")
    doc.save(str(path))

    with zipfile.ZipFile(path, "r") as zf_in:
        parts = {n: zf_in.read(n) for n in zf_in.namelist()}
    parts["word/document.xml"] = TRACKED_CHANGES_XML.encode("utf-8")
    with zipfile.ZipFile(path, "w") as zf_out:
        for name, data in parts.items():
            zf_out.writestr(name, data)

    paragraphs, markdown, diags = extract_text(path)
    all_text = " ".join(p.text for p in paragraphs)
    assert "old text" in all_text or "new text" in all_text


SPLIT_RUNS_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r><w:t>hel</w:t></w:r>
      <w:r><w:t>lo</w:t></w:r>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_text_split_runs_no_spurious_spaces(tmp_dir):
    """Split runs should be concatenated without inserting spaces."""
    from extract_text import extract_text

    path = tmp_dir / "split.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(path))

    with zipfile.ZipFile(path, "r") as zf_in:
        parts = {n: zf_in.read(n) for n in zf_in.namelist()}
    parts["word/document.xml"] = SPLIT_RUNS_XML.encode("utf-8")
    with zipfile.ZipFile(path, "w") as zf_out:
        for name, data in parts.items():
            zf_out.writestr(name, data)

    paragraphs, markdown, diags = extract_text(path)
    assert any(p.text == "hello" for p in paragraphs), (
        f"Expected 'hello' without spaces, got: {[p.text for p in paragraphs]}"
    )


def test_build_manifest_parent_dir_created(tmp_dir):
    """write_build_manifest should create parent directory if missing."""
    from build_docx import build_from_spec, write_build_manifest
    from models import BuildSpec

    spec = BuildSpec(title="Test", sections=[], items=[])
    nested_out = tmp_dir / "deep" / "nested" / "output.docx"
    diags = build_from_spec(spec, nested_out)
    manifest_path = write_build_manifest(nested_out, spec, diags)
    assert manifest_path.exists()


def test_pipe_escaping_in_tables(tmp_dir):
    """Table cells containing pipe characters should be escaped in Markdown."""
    from extract_text import _escape_pipe
    assert _escape_pipe("a|b") == "a\\|b"
    assert _escape_pipe("no pipes") == "no pipes"


def test_comment_model_has_ooxml_id():
    """Comment model should have ooxml_id field."""
    from models import Comment
    c = Comment(
        comment_id="C001", author="Alice", date="2026-01-01",
        reference_text="ref", comment_text="text",
    )
    assert c.ooxml_id is None
    c.ooxml_id = "42"
    assert c.ooxml_id == "42"


def test_apply_edits_replace(tmp_dir):
    """apply_edits should create tracked changes in the output."""
    from apply_edits import apply_edits
    from models import EditOperation, DiagnosticEntry
    from docx import Document

    # Create a simple docx
    doc = Document()
    doc.add_paragraph("First paragraph with old text here.")
    doc.add_paragraph("Second paragraph unchanged.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="replace",
        paragraph_index=0,
        old_text="old text",
        new_text="new text",
        author="TestBot",
        date="2026-05-12T00:00:00Z",
    )]

    diags = apply_edits(input_path, ops, output_path)
    assert output_path.exists()
    assert not any(d.level == "error" for d in diags)

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Find the first body paragraph
    body = tree.find(f".//{{{ns}}}body")
    para = body.findall(f"{{{ns}}}p")[0]
    children = list(para)

    # Verify structure: before-run, w:del, w:ins, after-run (in that order)
    child_tags = [etree.QName(c.tag).localname for c in children if c.tag != f"{{{ns}}}pPr"]
    assert "del" in child_tags, f"Expected w:del in paragraph children, got {child_tags}"
    assert "ins" in child_tags, f"Expected w:ins in paragraph children, got {child_tags}"
    del_idx = child_tags.index("del")
    ins_idx = child_tags.index("ins")
    assert del_idx < ins_idx, "w:del should come before w:ins"

    # Verify the deleted and inserted text
    del_el = [c for c in children if etree.QName(c.tag).localname == "del"][0]
    ins_el = [c for c in children if etree.QName(c.tag).localname == "ins"][0]
    dt = del_el.find(f".//{{{ns}}}delText")
    it = ins_el.find(f".//{{{ns}}}t")
    assert dt is not None and dt.text == "old text"
    assert it is not None and it.text == "new text"

    # Verify surrounding text is preserved (before: "First paragraph with ", after: " here.")
    runs = [c for c in children if etree.QName(c.tag).localname == "r"]
    run_texts = []
    for r in runs:
        for t in r.findall(f"{{{ns}}}t"):
            if t.text:
                run_texts.append(t.text)
    full_preserved = "".join(run_texts)
    assert "First paragraph with " in full_preserved or full_preserved.startswith("First")
    assert " here." in full_preserved or full_preserved.endswith("here.")


def test_apply_edits_insert(tmp_dir):
    """apply_edits insert should create a tracked insertion."""
    from apply_edits import apply_edits
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Existing paragraph.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="insert",
        paragraph_index=0,
        new_text=" Appended text.",
        author="InsertBot",
        date="2026-05-12T00:00:00Z",
    )]
    diags = apply_edits(input_path, ops, output_path)
    assert output_path.exists()

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ins_elements = list(tree.iter(f"{{{ns}}}ins"))
    assert len(ins_elements) > 0, "Expected w:ins elements for insert op"
    # Verify the inserted text
    t_el = ins_elements[0].find(f".//{{{ns}}}t")
    assert t_el is not None and t_el.text == " Appended text."


def test_apply_edits_delete(tmp_dir):
    """apply_edits delete should create a tracked deletion."""
    from apply_edits import apply_edits
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Remove this word please.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="delete",
        paragraph_index=0,
        old_text="this word",
        author="DeleteBot",
        date="2026-05-12T00:00:00Z",
    )]
    diags = apply_edits(input_path, ops, output_path)
    assert output_path.exists()

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    del_elements = list(tree.iter(f"{{{ns}}}del"))
    assert len(del_elements) > 0, "Expected w:del elements for delete op"
    dt_el = del_elements[0].find(f".//{{{ns}}}delText")
    assert dt_el is not None and dt_el.text == "this word"


def test_apply_edits_paragraph_index_out_of_range(tmp_dir):
    """apply_edits should warn when paragraph_index is out of range."""
    from apply_edits import apply_edits
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Only paragraph.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="replace",
        paragraph_index=99,
        old_text="Only",
        new_text="First",
        author="Bot",
    )]
    diags = apply_edits(input_path, ops, output_path)
    assert any("out of range" in d.message for d in diags)


def test_apply_edits_old_text_not_found(tmp_dir):
    """apply_edits should warn when old_text is not found in the target paragraph."""
    from apply_edits import apply_edits
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Some text here.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="replace",
        paragraph_index=0,
        old_text="nonexistent phrase",
        new_text="replacement",
        author="Bot",
    )]
    diags = apply_edits(input_path, ops, output_path)
    assert any("not found" in d.message for d in diags)


def test_apply_edits_author_and_date_metadata(tmp_dir):
    """Tracked change elements should carry correct author and date attributes."""
    from apply_edits import apply_edits
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Check metadata here.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="replace",
        paragraph_index=0,
        old_text="metadata",
        new_text="attributes",
        author="MetaBot",
        date="2026-03-15T14:30:00Z",
    )]
    apply_edits(input_path, ops, output_path)

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    del_el = list(tree.iter(f"{{{ns}}}del"))[0]
    assert del_el.get(f"{{{ns}}}author") == "MetaBot"
    assert del_el.get(f"{{{ns}}}date") == "2026-03-15T14:30:00Z"

    ins_el = list(tree.iter(f"{{{ns}}}ins"))[0]
    assert ins_el.get(f"{{{ns}}}author") == "MetaBot"
    assert ins_el.get(f"{{{ns}}}date") == "2026-03-15T14:30:00Z"


# ---------------------------------------------------------------------------
# Silent (non-tracked) edits
# ---------------------------------------------------------------------------


def test_silent_edits_replace(tmp_dir):
    """apply_edits_silent replace should change text with NO w:ins or w:del."""
    from apply_edits import apply_edits_silent
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph with old text here.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="replace",
        paragraph_index=0,
        old_text="old text",
        new_text="new text",
        author="TestBot",
        date="2026-05-12T00:00:00Z",
    )]

    diags = apply_edits_silent(input_path, ops, output_path)
    assert output_path.exists()
    assert not any(d.level == "error" for d in diags)

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # No tracked-change elements should be present
    assert list(tree.iter(f"{{{ns}}}del")) == [], "w:del should not exist in silent edit"
    assert list(tree.iter(f"{{{ns}}}ins")) == [], "w:ins should not exist in silent edit"

    # The replacement text should appear in the paragraph's runs
    body = tree.find(f".//{{{ns}}}body")
    para = body.findall(f"{{{ns}}}p")[0]
    run_texts = []
    for r in para.iter(f"{{{ns}}}t"):
        if r.text:
            run_texts.append(r.text)
    full = "".join(run_texts)
    assert "new text" in full, f"Expected 'new text' in paragraph, got: {full}"
    assert "old text" not in full, f"'old text' should have been replaced, got: {full}"


def test_silent_edits_insert(tmp_dir):
    """apply_edits_silent insert should add text with NO w:ins."""
    from apply_edits import apply_edits_silent
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Existing paragraph.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="insert",
        paragraph_index=0,
        new_text=" Appended text.",
        author="InsertBot",
        date="2026-05-12T00:00:00Z",
    )]

    diags = apply_edits_silent(input_path, ops, output_path)
    assert output_path.exists()
    assert not any(d.level == "error" for d in diags)

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # No tracked-change elements
    assert list(tree.iter(f"{{{ns}}}ins")) == [], "w:ins should not exist in silent insert"

    # The appended text should be present in a plain run
    body = tree.find(f".//{{{ns}}}body")
    para = body.findall(f"{{{ns}}}p")[0]
    run_texts = []
    for r in para.iter(f"{{{ns}}}t"):
        if r.text:
            run_texts.append(r.text)
    full = "".join(run_texts)
    assert " Appended text." in full, f"Expected appended text in paragraph, got: {full}"


def test_silent_edits_delete(tmp_dir):
    """apply_edits_silent delete should remove text with NO w:del."""
    from apply_edits import apply_edits_silent
    from models import EditOperation

    from docx import Document
    doc = Document()
    doc.add_paragraph("Remove this word please.")
    input_path = tmp_dir / "input.docx"
    doc.save(str(input_path))

    output_path = tmp_dir / "output.docx"
    ops = [EditOperation(
        operation="delete",
        paragraph_index=0,
        old_text="this word",
        author="DeleteBot",
        date="2026-05-12T00:00:00Z",
    )]

    diags = apply_edits_silent(input_path, ops, output_path)
    assert output_path.exists()
    assert not any(d.level == "error" for d in diags)

    from lxml import etree
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    tree = etree.fromstring(xml)
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # No tracked-change elements
    assert list(tree.iter(f"{{{ns}}}del")) == [], "w:del should not exist in silent delete"

    # The deleted text should be gone
    body = tree.find(f".//{{{ns}}}body")
    para = body.findall(f"{{{ns}}}p")[0]
    run_texts = []
    for r in para.iter(f"{{{ns}}}t"):
        if r.text:
            run_texts.append(r.text)
    full = "".join(run_texts)
    assert "this word" not in full, f"'this word' should have been removed, got: {full}"
    assert "Remove " in full, f"Text before deletion should remain, got: {full}"
    assert " please." in full, f"Text after deletion should remain, got: {full}"


# ---------------------------------------------------------------------------
# Move revision extraction tests
# ---------------------------------------------------------------------------

MOVE_REVISIONS_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:moveFrom w:id="10" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>moved text</w:t></w:r>
      </w:moveFrom>
    </w:p>
    <w:p>
      <w:moveTo w:id="11" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>moved text</w:t></w:r>
      </w:moveTo>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_move_revisions(tmp_dir):
    """OOXML extractor should detect moveFrom and moveTo revisions."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "moves.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", MOVE_REVISIONS_XML)

    revisions, diags = _extract_with_ooxml(path)
    assert len(revisions) == 2
    assert revisions[0].type == "move_from"
    assert revisions[0].text == "moved text"
    assert revisions[0].author == "Alice"
    assert revisions[0].paragraph_index == 0
    assert revisions[1].type == "move_to"
    assert revisions[1].text == "moved text"
    assert revisions[1].paragraph_index == 1


# ---------------------------------------------------------------------------
# Paragraph property change extraction tests
# ---------------------------------------------------------------------------

PPR_CHANGE_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:pPr>
        <w:jc w:val="center"/>
        <w:pPrChange w:id="20" w:author="Bob" w:date="2026-05-20T11:00:00Z">
          <w:pPr/>
        </w:pPrChange>
      </w:pPr>
      <w:r><w:t>Centered paragraph</w:t></w:r>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_paragraph_property_change(tmp_dir):
    """OOXML extractor should detect pPrChange revisions."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "ppr.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", PPR_CHANGE_XML)

    revisions, diags = _extract_with_ooxml(path)
    ppr_revs = [r for r in revisions if r.type == "paragraph_property_change"]
    assert len(ppr_revs) == 1
    assert ppr_revs[0].author == "Bob"
    assert "default" in ppr_revs[0].text
    assert "center" in ppr_revs[0].text
    assert ppr_revs[0].location == "pPr"


# ---------------------------------------------------------------------------
# Run property change extraction tests
# ---------------------------------------------------------------------------

RPR_CHANGE_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r>
        <w:rPr>
          <w:b/>
          <w:rPrChange w:id="30" w:author="Carol" w:date="2026-05-20T12:00:00Z">
            <w:rPr/>
          </w:rPrChange>
        </w:rPr>
        <w:t>bold text</w:t>
      </w:r>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_run_property_change(tmp_dir):
    """OOXML extractor should detect rPrChange revisions."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "rpr.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", RPR_CHANGE_XML)

    revisions, diags = _extract_with_ooxml(path)
    rpr_revs = [r for r in revisions if r.type == "run_property_change"]
    assert len(rpr_revs) == 1
    assert rpr_revs[0].author == "Carol"
    assert "bold text" in rpr_revs[0].text
    assert "b" in rpr_revs[0].text
    assert rpr_revs[0].location == "rPr"


# ---------------------------------------------------------------------------
# Paragraph mark deletion extraction tests
# ---------------------------------------------------------------------------

PARA_MARK_DEL_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:pPr>
        <w:rPr>
          <w:del w:id="40" w:author="Dave" w:date="2026-05-20T13:00:00Z"/>
        </w:rPr>
      </w:pPr>
      <w:r><w:t>First paragraph</w:t></w:r>
    </w:p>
    <w:p>
      <w:r><w:t>Second paragraph</w:t></w:r>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_paragraph_mark_deletion(tmp_dir):
    """OOXML extractor should detect paragraph mark deletions separately from inline dels."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "pmark.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", PARA_MARK_DEL_XML)

    revisions, diags = _extract_with_ooxml(path)
    pm_revs = [r for r in revisions if r.type == "paragraph_mark_deletion"]
    assert len(pm_revs) == 1
    assert pm_revs[0].author == "Dave"
    assert pm_revs[0].text == "¶"
    assert pm_revs[0].location == "pPr/rPr"
    assert pm_revs[0].paragraph_index == 0
    # Should NOT also produce an inline deletion
    inline_dels = [r for r in revisions if r.type == "deletion"]
    assert len(inline_dels) == 0


# ---------------------------------------------------------------------------
# Revision markdown renderer tests
# ---------------------------------------------------------------------------


def test_revisions_markdown_new_types():
    """revisions_to_markdown should render all new revision types."""
    from extract_revisions import revisions_to_markdown
    from models import Revision

    revisions = [
        Revision(revision_id="R001", type="move_from", author="A", date="d",
                 text="moved", paragraph_index=0, source="ooxml"),
        Revision(revision_id="R002", type="move_to", author="A", date="d",
                 text="moved", paragraph_index=1, source="ooxml"),
        Revision(revision_id="R003", type="paragraph_property_change", author="B",
                 date="d", text="from [default] to [align=center]",
                 paragraph_index=0, location="pPr", source="ooxml"),
        Revision(revision_id="R004", type="run_property_change", author="C",
                 date="d", text="'word' from [default] to [b]",
                 paragraph_index=0, location="rPr", source="ooxml"),
        Revision(revision_id="R005", type="paragraph_mark_deletion", author="D",
                 date="d", text="¶", paragraph_index=0, location="pPr/rPr",
                 source="ooxml"),
    ]
    md = revisions_to_markdown(revisions)
    assert "Move Source" in md
    assert "Move Destination" in md
    assert "Paragraph Property Change" in md
    assert "Run Property Change" in md
    assert "Paragraph Mark Deletion" in md
    assert "**Location:**" in md


# ---------------------------------------------------------------------------
# Combined revisions test (all types in one document)
# ---------------------------------------------------------------------------

COMBINED_REVISIONS_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:pPr>
        <w:jc w:val="right"/>
        <w:pPrChange w:id="1" w:author="Alice" w:date="2026-05-20T10:00:00Z">
          <w:pPr><w:jc w:val="left"/></w:pPr>
        </w:pPrChange>
        <w:rPr>
          <w:del w:id="2" w:author="Alice" w:date="2026-05-20T10:00:00Z"/>
        </w:rPr>
      </w:pPr>
      <w:ins w:id="3" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>inserted</w:t></w:r>
      </w:ins>
      <w:del w:id="4" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:delText>deleted</w:delText></w:r>
      </w:del>
      <w:r>
        <w:rPr>
          <w:i/>
          <w:rPrChange w:id="5" w:author="Alice" w:date="2026-05-20T10:00:00Z">
            <w:rPr/>
          </w:rPrChange>
        </w:rPr>
        <w:t>italic</w:t>
      </w:r>
      <w:moveFrom w:id="6" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>moved away</w:t></w:r>
      </w:moveFrom>
      <w:moveTo w:id="7" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>moved here</w:t></w:r>
      </w:moveTo>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_combined_revisions(tmp_dir):
    """OOXML extractor should extract all revision types from a single document."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "combined.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", COMBINED_REVISIONS_XML)

    revisions, diags = _extract_with_ooxml(path)
    types = [r.type for r in revisions]
    assert "paragraph_mark_deletion" in types
    assert "paragraph_property_change" in types
    assert "insertion" in types
    assert "deletion" in types
    assert "run_property_change" in types
    assert "move_from" in types
    assert "move_to" in types
    assert len(revisions) == 7


# ---------------------------------------------------------------------------
# rPrChange under pPr/rPr (paragraph-mark formatting change)
# ---------------------------------------------------------------------------

PPR_RPR_CHANGE_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:pPr>
        <w:rPr>
          <w:b/>
          <w:rPrChange w:id="50" w:author="Eve" w:date="2026-05-20T14:00:00Z">
            <w:rPr/>
          </w:rPrChange>
        </w:rPr>
      </w:pPr>
      <w:r><w:t>Normal text</w:t></w:r>
    </w:p>
  </w:body>
</w:document>"""


def test_extract_rpr_change_under_ppr(tmp_dir):
    """rPrChange under pPr/rPr should report paragraph-mark scope, not run scope."""
    from extract_revisions import _extract_with_ooxml

    path = tmp_dir / "ppr_rpr.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES_XML)
        zf.writestr("word/document.xml", PPR_RPR_CHANGE_XML)

    revisions, diags = _extract_with_ooxml(path)
    rpr_revs = [r for r in revisions if r.type == "run_property_change"]
    assert len(rpr_revs) == 1
    assert rpr_revs[0].location == "pPr/rPr"
    assert "¶" in rpr_revs[0].text
    assert rpr_revs[0].author == "Eve"


# ---------------------------------------------------------------------------
# Supplement path: docx-revisions primary + OOXML-only types merged
# ---------------------------------------------------------------------------


def test_extract_revisions_supplements_ooxml_only_types(tmp_dir):
    """extract_revisions() should supplement docx-revisions output with OOXML-only types."""
    from extract_revisions import extract_revisions

    # Build a docx that has both an insertion (docx-revisions can handle) and
    # a pPrChange (OOXML-only). The public API should return both.
    xml = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:pPr>
        <w:jc w:val="center"/>
        <w:pPrChange w:id="1" w:author="Alice" w:date="2026-05-20T10:00:00Z">
          <w:pPr/>
        </w:pPrChange>
      </w:pPr>
      <w:ins w:id="2" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>new text</w:t></w:r>
      </w:ins>
      <w:moveFrom w:id="3" w:author="Alice" w:date="2026-05-20T10:00:00Z">
        <w:r><w:t>moved</w:t></w:r>
      </w:moveFrom>
    </w:p>
  </w:body>
</w:document>"""

    from docx import Document
    doc = Document()
    doc.add_paragraph("placeholder")
    path = tmp_dir / "supplement.docx"
    doc.save(str(path))

    with zipfile.ZipFile(path, "r") as zf_in:
        parts = {n: zf_in.read(n) for n in zf_in.namelist()}
    parts["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(path, "w") as zf_out:
        for name, data in parts.items():
            zf_out.writestr(name, data)

    revisions, diags = extract_revisions(path)
    types = {r.type for r in revisions}
    assert "paragraph_property_change" in types, f"Expected pPrChange supplement, got types: {types}"
    assert "move_from" in types, f"Expected moveFrom supplement, got types: {types}"
