"""Generate the three reference .docx fixtures from scratch.

Run manually with `python tests/fixtures/build.py` after a fresh checkout,
or let `tests/conftest.py` auto-invoke this on first test run.
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

# Reuse the shared post-save zoom-percent patcher so fixtures and the
# build pipeline cannot diverge (codex audit issue 5).
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent / "scripts"))
from build_docx import _fix_zoom_percent  # noqa: E402

_DIR = Path(__file__).resolve().parent
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W_NS}


def build_clean() -> None:
    doc = Document()
    doc.add_heading("Clean Fixture", level=1)
    doc.add_paragraph("First paragraph of the clean fixture document.")
    doc.add_paragraph("Second paragraph — no comments, no revisions.")
    doc.add_paragraph("Third paragraph for testing paragraph_index lookups.")
    doc.save(_DIR / "clean.docx")
    # python-docx emits <w:zoom w:val="bestFit"/> without the required
    # w:percent attribute; patch it so the schema validator accepts the
    # fixture.
    _fix_zoom_percent(_DIR / "clean.docx")


def build_with_comments() -> None:
    """Two threaded comments — C0 with a child reply C1; C0 marked done."""
    doc = Document()
    doc.add_heading("Comments Fixture", level=1)
    doc.add_paragraph("Anchor paragraph zero for the first comment.")
    doc.add_paragraph("Anchor paragraph one for a separate comment thread.")
    doc.save(_DIR / "with_comments.docx")
    _inject_comments(_DIR / "with_comments.docx")


def build_with_revisions() -> None:
    """Insertion by author 'Alice' and deletion by author 'Bob' across two paragraphs."""
    doc = Document()
    doc.add_heading("Revisions Fixture", level=1)
    doc.add_paragraph("Paragraph zero is unchanged.")
    doc.add_paragraph("Paragraph one will gain a tracked insertion.")
    doc.add_paragraph("Paragraph two has text scheduled for tracked deletion here.")
    doc.save(_DIR / "with_revisions.docx")
    _inject_revisions(_DIR / "with_revisions.docx")


def _inject_comments(path: Path) -> None:
    """Inject a parent comment C0 + a reply C1 + a separate comment C2.

    Manipulates word/document.xml, word/comments.xml, word/commentsExtended.xml,
    and updates [Content_Types].xml and word/_rels/document.xml.rels.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp) / "out.docx"
        shutil.copy(path, tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zin:
            members = {n: zin.read(n) for n in zin.namelist()}

        # 1. document.xml — add comment range markers around runs
        doc_tree = etree.fromstring(members["word/document.xml"])
        body = doc_tree.find(f"{{{W_NS}}}body")
        paragraphs = body.findall(f"{{{W_NS}}}p")
        # paragraphs[0] is the heading; [1] is "Anchor paragraph zero...";
        # [2] is "Anchor paragraph one..."
        _wrap_with_comment(paragraphs[1], comment_id="0")
        _wrap_with_comment(paragraphs[1], comment_id="1", nest_inside="0")
        _wrap_with_comment(paragraphs[2], comment_id="2")
        members["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

        # 2. comments.xml
        members["word/comments.xml"] = _build_comments_xml([
            ("0", "Alice", "2026-05-01T12:00:00Z", "Parent comment text."),
            ("1", "Bob",   "2026-05-01T12:05:00Z", "Reply to parent."),
            ("2", "Alice", "2026-05-01T12:10:00Z", "Unrelated comment."),
        ])

        # 3. commentsExtended.xml — parent/child + done attribute
        members["word/commentsExtended.xml"] = _build_comments_extended_xml([
            {"paraId": "00000001", "doneVal": "1"},                      # comment 0 (done)
            {"paraId": "00000002", "doneVal": "0", "parentParaId": "00000001"},  # reply
            {"paraId": "00000003", "doneVal": "0"},                      # comment 2
        ])

        # 4. [Content_Types].xml — declare comment parts
        members["[Content_Types].xml"] = _ensure_content_types(members["[Content_Types].xml"])

        # 5. word/_rels/document.xml.rels — link the comment parts
        members["word/_rels/document.xml.rels"] = _ensure_rels(members["word/_rels/document.xml.rels"])

        # Rewrite the .docx
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in members.items():
                zout.writestr(name, data)


def _inject_revisions(path: Path) -> None:
    """Inject a tracked insertion in paragraph 1 and a tracked deletion in paragraph 2."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp) / "out.docx"
        shutil.copy(path, tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zin:
            members = {n: zin.read(n) for n in zin.namelist()}

        doc_tree = etree.fromstring(members["word/document.xml"])
        body = doc_tree.find(f"{{{W_NS}}}body")
        paragraphs = body.findall(f"{{{W_NS}}}p")
        # paragraphs[0] heading; [1] is "Paragraph zero...";
        # [2] gets the tracked insertion; [3] gets the tracked deletion
        _add_tracked_insertion(paragraphs[2], "INSERTED TEXT", author="Alice", date="2026-05-01T13:00:00Z", w_id="100")
        _add_tracked_deletion(paragraphs[3], "scheduled for tracked deletion ", author="Bob", date="2026-05-01T13:05:00Z", w_id="101")

        members["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in members.items():
                zout.writestr(name, data)


def _wrap_with_comment(paragraph, comment_id: str, nest_inside: str | None = None) -> None:
    """Insert <w:commentRangeStart>, <w:commentRangeEnd>, and the reference run on a paragraph.

    The range markers bracket the original text run(s) of the paragraph so the comment
    has a non-empty anchor span. If nest_inside is given, place this comment's range
    markers strictly inside the parent's range markers, and place this reference run
    after the parent's reference run (per Anthropic SKILL.md reply nesting rule).
    """
    W = f"{{{W_NS}}}"
    original_runs = [child for child in paragraph if child.tag == f"{W}r"]
    start = etree.Element(f"{W}commentRangeStart")
    start.set(f"{W}id", comment_id)
    end = etree.Element(f"{W}commentRangeEnd")
    end.set(f"{W}id", comment_id)
    ref_run = etree.Element(f"{W}r")
    rpr = etree.SubElement(ref_run, f"{W}rPr")
    rstyle = etree.SubElement(rpr, f"{W}rStyle")
    rstyle.set(f"{W}val", "CommentReference")
    cref = etree.SubElement(ref_run, f"{W}commentReference")
    cref.set(f"{W}id", comment_id)

    if nest_inside is None:
        if original_runs:
            original_runs[0].addprevious(start)
            original_runs[-1].addnext(end)
            end.addnext(ref_run)
        else:
            paragraph.append(start)
            paragraph.append(end)
            paragraph.append(ref_run)
    else:
        # Find parent's commentRangeStart and commentRangeEnd; place inside them.
        parent_start = None
        parent_end = None
        for el in paragraph.iter():
            tag = etree.QName(el).localname
            wid = el.get(f"{W}id")
            if tag == "commentRangeStart" and wid == nest_inside:
                parent_start = el
            elif tag == "commentRangeEnd" and wid == nest_inside:
                parent_end = el
        if parent_start is not None and parent_end is not None:
            parent_start.addnext(start)
            parent_end.addprevious(end)
            # Reply's reference run goes after the parent comment's reference run.
            parent_ref_run = None
            for r in paragraph.iter(f"{W}r"):
                cr = r.find(f"{W}commentReference")
                if cr is not None and cr.get(f"{W}id") == nest_inside:
                    parent_ref_run = r
                    break
            if parent_ref_run is not None:
                parent_ref_run.addnext(ref_run)
            else:
                paragraph.append(ref_run)


def _add_tracked_insertion(paragraph, text: str, author: str, date: str, w_id: str) -> None:
    ins = etree.SubElement(paragraph, f"{{{W_NS}}}ins")
    ins.set(f"{{{W_NS}}}id", w_id)
    ins.set(f"{{{W_NS}}}author", author)
    ins.set(f"{{{W_NS}}}date", date)
    r = etree.SubElement(ins, f"{{{W_NS}}}r")
    t = etree.SubElement(r, f"{{{W_NS}}}t")
    t.text = text


def _add_tracked_deletion(paragraph, text: str, author: str, date: str, w_id: str) -> None:
    """Wrap a literal substring in paragraph's first matching run with a tracked deletion."""
    runs = paragraph.findall(f"{{{W_NS}}}r")
    for run in runs:
        ts = run.findall(f"{{{W_NS}}}t")
        for t in ts:
            if t.text and text in t.text:
                idx = t.text.index(text)
                before, after = t.text[:idx], t.text[idx + len(text):]
                t.text = before

                # Build the deletion element
                del_el = etree.Element(f"{{{W_NS}}}del")
                del_el.set(f"{{{W_NS}}}id", w_id)
                del_el.set(f"{{{W_NS}}}author", author)
                del_el.set(f"{{{W_NS}}}date", date)
                del_r = etree.SubElement(del_el, f"{{{W_NS}}}r")
                del_t = etree.SubElement(del_r, f"{{{W_NS}}}delText")
                del_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                del_t.text = text

                # Insert the deletion immediately after the truncated run
                run.addnext(del_el)

                # If there's leftover text, place it in a fresh run AFTER the deletion
                if after:
                    new_r = etree.Element(f"{{{W_NS}}}r")
                    new_t = etree.SubElement(new_r, f"{{{W_NS}}}t")
                    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    new_t.text = after
                    del_el.addnext(new_r)
                return


def _build_comments_xml(rows: list[tuple[str, str, str, str]]) -> bytes:
    root = etree.Element(f"{{{W_NS}}}comments", nsmap=NSMAP)
    for cid, author, date, text in rows:
        c = etree.SubElement(root, f"{{{W_NS}}}comment")
        c.set(f"{{{W_NS}}}id", cid)
        c.set(f"{{{W_NS}}}author", author)
        c.set(f"{{{W_NS}}}date", date)
        c.set(f"{{{W_NS}}}initials", "".join(w[0] for w in author.split()))
        p = etree.SubElement(c, f"{{{W_NS}}}p")
        r = etree.SubElement(p, f"{{{W_NS}}}r")
        t = etree.SubElement(r, f"{{{W_NS}}}t")
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_comments_extended_xml(entries: list[dict]) -> bytes:
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    nsmap = {"w15": w15_ns, "w": W_NS}
    root = etree.Element(f"{{{w15_ns}}}commentsEx", nsmap=nsmap)
    for entry in entries:
        cx = etree.SubElement(root, f"{{{w15_ns}}}commentEx")
        cx.set(f"{{{w15_ns}}}paraId", entry["paraId"])
        cx.set(f"{{{w15_ns}}}done", entry.get("doneVal", "0"))
        if "parentParaId" in entry:
            cx.set(f"{{{w15_ns}}}paraIdParent", entry["parentParaId"])
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _ensure_content_types(data: bytes) -> bytes:
    tree = etree.fromstring(data)
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    overrides_present = {ov.get("PartName") for ov in tree.findall(f"{{{ct_ns}}}Override")}
    additions = [
        ("/word/comments.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"),
        ("/word/commentsExtended.xml", "application/vnd.ms-word.commentsExtended+xml"),
    ]
    for part, content_type in additions:
        if part not in overrides_present:
            ov = etree.SubElement(tree, f"{{{ct_ns}}}Override")
            ov.set("PartName", part)
            ov.set("ContentType", content_type)
    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _ensure_rels(data: bytes) -> bytes:
    tree = etree.fromstring(data)
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    rels_present = {r.get("Target") for r in tree.findall(f"{{{rel_ns}}}Relationship")}
    additions = [
        ("comments.xml", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"),
        ("commentsExtended.xml", "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"),
    ]
    next_id = 100
    for target, rel_type in additions:
        if target not in rels_present:
            r = etree.SubElement(tree, f"{{{rel_ns}}}Relationship")
            r.set("Id", f"rId{next_id}")
            r.set("Type", rel_type)
            r.set("Target", target)
            next_id += 1
    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


if __name__ == "__main__":
    build_clean()
    build_with_comments()
    build_with_revisions()
    print(f"Generated 3 fixtures under {_DIR}")
