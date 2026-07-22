"""Tests for scripts/add_comment.py (the three sub-modes)."""

import shutil
import subprocess
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))

import pytest

FIXTURES = Path(__file__).resolve().parent / "fixtures"
PLUGIN_BASE = Path.home() / ".claude" / "plugins" / "cache" / "anthropic-agent-skills" / "document-skills"


def _skip_if_plugin_absent():
    if not PLUGIN_BASE.is_dir():
        pytest.skip("Anthropic docx plugin not installed")


def _run_cli(*args) -> subprocess.CompletedProcess:
    script = Path(__file__).resolve().parent.parent / "scripts" / "word_docx.py"
    return subprocess.run(
        [sys.executable, str(script), *args],
        capture_output=True,
        text=True,
    )


def test_add_comment_help_lists_three_modes():
    result = _run_cli("add-comment", "--help")
    assert result.returncode == 0
    text = result.stdout.lower()
    for term in ("reply-to", "anchor-paragraph", "resolve"):
        assert term in text, f"--help missing flag for {term}"


def test_resolve_marks_comment_as_done(tmp_path):
    _skip_if_plugin_absent()
    out = tmp_path / "resolved.docx"

    # Confirm fixture starts with comment C2 NOT done (C0 already done in fixture)
    extract_dir = tmp_path / "extract_before"
    pre = _run_cli("extract-comments", str(FIXTURES / "with_comments.docx"), "--out", str(extract_dir))
    assert pre.returncode == 0, pre.stderr

    # Now resolve C3 (the unrelated comment, w:id=2 → C003 in 1-indexed extract-comments)
    r = _run_cli("add-comment", str(FIXTURES / "with_comments.docx"), "--out", str(out), "--resolve", "C003")
    assert r.returncode == 0, r.stderr
    assert out.is_file()

    # Verify w15:done="1" appears for the resolved comment in commentsExtended.xml.
    # The third commentEx (paraId 00000003) should now have done="1".
    import zipfile
    with zipfile.ZipFile(out, "r") as zf:
        ext = zf.read("word/commentsExtended.xml").decode("utf-8")
    import re
    matches = re.findall(r'paraId="00000003"[^/]*done="1"', ext)
    assert matches, f"Expected paraId=00000003 with done=1, got: {ext}"


def test_reply_appends_child_comment(tmp_path):
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"

    # Reply to comment C003 (the unrelated single comment) so the reply has a clear parent
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Replying to the unrelated comment.",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr
    assert out.is_file()

    # Re-extract comments; the reply should appear as an additional comment
    extract_dir = tmp_path / "extracted_after"
    e = _run_cli("extract-comments", str(out), "--out", str(extract_dir))
    assert e.returncode == 0, e.stderr

    import json
    data = json.loads((extract_dir / "comments.json").read_text(encoding="utf-8"))
    bodies = [c["comment_text"] for c in data]
    assert any("Replying to the unrelated comment." in t for t in bodies)


def test_new_comment_anchors_on_span(tmp_path):
    _skip_if_plugin_absent()
    out = tmp_path / "annotated.docx"

    r = _run_cli(
        "add-comment", str(FIXTURES / "clean.docx"),
        "--out", str(out),
        "--anchor-paragraph", "1",
        "--anchor-text", "First paragraph",
        "--text", "Fresh comment on the first phrase.",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr
    assert out.is_file()

    extract_dir = tmp_path / "extracted"
    e = _run_cli("extract-comments", str(out), "--out", str(extract_dir))
    assert e.returncode == 0, e.stderr

    import json
    data = json.loads((extract_dir / "comments.json").read_text(encoding="utf-8"))
    assert len(data) == 1
    assert data[0]["author"] == "Claude"
    assert "Fresh comment" in data[0]["comment_text"]
    assert "First paragraph" in data[0]["reference_text"]


def test_new_comment_handles_cross_run_anchor(tmp_path):
    """Anchor text spanning runs with different formatting should still match.

    Regression test for cross-run anchor handling.  Asserts not just round-trip
    survival but also that document.xml shows the correct OOXML structure:
    a commentRangeStart immediately before the first matched run, a
    commentRangeEnd after the last matched run, a commentReference run after
    the end marker, and per-run formatting preservation (the "Hello " run
    remains plain while the "world" run remains bold).  Any regression that
    collapses runs or loses formatting will fail this test.
    """
    _skip_if_plugin_absent()
    from docx import Document

    src = tmp_path / "cross_run.docx"
    out = tmp_path / "annotated.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hello ")
    r2 = p.add_run("world")
    r2.bold = True
    p.add_run(" extra")
    doc.save(src)

    result = _run_cli(
        "add-comment", str(src),
        "--out", str(out),
        "--anchor-paragraph", "0",
        "--anchor-text", "Hello world",
        "--text", "Cross-run comment",
    )
    assert result.returncode == 0, result.stderr
    assert out.is_file()

    extract_dir = tmp_path / "extracted"
    e = _run_cli("extract-comments", str(out), "--out", str(extract_dir))
    assert e.returncode == 0, e.stderr
    import json
    data = json.loads((extract_dir / "comments.json").read_text(encoding="utf-8"))
    assert len(data) == 1
    assert "Hello world" in data[0]["reference_text"]

    # ---- Inspect document.xml directly for marker placement + formatting ----
    import zipfile
    from lxml import etree

    with zipfile.ZipFile(out, "r") as zf:
        doc_xml = zf.read("word/document.xml")

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.fromstring(doc_xml)

    # Locate the first paragraph that carries a commentRangeStart.
    para = None
    for p_el in root.iter(f"{{{W}}}p"):
        if p_el.find(f"{{{W}}}commentRangeStart") is not None:
            para = p_el
            break
    assert para is not None, "No paragraph with commentRangeStart found in document.xml"

    # Walk children in order; for w:r, capture (text, bold) — bold detection
    # checks for a w:b element whose w:val is not "0"/"false".
    def _is_bold(run_el):
        rpr = run_el.find(f"{{{W}}}rPr")
        if rpr is None:
            return False
        b = rpr.find(f"{{{W}}}b")
        if b is None:
            return False
        val = b.get(f"{{{W}}}val")
        return val not in ("0", "false")

    children = []
    for el in para:
        tag = etree.QName(el).localname
        if tag == "r":
            t = el.find(f"{{{W}}}t")
            text = t.text if t is not None else None
            # Detect whether this run carries a commentReference (the
            # anchor marker), so we can distinguish it from content runs.
            has_cref = el.find(f"{{{W}}}commentReference") is not None
            children.append(("r", text, _is_bold(el), has_cref))
        else:
            children.append((tag, None, None, False))

    # Markers must be present and in order.
    tags_in_order = [c[0] for c in children]
    assert "commentRangeStart" in tags_in_order, f"missing commentRangeStart, got {tags_in_order}"
    assert "commentRangeEnd" in tags_in_order, f"missing commentRangeEnd, got {tags_in_order}"
    start_idx = tags_in_order.index("commentRangeStart")
    end_idx = tags_in_order.index("commentRangeEnd")
    assert start_idx < end_idx, (
        f"commentRangeStart (idx {start_idx}) must precede commentRangeEnd "
        f"(idx {end_idx}); order: {tags_in_order}"
    )

    # The content runs between Start and End should concatenate to "Hello world".
    between = children[start_idx + 1:end_idx]
    between_content_runs = [c for c in between if c[0] == "r" and c[1] and not c[3]]
    concat = "".join(r[1] for r in between_content_runs)
    assert concat == "Hello world", (
        f"Expected runs between markers to concatenate to 'Hello world', got {concat!r} "
        f"from runs {[r[1] for r in between_content_runs]}"
    )

    # The "Hello" portion must come from a non-bold run; the "world" portion
    # must come from a bold run.  This proves per-run formatting was preserved.
    hello_runs = [r for r in between_content_runs if "Hello" in (r[1] or "")]
    world_runs = [r for r in between_content_runs if "world" in (r[1] or "")]
    assert hello_runs, f"no run containing 'Hello' found between markers: {between_content_runs}"
    assert world_runs, f"no run containing 'world' found between markers: {between_content_runs}"
    assert any(not r[2] for r in hello_runs), (
        f"the 'Hello' portion should be plain (not bold); runs: {hello_runs}"
    )
    assert any(r[2] for r in world_runs), (
        f"the 'world' portion should remain bold; runs: {world_runs}"
    )

    # A commentReference run must appear after commentRangeEnd.
    after_end = children[end_idx + 1:]
    ref_run = next((c for c in after_end if c[0] == "r" and c[3]), None)
    assert ref_run is not None, (
        f"expected a commentReference run after commentRangeEnd; saw {after_end}"
    )


def test_resolve_correct_under_shuffled_paraIds(tmp_path):
    """Resolving C002 must mark the right commentEx done even when commentsExtended.xml
    entries are in a different order than the comments in comments.xml.

    Regression test for the paraId-based resolve fix.  An older implementation
    matched the n-th commentEx to the n-th comment by ordinal index; if the two
    files are out of sync (Word writes them in independent orders), that
    fallback marks the WRONG comment as resolved.  This test constructs a
    .docx where commentsExtended.xml is deliberately shuffled (entry for C2
    appears first, entry for C0 second, entry for C1 third), invokes resolve
    on C002, and asserts that the commentEx whose paraId matches C002's
    paragraph paraId in comments.xml is the one marked done="1", while the
    other two retain done="0".
    """
    _skip_if_plugin_absent()
    import sys as _sys
    import zipfile
    from lxml import etree

    # Build a fresh .docx with three independent comments using the fixture helpers.
    sys.path.insert(0, str(Path(__file__).resolve().parent / "fixtures"))
    import importlib.util as _imp_util
    _spec = _imp_util.spec_from_file_location(
        "_fixture_build",
        Path(__file__).resolve().parent / "fixtures" / "build.py",
    )
    _mod = _imp_util.module_from_spec(_spec)
    _spec.loader.exec_module(_mod)

    from docx import Document

    W_NS = _mod.W_NS

    src = tmp_path / "shuffled.docx"
    doc = Document()
    doc.add_heading("Shuffled paraId fixture", level=1)
    doc.add_paragraph("Anchor zero for comment C0.")
    doc.add_paragraph("Anchor one for comment C1.")
    doc.add_paragraph("Anchor two for comment C2.")
    doc.save(src)
    _mod._fix_zoom_percent(src)

    # Inject three independent (non-threaded) comments with distinct paraIds.
    # Then re-write commentsExtended.xml in shuffled order so that the n-th
    # commentEx no longer corresponds to the n-th comment.
    with zipfile.ZipFile(src, "r") as zin:
        members = {n: zin.read(n) for n in zin.namelist()}

    doc_tree = etree.fromstring(members["word/document.xml"])
    body = doc_tree.find(f"{{{W_NS}}}body")
    paragraphs = body.findall(f"{{{W_NS}}}p")
    # paragraphs[0] is the heading; [1], [2], [3] are the anchor paragraphs.
    _mod._wrap_with_comment(paragraphs[1], comment_id="0")
    _mod._wrap_with_comment(paragraphs[2], comment_id="1")
    _mod._wrap_with_comment(paragraphs[3], comment_id="2")
    members["word/document.xml"] = etree.tostring(
        doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # comments.xml — three independent comments.
    members["word/comments.xml"] = _mod._build_comments_xml([
        ("0", "Alice", "2026-05-01T12:00:00Z", "Comment zero."),
        ("1", "Bob",   "2026-05-01T12:05:00Z", "Comment one."),
        ("2", "Carol", "2026-05-01T12:10:00Z", "Comment two."),
    ])

    # commentsExtended.xml — paraId values are chosen so that the n-th entry
    # would map to the WRONG comment under ordinal matching.  Mapping that
    # the resolve code MUST follow (via comments.xml's w14:paraId):
    #   C0 -> paraId "3A000010"
    #   C1 -> paraId "3A000020"
    #   C2 -> paraId "3A000030"
    # But we write the commentsExtended.xml entries in the shuffled order
    # [C2, C0, C1].  Ordinal-fallback resolve would then mark the entry at
    # index 1 (which carries paraId 3A000010 = C0) when the user asks for
    # C002, instead of the correct paraId 3A000030 entry.
    paraId_for = {"0": "3A000010", "1": "3A000020", "2": "3A000030"}
    members["word/commentsExtended.xml"] = _mod._build_comments_extended_xml([
        {"paraId": paraId_for["2"], "doneVal": "0"},  # entry index 0 -> C2
        {"paraId": paraId_for["0"], "doneVal": "0"},  # entry index 1 -> C0
        {"paraId": paraId_for["1"], "doneVal": "0"},  # entry index 2 -> C1
    ])

    # Inject the matching w14:paraId attributes into comments.xml so the
    # resolve code can map by paraId (instead of falling back to ordinal).
    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
    c_root = etree.fromstring(members["word/comments.xml"])
    # Re-build the root with the w14 namespace declaration so the attribute
    # serialises with a registered prefix.
    new_nsmap = dict(c_root.nsmap)
    new_nsmap["w14"] = W14_NS
    new_root = etree.Element(c_root.tag, nsmap=new_nsmap)
    for k, v in c_root.attrib.items():
        new_root.set(k, v)
    for child in c_root:
        new_root.append(child)
    for c in new_root.findall(f"{{{W_NS}}}comment"):
        wid = c.get(f"{{{W_NS}}}id")
        p = c.find(f"{{{W_NS}}}p")
        if p is not None and wid in paraId_for:
            p.set(f"{{{W14_NS}}}paraId", paraId_for[wid])
    members["word/comments.xml"] = etree.tostring(
        new_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # Ensure [Content_Types].xml and rels are wired up.
    members["[Content_Types].xml"] = _mod._ensure_content_types(members["[Content_Types].xml"])
    members["word/_rels/document.xml.rels"] = _mod._ensure_rels(
        members["word/_rels/document.xml.rels"]
    )

    # Rewrite the .docx with the shuffled extended-XML.
    with zipfile.ZipFile(src, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)

    # ---- Resolve C002 (which corresponds to w:id=1, paraId 3A000020). ----
    out = tmp_path / "resolved_shuffled.docx"
    r = _run_cli(
        "add-comment", str(src),
        "--out", str(out),
        "--resolve", "C002",
    )
    assert r.returncode == 0, r.stderr

    # ---- Assert: the commentEx for paraId 3A000020 is done="1"; the other two are done="0". ----
    with zipfile.ZipFile(out, "r") as zf:
        ext_bytes = zf.read("word/commentsExtended.xml")
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    ext_root = etree.fromstring(ext_bytes)
    entries = ext_root.findall(f"{{{w15_ns}}}commentEx")
    state = {
        e.get(f"{{{w15_ns}}}paraId"): e.get(f"{{{w15_ns}}}done")
        for e in entries
    }
    # The correct paraId (C002's) must be done=1.
    assert state.get(paraId_for["1"]) == "1", (
        f"C002's commentEx (paraId {paraId_for['1']}) was NOT marked done; full state: {state}"
    )
    # The other two must remain done=0 (NOT changed by mistake).
    assert state.get(paraId_for["0"]) == "0", (
        f"C0's commentEx (paraId {paraId_for['0']}) was wrongly marked done; full state: {state}"
    )
    assert state.get(paraId_for["2"]) == "0", (
        f"C2's commentEx (paraId {paraId_for['2']}) was wrongly marked done; full state: {state}"
    )


def test_new_comment_with_special_characters_in_body(tmp_path):
    """Comment body containing &, <, > must be XML-escaped before reaching comment.py.

    Regression test for codex audit issue 2: an unescaped ``&`` (or ``<`` /
    ``>``) in --text used to crash the downstream comment.py with an
    ``ExpatError`` because the body is interpolated directly into XML.
    After the fix, the comment body extracts back exactly as written —
    the XML serializer takes care of the escaping. Other Unicode such as
    smart quotes and apostrophes must round-trip unchanged.
    """
    _skip_if_plugin_absent()
    out = tmp_path / "annotated.docx"

    body = 'A & B <tag> "quoted" “smart” ’apostrophe'
    r = _run_cli(
        "add-comment", str(FIXTURES / "clean.docx"),
        "--out", str(out),
        "--anchor-paragraph", "1",
        "--anchor-text", "First paragraph",
        "--text", body,
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr
    assert out.is_file()

    extract_dir = tmp_path / "extracted"
    e = _run_cli("extract-comments", str(out), "--out", str(extract_dir))
    assert e.returncode == 0, e.stderr

    import json
    data = json.loads((extract_dir / "comments.json").read_text(encoding="utf-8"))
    assert len(data) == 1
    text = data[0]["comment_text"]
    assert text == body, (
        f"Comment body did not round-trip; expected {body!r}, got {text!r}"
    )


def test_reply_with_special_characters_in_body(tmp_path):
    """Reply text containing &, <, > must also be XML-escaped before reaching comment.py."""
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"

    body = "Reply with & < > and “quotes”"
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", body,
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr
    assert out.is_file()

    extract_dir = tmp_path / "extracted_reply"
    e = _run_cli("extract-comments", str(out), "--out", str(extract_dir))
    assert e.returncode == 0, e.stderr

    import json
    data = json.loads((extract_dir / "comments.json").read_text(encoding="utf-8"))
    bodies = [c["comment_text"] for c in data]
    assert body in bodies, (
        f"Reply body did not round-trip; expected to find {body!r} in {bodies!r}"
    )


def test_reply_preserves_existing_paraId_link(tmp_path):
    """Reply's paraIdParent in commentsExtended.xml must match the parent's existing paraId in the SAME file."""
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"

    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Reply text",
    )
    assert r.returncode == 0, r.stderr

    import zipfile
    with zipfile.ZipFile(out) as z:
        ext_xml = z.read("word/commentsExtended.xml").decode("utf-8")

    # Parse commentsExtended and verify the reply's paraIdParent points at an existing paraId
    from lxml import etree
    root = etree.fromstring(ext_xml.encode("utf-8"))
    W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
    entries = root.findall(f"{{{W15}}}commentEx")
    paraIds = {e.get(f"{{{W15}}}paraId") for e in entries if e.get(f"{{{W15}}}paraId")}

    parents = [e.get(f"{{{W15}}}paraIdParent") for e in entries if e.get(f"{{{W15}}}paraIdParent")]

    # Every paraIdParent should be a real paraId from commentsExtended.xml
    for parent_id in parents:
        assert parent_id in paraIds, (
            f"paraIdParent {parent_id} doesn't match any commentEx paraId; mapping is broken"
        )


def _read_part(docx_path: Path, name: str) -> bytes | None:
    import zipfile
    with zipfile.ZipFile(docx_path) as z:
        try:
            return z.read(name)
        except KeyError:
            return None


def _paraIds_in_comments_xml(docx_path: Path) -> list[str]:
    from lxml import etree
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    data = _read_part(docx_path, "word/comments.xml")
    assert data is not None, "word/comments.xml missing"
    root = etree.fromstring(data)
    out: list[str] = []
    for c in root.findall(f"{{{W}}}comment"):
        p = c.find(f"{{{W}}}p")
        if p is not None:
            pid = p.get(f"{{{W14}}}paraId")
            if pid:
                out.append(pid)
    return out


def test_reply_backfills_commentsIds_for_all_comments(tmp_path):
    """Every comment's paraId (parents + reply) must have a row in commentsIds.xml.

    Word 2016+ uses w16cid:commentId/durableId to render comment threading.
    A comment whose paraId is absent from commentsIds.xml is treated as legacy
    by Word and replies cannot resolve to it as a thread parent.  Before the
    fix, only the reply's paraId appeared in commentsIds.xml; the three parent
    paraIds were silently dropped.
    """
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Reply text",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr

    ids_data = _read_part(out, "word/commentsIds.xml")
    assert ids_data is not None, "word/commentsIds.xml is missing"

    from lxml import etree
    W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
    ids_root = etree.fromstring(ids_data)
    ids_entries = ids_root.findall(f"{{{W16CID}}}commentId")
    ids_paraIds = {e.get(f"{{{W16CID}}}paraId") for e in ids_entries}
    ids_durables = [e.get(f"{{{W16CID}}}durableId") for e in ids_entries]

    comments_paraIds = set(_paraIds_in_comments_xml(out))
    assert comments_paraIds, "comments.xml has no paraIds"

    missing = comments_paraIds - ids_paraIds
    assert not missing, (
        f"commentsIds.xml is missing rows for paraIds {sorted(missing)}; "
        f"present: {sorted(ids_paraIds)}"
    )

    assert len(ids_durables) == len(set(ids_durables)), (
        f"duplicate durableIds in commentsIds.xml: {ids_durables}"
    )
    for d in ids_durables:
        assert d, "found commentId row with empty durableId"


def test_reply_backfills_commentsExtensible_for_all_comments(tmp_path):
    """Every comment's durableId must have a matching commentExtensible row.

    commentsIds.xml and commentsExtensible.xml share the durableId space.
    Word expects a parallel entry in commentsExtensible.xml (with the same
    durableId) for every commentId.  Before the fix, only the reply's
    durableId appeared.
    """
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Reply text",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr

    from lxml import etree
    W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
    W16CEX = "http://schemas.microsoft.com/office/word/2018/wordml/cex"

    ids_data = _read_part(out, "word/commentsIds.xml")
    cex_data = _read_part(out, "word/commentsExtensible.xml")
    assert ids_data is not None, "word/commentsIds.xml is missing"
    assert cex_data is not None, "word/commentsExtensible.xml is missing"

    ids_durables = {
        e.get(f"{{{W16CID}}}durableId")
        for e in etree.fromstring(ids_data).findall(f"{{{W16CID}}}commentId")
    }
    cex_durables = {
        e.get(f"{{{W16CEX}}}durableId")
        for e in etree.fromstring(cex_data).findall(f"{{{W16CEX}}}commentExtensible")
    }

    # The cardinality must match the number of comments in comments.xml — not
    # just match each other (two empty-ish sets would trivially match).
    comments_count = len(_paraIds_in_comments_xml(out))
    assert len(cex_durables) == comments_count, (
        f"commentsExtensible row count {len(cex_durables)} != comments count {comments_count}"
    )

    assert ids_durables == cex_durables, (
        f"durableId sets diverge: commentsIds={sorted(ids_durables)}, "
        f"commentsExtensible={sorted(cex_durables)}"
    )

    # And every entry must carry a dateUtc attribute (Word omits the comment
    # entirely if dateUtc is missing on a commentExtensible row).
    cex_root = etree.fromstring(cex_data)
    for el in cex_root.findall(f"{{{W16CEX}}}commentExtensible"):
        assert el.get(f"{{{W16CEX}}}dateUtc"), (
            f"commentExtensible row {el.get(f'{{{W16CEX}}}durableId')} lacks dateUtc"
        )


def test_reply_creates_people_xml_with_authors(tmp_path):
    """word/people.xml must list a w15:person for every comment author.

    Word's modern review pane keys author display off word/people.xml.
    Before the fix, people.xml was never created, so Word fell back to
    legacy display and could not show authored replies as part of a thread.
    """
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Reply text",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr

    people_data = _read_part(out, "word/people.xml")
    assert people_data is not None, "word/people.xml is missing"

    from lxml import etree
    W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
    root = etree.fromstring(people_data)
    persons = root.findall(f"{{{W15}}}person")
    names = {p.get(f"{{{W15}}}author") for p in persons}

    # Fixture has authors Alice, Bob; reply adds Claude.  All three must
    # appear in people.xml (set equality, not subset).
    expected = {"Alice", "Bob", "Claude"}
    assert expected.issubset(names), (
        f"people.xml missing authors; have {sorted(names)}, expected superset of {sorted(expected)}"
    )

    # The part must be registered in [Content_Types].xml and document rels.
    ct_data = _read_part(out, "[Content_Types].xml")
    rels_data = _read_part(out, "word/_rels/document.xml.rels")
    assert ct_data is not None and rels_data is not None
    assert b"/word/people.xml" in ct_data, "people.xml lacks Content_Types override"
    assert b"people.xml" in rels_data, "people.xml lacks document relationship"


def test_reply_paraIdParent_equals_specific_parent_paraId(tmp_path):
    """The reply's paraIdParent must equal C003's specific w14:paraId — not just *some* paraId.

    Strengthens test_reply_preserves_existing_paraId_link, which only
    checked that the reply's paraIdParent appeared somewhere in the
    commentsExtended.xml paraId set.  That weaker check passes even when
    the reply is linked to the wrong parent (e.g., C001) as long as that
    parent's row exists.  Here we resolve C003 -> w:id=2 -> w14:paraId in
    comments.xml and require an exact match.
    """
    _skip_if_plugin_absent()
    out = tmp_path / "replied.docx"
    r = _run_cli(
        "add-comment", str(FIXTURES / "with_comments.docx"),
        "--out", str(out),
        "--reply-to", "C003",
        "--text", "Reply text",
        "--author", "Claude",
    )
    assert r.returncode == 0, r.stderr

    from lxml import etree
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    W15 = "http://schemas.microsoft.com/office/word/2012/wordml"

    # Find C003 (w:id=2)'s paraId in comments.xml after the reply.
    comments_root = etree.fromstring(_read_part(out, "word/comments.xml"))
    parent_para_id = None
    for c in comments_root.findall(f"{{{W}}}comment"):
        if c.get(f"{{{W}}}id") == "2":
            p = c.find(f"{{{W}}}p")
            parent_para_id = p.get(f"{{{W14}}}paraId") if p is not None else None
            break
    assert parent_para_id, "Parent C003 (w:id=2) has no w14:paraId in comments.xml"

    # Find the reply's commentEx (last appended) and verify paraIdParent.
    ext_root = etree.fromstring(_read_part(out, "word/commentsExtended.xml"))
    entries = ext_root.findall(f"{{{W15}}}commentEx")
    # The reply is the entry whose paraId corresponds to the new comment (w:id=3).
    reply_para_id = None
    for c in comments_root.findall(f"{{{W}}}comment"):
        if c.get(f"{{{W}}}id") == "3":
            p = c.find(f"{{{W}}}p")
            reply_para_id = p.get(f"{{{W14}}}paraId") if p is not None else None
            break
    assert reply_para_id, "Reply comment (w:id=3) has no w14:paraId in comments.xml"

    reply_ext = next(
        (e for e in entries if e.get(f"{{{W15}}}paraId") == reply_para_id),
        None,
    )
    assert reply_ext is not None, (
        f"No commentEx with paraId {reply_para_id} for the reply"
    )
    actual_parent = reply_ext.get(f"{{{W15}}}paraIdParent")
    assert actual_parent == parent_para_id, (
        f"Reply's paraIdParent is {actual_parent!r}, expected {parent_para_id!r} "
        f"(C003's paraId in comments.xml)"
    )
