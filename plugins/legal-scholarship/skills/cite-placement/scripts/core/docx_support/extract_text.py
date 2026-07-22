"""Extract paragraphs and tables from a .docx file as structured JSON and Markdown.

Imported from word-docx skill for standalone use. Reads a Word document via
python-docx, producing ParagraphRecord objects and a Markdown string.
"""

from __future__ import annotations

import json
import logging
import sys
from pathlib import Path
from typing import TYPE_CHECKING

sys.path.insert(0, str(Path(__file__).resolve().parent))

from models import DiagnosticEntry, ParagraphRecord

if TYPE_CHECKING:
    from docx.table import Table

logger = logging.getLogger(__name__)

_HEADING_MAP: dict[str, str] = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Title": "#",
}


def _style_to_md(style_name: str, text: str) -> str:
    if style_name in _HEADING_MAP:
        return f"{_HEADING_MAP[style_name]} {text}"
    if style_name == "Subtitle":
        return f"*{text}*"
    if style_name.startswith("List"):
        return f"- {text}"
    return text


def _all_text_from_element(element) -> str:
    """Extract all text from an OOXML element, including tracked changes."""
    from docx.oxml.ns import qn

    parts: list[str] = []
    for node in element.iter():
        if node.tag in (qn("w:t"), qn("w:delText")):
            if node.text:
                parts.append(node.text)
    return "".join(parts).strip()


def _escape_pipe(text: str) -> str:
    return text.replace("|", "\\|")


def _table_to_md(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_escape_pipe(" ".join(_all_text_from_element(c._element).split())) for c in row.cells]
        rows.append(cells)
    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")

    widths = [max(len(rows[ri][ci]) for ri in range(len(rows))) for ci in range(col_count)]
    widths = [max(w, 3) for w in widths]

    def _fmt_row(cells: list[str]) -> str:
        padded = [cells[i].ljust(widths[i]) for i in range(col_count)]
        return "| " + " | ".join(padded) + " |"

    lines: list[str] = [_fmt_row(rows[0])]
    sep = "| " + " | ".join("-" * w for w in widths) + " |"
    lines.append(sep)
    for row in rows[1:]:
        lines.append(_fmt_row(row))
    return "\n".join(lines)


def extract_text(
    input_path: Path,
) -> tuple[list[ParagraphRecord], str, list[DiagnosticEntry]]:
    """Extract paragraphs and tables from *input_path*."""
    from docx import Document

    diagnostics: list[DiagnosticEntry] = []
    input_path = Path(input_path).resolve()

    if not input_path.exists():
        diagnostics.append(DiagnosticEntry(
            level="error", source="extract_text",
            message=f"File not found: {input_path}",
        ))
        return [], "", diagnostics

    try:
        doc = Document(str(input_path))
    except Exception as exc:
        diagnostics.append(DiagnosticEntry(
            level="error", source="extract_text",
            message=f"Failed to open document: {exc}",
        ))
        return [], "", diagnostics

    paragraphs: list[ParagraphRecord] = []
    md_parts: list[str] = []

    from docx.oxml.ns import qn

    table_elements = {tbl._element: tbl for tbl in doc.tables}

    para_idx = 0
    for child in doc.element.body:
        tag = child.tag

        if tag == qn("w:p"):
            if para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1

            text = _all_text_from_element(para._element).strip()
            style_name = para.style.name if para.style else "Normal"

            if text:
                paragraphs.append(ParagraphRecord(
                    index=len(paragraphs), style=style_name, text=text,
                ))
                md_parts.append(_style_to_md(style_name, text))
            else:
                md_parts.append("")

        elif tag == qn("w:tbl"):
            tbl_obj = table_elements.get(child)
            if tbl_obj is not None:
                try:
                    md_table = _table_to_md(tbl_obj)
                    if md_table:
                        md_parts.append("")
                        md_parts.append(md_table)
                        md_parts.append("")
                except Exception as exc:
                    diagnostics.append(DiagnosticEntry(
                        level="warning", source="extract_text",
                        message=f"Could not render table: {exc}",
                    ))

    markdown = "\n".join(md_parts).strip() + "\n"

    diagnostics.append(DiagnosticEntry(
        level="info", source="extract_text",
        message=(
            f"Extracted {len(paragraphs)} paragraph(s) and "
            f"{len(doc.tables)} table(s) from {input_path.name}"
        ),
    ))
    return paragraphs, markdown, diagnostics


def write_text(paragraphs: list[ParagraphRecord], markdown: str, out_dir: Path) -> None:
    """Write extracted content to *out_dir* as paragraphs.json and document.md."""
    out_dir = Path(out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    para_path = out_dir / "paragraphs.json"
    md_path = out_dir / "document.md"

    para_data = [p.model_dump() for p in paragraphs]
    para_path.write_text(json.dumps(para_data, indent=2, ensure_ascii=False), encoding="utf-8")
    md_path.write_text(markdown, encoding="utf-8")
    logger.info("Wrote %s and %s", para_path, md_path)
