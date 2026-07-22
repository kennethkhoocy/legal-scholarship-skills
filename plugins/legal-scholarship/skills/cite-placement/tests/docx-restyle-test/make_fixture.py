#!/usr/bin/env python3
"""Build a fully synthetic ``input.docx`` fixture for the docx-restyle tests.

The fixture contains a title, an author line, six neutral academic body
paragraphs, and five numeric footnotes carrying fake Bluebook-style texts.
Everything is invented; there is no personal data.

Idempotent: running this module overwrites ``input.docx`` deterministically
(same paragraphs, same footnotes, same display numbering on every run).

Run standalone::

    python make_fixture.py
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

# Put scripts/core on sys.path so docx_support imports as a package.
_HERE = Path(__file__).resolve().parent                       # .../tests/docx-restyle-test
_CORE = _HERE.parent.parent / "scripts" / "core"              # .../scripts/core
if str(_CORE) not in sys.path:
    sys.path.insert(0, str(_CORE))

from docx import Document  # noqa: E402
from lxml import etree  # noqa: E402

from docx_support.footnotes import insert_footnote  # noqa: E402

FIXTURE = _HERE / "input.docx"

TITLE = "The Economics of Widget Regulation"
AUTHOR_LINE = "Jane Q. Author, Example University, jane.author@example.edu"

# Six neutral academic body paragraphs. Each carries a unique anchor phrase
# (used to locate it for footnote insertion and by test_insert).
BODY_PARAGRAPHS = [
    "Widget markets have long attracted the attention of economists studying "
    "regulatory intervention and its welfare consequences.",
    "Early empirical work suggested that price controls on widgets reduced "
    "consumer surplus over the short run.",
    "Subsequent studies revisited these findings using richer panel data and "
    "more credible identification.",
    "The theoretical mechanism connecting entry barriers to observed markups "
    "remains contested throughout the literature.",
    "Recent policy proposals emphasize transparency requirements and "
    "standardized disclosure for widget manufacturers.",
    "This article proceeds in four parts, beginning with the institutional "
    "background of widget regulation.",
]

# ponytail: numeric footnotes only -- symbol/customMarkFollows footnotes not
# reproducible via insert_footnote; extend fixture if display-map offsets ever regress.
#
# Each tuple is (anchor_phrase_in_body, footnote_text_with_markers). The
# anchor phrase identifies the target paragraph; footnotes are inserted in
# document order so their displayed numbers are 1..5, which is what the fake
# short-form footnotes (Id. / supra note 1) assume.
FOOTNOTE_PLAN = [
    (
        "welfare consequences",
        "Jane Q. Author & John B. Scholar, *The Economics of Widget Regulation*, "
        "12 J. Widget Stud. 345 (2020).",
    ),
    (
        "consumer surplus",
        "*Id.* at 350.",
    ),
    (
        "credible identification",
        "Author & Scholar, *supra* note 1, at 347.",
    ),
    (
        "observed markups",
        "The magnitude of these effects depends on identifying assumptions that "
        "the empirical design cannot fully test.",
    ),
    (
        "standardized disclosure",
        "For the underlying dataset, see https://example.edu/widget-data "
        "(last visited Jan. 1, 2020).",
    ),
]

# The footnote texts as extract_footnotes returns them: format markers
# (*italic*, **bold**) are consumed into runs and stripped from the plain text.
EXPECTED_TEXTS = [
    "Jane Q. Author & John B. Scholar, The Economics of Widget Regulation, "
    "12 J. Widget Stud. 345 (2020).",
    "Id. at 350.",
    "Author & Scholar, supra note 1, at 347.",
    "The magnitude of these effects depends on identifying assumptions that "
    "the empirical design cannot fully test.",
    "For the underlying dataset, see https://example.edu/widget-data "
    "(last visited Jan. 1, 2020).",
]

# Anchor used by test_insert.py to place a new footnote.
INSERT_ANCHOR = "institutional background"

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _qn(tag: str) -> str:
    return f"{{{_W}}}{tag.split(':', 1)[1]}"


def para_index_of(docx_path: Path, phrase: str) -> int:
    """Return the 0-based body-paragraph index whose text contains *phrase*.

    Scans w:p children of the body in document order, mirroring
    extract_footnote_locations. Robust to leading/trailing empty paragraphs
    the default template may add.
    """
    with zipfile.ZipFile(docx_path, "r") as zf:
        tree = etree.fromstring(zf.read("word/document.xml"))
    body = tree.find(_qn("w:body"))
    if body is None:
        raise ValueError("document has no body")
    idx = 0
    for child in body:
        if child.tag != _qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(_qn("w:t")))
        if phrase in text:
            return idx
        idx += 1
    raise ValueError(f"anchor phrase not found: {phrase!r}")


def build() -> Path:
    """Build (or overwrite) the fixture and return its path."""
    doc = Document()

    try:
        doc.add_heading(TITLE, level=0)
    except KeyError:
        # Default template missing the Title style: fall back to a normal paragraph.
        doc.add_paragraph(TITLE)

    doc.add_paragraph(AUTHOR_LINE)

    for para in BODY_PARAGRAPHS:
        doc.add_paragraph(para)

    doc.save(str(FIXTURE))

    # Insert footnotes in document order so displayed numbers are 1..5.
    for anchor, text in FOOTNOTE_PLAN:
        idx = para_index_of(FIXTURE, anchor)
        insert_footnote(FIXTURE, paragraph_index=idx, text=text, after_text=None)

    return FIXTURE


if __name__ == "__main__":
    path = build()
    print(f"Wrote fixture: {path}")
