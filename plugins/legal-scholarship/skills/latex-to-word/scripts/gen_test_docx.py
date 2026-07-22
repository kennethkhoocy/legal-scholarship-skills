"""Generate a test .docx with academic legal prose and proper Word footnotes.

Uses direct OOXML manipulation to create real footnotes (not inline text).
"""

import os
import zipfile
import shutil
import tempfile
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Track footnote IDs globally per document generation
_footnote_counter = 1  # starts at 1; 0 and -1 are reserved


def reset_counter():
    global _footnote_counter
    _footnote_counter = 1


def add_footnote_ref(paragraph, run_text, footnote_id):
    """Add body text + a footnote reference mark to a paragraph."""
    # Body text
    run = paragraph.add_run(run_text)

    # Footnote reference superscript
    ref_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    ref_run.append(rPr)

    ref_el = OxmlElement('w:footnoteReference')
    ref_el.set(qn('w:id'), str(footnote_id))
    ref_run.append(ref_el)
    paragraph._element.append(ref_run)


def build_footnote_xml(footnote_id, text):
    """Build a single <w:footnote> element."""
    return (
        f'<w:footnote w:id="{footnote_id}" '
        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/>'
        f'<w:vertAlign w:val="superscript"/></w:rPr>'
        f'<w:footnoteRef/></w:r>'
        f'<w:r><w:t xml:space="preserve"> </w:t></w:r>'
        f'<w:r><w:t xml:space="preserve">{escape_xml(text)}</w:t></w:r>'
        f'</w:p>'
        f'</w:footnote>'
    )


def escape_xml(text):
    """Escape XML special characters."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def fn(paragraph, run_text, footnote_text):
    """Shorthand: add body text with footnote. Returns (footnote_id, footnote_text) for later."""
    global _footnote_counter
    fid = _footnote_counter
    _footnote_counter += 1
    add_footnote_ref(paragraph, run_text, fid)
    return (fid, footnote_text)


def inject_footnotes(docx_path, footnotes):
    """Post-process: inject footnotes XML into the .docx ZIP."""
    # Build the full footnotes.xml
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        # Separator footnotes (required by Word)
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>',
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>',
    ]
    for fid, text in footnotes:
        parts.append(build_footnote_xml(fid, text))
    parts.append('</w:footnotes>')
    footnotes_xml = '\n'.join(parts)

    # Manipulate the ZIP
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp, 'w') as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                # Ensure footnotes relationship exists
                zout.writestr(item, data)
            elif item.filename == '[Content_Types].xml':
                # Add footnotes content type if missing
                ct = data.decode('utf-8')
                if 'footnotes.xml' not in ct:
                    ct = ct.replace(
                        '</Types>',
                        '<Override PartName="/word/footnotes.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        '</Types>'
                    )
                zout.writestr(item, ct)
            elif item.filename == 'word/_rels/document.xml.rels':
                # Add footnotes relationship if missing
                rels = data.decode('utf-8')
                if 'footnotes.xml' not in rels:
                    rels = rels.replace(
                        '</Relationships>',
                        '<Relationship Id="rIdFootnotes" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
                        'Target="footnotes.xml"/>'
                        '</Relationships>'
                    )
                zout.writestr(item, rels)
            elif item.filename == 'word/styles.xml':
                # Inject FootnoteReference + FootnoteText styles if missing
                styles = data.decode('utf-8')
                if 'FootnoteReference' not in styles:
                    fn_ref_style = (
                        '<w:style w:type="character" w:styleId="FootnoteReference" '
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:name w:val="footnote reference"/>'
                        '<w:uiPriority w:val="99"/>'
                        '<w:unhideWhenUsed/>'
                        '<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        '</w:style>'
                    )
                    styles = styles.replace('</w:styles>', fn_ref_style + '</w:styles>')
                if 'FootnoteText' not in styles:
                    fn_text_style = (
                        '<w:style w:type="paragraph" w:styleId="FootnoteText" '
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:name w:val="footnote text"/>'
                        '<w:basedOn w:val="Normal"/>'
                        '<w:uiPriority w:val="99"/>'
                        '<w:unhideWhenUsed/>'
                        '<w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr>'
                        '</w:style>'
                    )
                    styles = styles.replace('</w:styles>', fn_text_style + '</w:styles>')
                zout.writestr(item, styles)
            else:
                zout.writestr(item, data)

        # Write footnotes.xml
        zout.writestr('word/footnotes.xml', footnotes_xml.encode('utf-8'))

    shutil.move(tmp, docx_path)


def main():
    output_path = os.path.join(ROOT, "input", "input_v1.docx")
    os.makedirs(os.path.join(ROOT, "input"), exist_ok=True)
    reset_counter()

    doc = Document()
    footnotes = []  # collect (id, text) tuples

    # -- Title --
    doc.add_heading(
        'Regulatory Capture and Administrative Law\u2019s Accountability Deficit',
        level=0
    )

    # -- Section I --
    doc.add_heading('I. Introduction', level=1)

    p = doc.add_paragraph()
    footnotes.append(fn(p,
        'The phenomenon of regulatory capture\u2014whereby an agency tasked with '
        'regulating an industry instead advances that industry\u2019s interests\u2014has '
        'long troubled scholars of administrative law.',
        'George J. Stigler, \u201cThe Theory of Economic Regulation,\u201d 2 Bell J. Econ. '
        '& Mgmt. Sci. 3, 3\u201321 (1971).'
    ))
    footnotes.append(fn(p,
        ' The concept, first articulated in the economic literature of the early 1970s, '
        'has since migrated into legal scholarship, where it has been deployed to explain '
        'phenomena ranging from the leniency of financial regulation to the persistence '
        'of occupational licensing regimes.',
        'See generally Daniel Carpenter & David A. Moss, Preventing Regulatory Capture: '
        'Special Interest Influence and How to Limit It 12\u201345 (2013) (surveying the '
        'interdisciplinary literature on capture).'
    ))

    p2 = doc.add_paragraph()
    footnotes.append(fn(p2,
        'This Article argues that the conventional account of regulatory capture is '
        'incomplete. While existing scholarship has focused primarily on the mechanisms '
        'by which industry actors influence agency decision-making\u2014lobbying, revolving '
        'doors, and information asymmetries\u2014it has paid insufficient attention to the '
        'structural features of administrative law that facilitate capture.',
        'For a comprehensive taxonomy of capture mechanisms, see Stigler, supra note 1, '
        'at 10\u201315; see also Michael E. Levine & Jennifer L. Forrence, \u201cRegulatory '
        'Capture, Public Interest, and the Public Agenda,\u201d 6 J.L. Econ. & Org. 167, '
        '170\u201378 (1990).'
    ))
    footnotes.append(fn(p2,
        ' In particular, the Article contends that the notice-and-comment process '
        'under \u00a7 553 of the Administrative Procedure Act, far from ensuring democratic '
        'accountability, often serves as a venue for well-resourced industry participants '
        'to exert disproportionate influence over rulemaking outcomes.',
        '5 U.S.C. \u00a7 553 (2018). The notice-and-comment requirement was enacted as '
        'part of the Administrative Procedure Act of 1946, Pub. L. No. 79-404, 60 Stat. '
        '237. For historical background, see Walter Gellhorn, \u201cThe Administrative '
        'Procedure Act: The Beginning,\u201d 72 Va. L. Rev. 219 (1986).'
    ))

    # -- Section II --
    doc.add_heading('II. The Structure of Regulatory Capture', level=1)

    p3 = doc.add_paragraph()
    footnotes.append(fn(p3,
        'Regulatory capture manifests along a spectrum. At one extreme lies what '
        'Carpenter and Moss have termed \u201cstrong capture,\u201d in which an agency\u2019s '
        'actions consistently reflect the preferences of the regulated industry rather '
        'than the public interest.',
        'Carpenter & Moss, supra note 2, at 13\u201314 (distinguishing \u201cstrong\u201d from '
        '\u201cweak\u201d capture based on the degree of policy distortion).'
    ))
    footnotes.append(fn(p3,
        ' At the other end of the spectrum lies \u201cweak capture,\u201d characterized by '
        'more subtle forms of influence\u2014cognitive biases, cultural affinity between '
        'regulators and the regulated, and the gradual normalization of industry '
        'perspectives within the agency.',
        'This weaker form of capture is particularly insidious because it operates '
        'below the threshold of conscious awareness. As one former commissioner of the '
        'Securities and Exchange Commission observed, \u201c[y]ou don\u2019t have to be '
        'corrupt to be captured.\u201d See Rachel E. Barkow, \u201cInsulating Agencies: '
        'Avoiding Capture Through Institutional Design,\u201d 89 Tex. L. Rev. 15, 22 (2010).'
    ))

    p4 = doc.add_paragraph(
        'The empirical evidence for capture is substantial, though methodologically '
        'contested. Studies of the Federal Communications Commission, the Interstate '
        'Commerce Commission, and the Environmental Protection Agency have all '
        'documented patterns consistent with the capture hypothesis. Yet isolating '
        'the causal mechanisms\u2014distinguishing, for example, between an agency that '
        'has been \u201ccaptured\u201d and one that has simply reached a different policy '
        'conclusion on the merits\u2014remains a formidable analytical challenge.'
    )

    # -- Section III --
    doc.add_heading('III. Administrative Law\u2019s Structural Vulnerabilities', level=1)

    p5 = doc.add_paragraph()
    footnotes.append(fn(p5,
        'Several features of contemporary administrative law create conditions '
        'conducive to regulatory capture. First, the notice-and-comment process, '
        'despite its democratic aspirations, systematically favors organized interests '
        'over diffuse public beneficiaries.',
        'See Wendy E. Wagner, \u201cAdministrative Law, Filter Failure, and the Crisis '
        'in Administrative Legitimacy,\u201d 59 Duke L.J. 1321, 1340\u201355 (2010) '
        '(documenting the dominance of industry comments in EPA rulemakings under the '
        'Clean Air Act).'
    ))

    p6 = doc.add_paragraph()
    footnotes.append(fn(p6,
        'Second, judicial review under the \u201carbitrary and capricious\u201d standard '
        'provides only limited protection against capture. Courts reviewing agency '
        'action under \u00a7 706(2)(A) of the APA focus primarily on the rationality of '
        'the agency\u2019s explanation, not on whether the decision was unduly influenced '
        'by private interests.',
        '5 U.S.C. \u00a7 706(2)(A) (2018); see Motor Vehicle Mfrs. Ass\u2019n v. State Farm '
        'Mut. Auto. Ins. Co., 463 U.S. 29, 43 (1983) (\u201c[T]he agency must examine '
        'the relevant data and articulate a satisfactory explanation for its action.\u201d).'
    ))
    footnotes.append(fn(p6,
        ' While \u201chard look\u201d review has occasionally served as a check on agency '
        'overreach, it is poorly calibrated to detect the subtle distortions '
        'characteristic of weak capture.',
        'See Mark Seidenfeld, \u201cDemystifying Deossification: Rethinking Recent '
        'Proposals to Modify Judicial Review of Notice and Comment Rulemaking,\u201d '
        '75 Tex. L. Rev. 483, 490\u2013514 (1997) (arguing that hard-look review may '
        'paradoxically worsen capture by increasing the costs of rulemaking).'
    ))

    # -- Section IV --
    doc.add_heading('IV. Toward Institutional Reform', level=1)

    p7 = doc.add_paragraph()
    footnotes.append(fn(p7,
        'If the diagnosis offered in this Article is correct\u2014that regulatory capture '
        'is facilitated not merely by the strategic behavior of industry actors but '
        'by structural features of the administrative state itself\u2014then the remedy '
        'must be structural as well. Several reforms merit consideration.',
        'The reforms proposed here build on, but depart from, earlier proposals. '
        'Compare Barkow, supra note 6, at 30\u201350 (emphasizing insulation through '
        'structural design), with Nicholas Bagley, \u201cThe Procedure Fetish,\u201d '
        '118 Mich. L. Rev. 345, 380\u201395 (2019) (arguing for reduced procedural '
        'requirements rather than additional structural constraints).'
    ))

    p8 = doc.add_paragraph(
        'First, agencies should be required to publish \u201cregulatory impact '
        'assessments\u201d that explicitly address the risk of capture. Such assessments '
        'would compel agencies to identify the distributional consequences of proposed '
        'rules and to explain why the chosen regulatory approach serves the public '
        'interest rather than the interests of a concentrated group of beneficiaries. '
        'Second, Congress should consider establishing an independent office of '
        'regulatory review\u2014modeled loosely on the Congressional Budget Office\u2014tasked '
        'with evaluating whether agency rules reflect evidence of capture.'
    )
    footnotes.append(fn(p8,
        ' A similar proposal was advanced by the Administrative Conference of the '
        'United States in its 2012 recommendations.',
        'See Admin. Conf. of the U.S., Recommendation 2012-1, Regulatory Analysis '
        'Requirements, 77 Fed. Reg. 47,802 (Aug. 10, 2012).'
    ))

    p9 = doc.add_paragraph(
        'These proposals are admittedly imperfect. Institutional design is an '
        'inherently second-best enterprise, and no structural reform can eliminate '
        'the risk of capture entirely. But the current framework\u2019s failure to address '
        'the structural dimensions of the problem suggests that even incremental '
        'reform could yield significant benefits for the accountability and legitimacy '
        'of the administrative state.'
    )

    # Save the base document
    doc.save(output_path)

    # Post-process: inject real footnotes into the ZIP
    inject_footnotes(output_path, footnotes)

    print(f"Created: {output_path}")
    print(f"Footnotes: {len(footnotes)}")


if __name__ == "__main__":
    main()
