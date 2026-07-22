# -*- coding: utf-8 -*-
"""Convert estout/booktabs LaTeX table fragments into native, editable Word tables
matching the compiled LaTeX look. Booktabs rules -> horizontal cell borders;
\\multicolumn -> merged cells; \\cmidrule -> partial under-rules; \\sym{}/^{}/_{} ->
super/subscripts; \\makecell{a\\\\b} protected from row-splitting; math symbols and
\\cref{} resolved. add_estout_table() populates a doc (batch); build() = one-table doc."""
import re
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from siunitx_expand import expand_siunitx_commands
from mhchem_unicode import replace_ce_commands

FONT = "Linux Libertine G"
CREF_PREFIXES = {
    "tab": ("Table", "Tables"),
    "tbl": ("Table", "Tables"),
    "table": ("Table", "Tables"),
    "fig": ("Figure", "Figures"),
    "figure": ("Figure", "Figures"),
    "sec": ("Section", "Sections"),
    "section": ("Section", "Sections"),
    "subsec": ("Section", "Sections"),
    "subsection": ("Section", "Sections"),
    "subsubsec": ("Section", "Sections"),
    "subsubsection": ("Section", "Sections"),
    "eq": ("Equation", "Equations"),
    "equation": ("Equation", "Equations"),
    "app": ("Appendix", "Appendices"),
    "appendix": ("Appendix", "Appendices"),
}
HEAVY = 10  # eighths of a pt -> 1.25pt (booktabs \toprule/\bottomrule)
LIGHT = 5   # 0.625pt          (\midrule / \cmidrule)

SYMBOLS = {
    r'\Delta': 'Δ', r'\delta': 'δ', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
    r'\mu': 'μ', r'\sigma': 'σ', r'\rho': 'ρ', r'\theta': 'θ', r'\lambda': 'λ',
    r'\times': '×', r'\pm': '±', r'\mp': '∓', r'\leq': '≤', r'\geq': '≥',
    r'\neq': '≠', r'\approx': '≈', r'\to': '→', r'\Rightarrow': '⇒',
    r'\cdot': '·', r'\ldots': '…', r'\dots': '…', r'\infty': '∞',
}


def set_font(name):
    global FONT
    FONT = name


def set_cref_prefixes(prefixes):
    global CREF_PREFIXES
    CREF_PREFIXES = dict(prefixes or CREF_PREFIXES)


def _border_tc(tc, edge, sz):
    tcPr = tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    e = b.find(qn('w:' + edge))
    if e is None:
        e = OxmlElement('w:' + edge); b.append(e)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')


def _run(p, text, pt, italic=False, bold=False, sup=False, sub=False):
    r = p.add_run(text)
    r.font.size = Pt(pt); r.font.italic = italic; r.font.bold = bold
    r.font.superscript = sup; r.font.subscript = sub
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), FONT)
    return r


def _brace(s, i):
    """s[i]=='{'; return (inner, index_after_closing)."""
    d = 0
    for k in range(i, len(s)):
        if s[k] == '{':
            d += 1
        elif s[k] == '}':
            d -= 1
            if d == 0:
                return s[i + 1:k], k + 1
    return s[i + 1:], len(s)


def _protect_makecell(frag):
    """Inline \\makecell[..]{a\\\\b} -> 'a b' so its internal \\\\ doesn't break row split."""
    out, i = [], 0
    while i < len(frag):
        m = re.match(r'\\makecell(\[[^\]]*\])?\{', frag[i:])
        if m:
            inner, end = _brace(frag, i + m.end() - 1)
            out.append(inner.replace('\\\\', ' '))
            i = end
        else:
            out.append(frag[i]); i += 1
    return ''.join(out)


def _plain_len(text):
    """Approximate rendered character width of a cell (for content-aware sizing)."""
    t = replace_ce_commands(re.sub(r'\\sym\{([^}]*)\}', r'\1', text))
    t = re.sub(r'\\[a-zA-Z]+', '', t)
    for ch in '${}^_\\':
        t = t.replace(ch, '')
    return max(1, len(t.strip()))


def _resolve_crefs(s, numbermap):
    def repl(m):
        cmd, lab = m.group(1), m.group(2)
        n = (numbermap or {}).get(lab, '')
        if cmd == 'eqref':
            return f'({n})' if n else ''
        if cmd == 'ref':
            return n
        key = lab.split(':', 1)[0]
        pfx = (CREF_PREFIXES.get(key) or CREF_PREFIXES.get(key.lower()) or ('', ''))[0]
        pfx = (pfx + ' ') if pfx else ''
        return (pfx + n).strip()
    return re.sub(r'\\(cref|Cref|ref|eqref|autoref)\{([^}]*)\}', repl, s)


def inline_to_runs(p, s, pt, numbermap=None):
    """Render an inline LaTeX snippet onto paragraph p as styled runs."""
    if not s:
        return
    s = expand_siunitx_commands(s)
    s = replace_ce_commands(s)
    s = _resolve_crefs(s, numbermap)
    s = (s.replace('---', '—').replace('--', '–')
         .replace('\\$', '\x00').replace('\\%', '%').replace('\\&', '&').replace('\\#', '#')
         .replace('\\,', ' ').replace('\\;', ' ').replace('\\ ', ' ').replace('\\!', '')
         .replace('~', ' '))
    for k, v in SYMBOLS.items():
        s = s.replace(k, v)
    for _ in range(3):  # unwrap nested simple wrappers
        s2 = re.sub(r'\\(?:text|mathrm|mathbf|mathbb|mathit|hat|bar|tilde|vec|operatorname|ensuremath)\{([^{}]*)\}', r'\1', s)
        if s2 == s:
            break
        s = s2
    s = re.sub(r'\\textsuperscript\{([^{}]*)\}', r'^{\1}', s)
    s = re.sub(r'\\textsubscript\{([^{}]*)\}', r'_{\1}', s)
    s = re.sub(r'\\(log|exp|ln|min|max|sup|inf|det|var|cov)\b', r'\1', s)
    s = s.replace('$', '')

    runs, buf, i, n = [], '', 0, len(s)

    def emit(t, **kw):
        if t:
            runs.append((t, kw))

    while i < n:
        c = s[i]
        if c in '^_':
            emit(buf); buf = ''
            issup = c == '^'; i += 1
            if i < n and s[i] == '{':
                j = s.find('}', i)
                arg = s[i + 1:j] if j >= 0 else s[i + 1:]
                i = j + 1 if j >= 0 else n
            else:
                arg = s[i] if i < n else ''; i += 1
            emit(arg, sup=issup, sub=not issup)
        elif c == '\\':
            m = re.match(r'\\(textit|emph|textsc|textbf)\{([^}]*)\}', s[i:])
            if m:
                emit(buf); buf = ''
                emit(m.group(2), italic=m.group(1) != 'textbf', bold=m.group(1) == 'textbf')
                i += m.end()
            else:
                m2 = re.match(r'\\[a-zA-Z]+', s[i:])
                i += m2.end() if m2 else 1
        else:
            buf += c; i += 1
    emit(buf)

    for t, kw in runs:
        t = t.replace('\x00', '$').replace('{', '').replace('}', '')
        if t:
            _run(p, t, pt, italic=kw.get('italic', False), bold=kw.get('bold', False),
                 sup=kw.get('sup', False), sub=kw.get('sub', False))


def parse_estout(frag, ncols):
    frag = _protect_makecell(frag)
    gridcols = 1 + ncols
    rows, top_light, space_before, bot_partial = [], set(), set(), {}
    pending_cmid = []
    midrule_next = addspace_next = False
    for seg in re.split(r'\\\\', frag):
        s = seg.strip()
        if not s:
            continue
        while True:
            mc = re.match(r'\\cmidrule(\([^)]*\))?\{(\d+)-(\d+)\}\s*', s)
            if mc:
                pending_cmid.append((int(mc.group(2)) - 1, int(mc.group(3)) - 1))
                s = s[mc.end():]; continue
            m = re.match(r'\\(toprule|midrule|bottomrule|addlinespace)(\[[^\]]*\])?\s*', s)
            if m:
                if m.group(1) in ('midrule', 'toprule'):
                    midrule_next = True
                elif m.group(1) == 'addlinespace':
                    addspace_next = True
                s = s[m.end():]; continue
            break
        if pending_cmid and rows:
            bot_partial.setdefault(len(rows) - 1, []).extend(pending_cmid); pending_cmid = []
        if not s.strip():
            continue
        parsed = []
        for c in s.split('&'):
            cm = re.search(r'\\multicolumn\{(\d+)\}\{([^}]*)\}\{(.*)\}', c.strip(), re.S)
            if cm:
                parsed.append((cm.group(3).strip(), int(cm.group(1)), cm.group(2).strip().strip('|')))
            else:
                parsed.append((c.strip(), 1, None))
        ridx = len(rows)
        rows.append(parsed)
        if midrule_next:
            top_light.add(ridx); midrule_next = False
        if addspace_next:
            space_before.add(ridx); addspace_next = False
    if pending_cmid and rows:
        bot_partial.setdefault(len(rows) - 1, []).extend(pending_cmid)
    return rows, top_light, space_before, bot_partial, gridcols


def _setup_doc(doc):
    sec = doc.sections[0]
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Inches(1))
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(11)


def _cell(para, text, pt, numbermap):
    t = text.strip()
    if t == '$R^2$':
        _run(para, 'R', pt, italic=True); _run(para, '2', pt, italic=True, sup=True)
    elif t == '$N$':
        _run(para, 'N', pt, italic=True)
    else:
        inline_to_runs(para, text, pt, numbermap)


def normalize_note_text(s: str) -> str:
    """Normalize author-introduced spacing in table-notes text only.

    Collapses runs of two or more spaces to one, removes spaces before
    ``,`` ``.`` ``;`` ``:`` ``)``, removes a space before an *escaped* ``\\%``
    (a literal percent sign), and removes a space after ``(``. A bare
    (unescaped) ``%`` starts a LaTeX comment at the source level, so spacing
    before it is left untouched. Spacing around the relational operators
    ``<`` ``>`` ``=`` is likewise preserved, so tokens such as ``p < 0.01``
    survive intact.

    Examples::

        normalize_note_text('*** p < 0.01 , robust .  ( see notes )')
            -> '*** p < 0.01, robust. (see notes)'
        normalize_note_text('50 \\% level , foo % bar')
            -> '50\\% level, foo % bar'
    """
    if not s:
        return s
    # Remove spaces before closing/trailing punctuation (not the comment %).
    s = re.sub(r"[ \t]+([,.;:)])", r"\1", s)
    # Remove a space before an escaped percent sign (\%), leaving a bare % alone.
    s = re.sub(r"[ \t]+(\\%)", r"\1", s)
    # Remove spaces after an opening parenthesis.
    s = re.sub(r"(\()[ \t]+", r"\1", s)
    # Collapse runs of two or more spaces to a single space.
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s


def add_estout_table(doc, frag_path, caption, notes_tex, ncols,
                     numbermap=None, avail_in=6.4, data_pt=9.5):
    frag = open(frag_path, encoding='utf-8').read()
    rows, top_light, space_before, bot_partial, gridcols = parse_estout(frag, ncols)

    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption matches the table body font size.
    inline_to_runs(cp, caption, data_pt, numbermap)

    # content-aware column widths (like LaTeX: each column sized to its widest cell)
    colmax = [1] * gridcols
    for parsed in rows:
        ci = 0
        for (text, span, align) in parsed:
            if span == 1:
                colmax[ci] = max(colmax[ci], _plain_len(text))
            ci += span
    CHAR_IN, PAD = 0.066, 0.18
    widths_in = [max(0.45, min(2.3 if i == 0 else 1.5, colmax[i] * CHAR_IN + PAD))
                 for i in range(gridcols)]
    # Fill the full available width (scale up or down) so every table spans
    # margin-to-margin (landscape tables fill the wider sidewaystable page).
    _wsum = sum(widths_in)
    if _wsum > 0:
        sc = avail_in / _wsum
        widths_in = [w * sc for w in widths_in]
    tbl = doc.add_table(rows=len(rows), cols=gridcols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    grid = tbl._tbl.find(qn('w:tblGrid'))
    for i, gc in enumerate(grid.findall(qn('w:gridCol'))):
        gc.set(qn('w:w'), str(int(widths_in[i] * 1440)))
    for r, parsed in enumerate(rows):
        cidx = 0
        for (text, span, align) in parsed:
            cell = tbl.cell(r, cidx)
            if span > 1:
                cell = cell.merge(tbl.cell(r, cidx + span - 1))
            cell.width = Emu(int(sum(widths_in[cidx:cidx + span]) * 914400))
            para = cell.paragraphs[0]
            # Tight, uniformly single-spaced rows (drop the \addlinespace 4pt gap).
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if (cidx == 0 and span == 1)
                              else WD_ALIGN_PARAGRAPH.CENTER)
            _cell(para, text, data_pt, numbermap)
            cidx += span

    def row_border(r, edge, sz):
        for tc in tbl.rows[r]._tr.tc_lst:
            _border_tc(tc, edge, sz)

    row_border(0, 'top', HEAVY)
    row_border(len(rows) - 1, 'bottom', HEAVY)
    for r in top_light:
        row_border(r, 'top', LIGHT)
    for r, spans in bot_partial.items():
        if r + 1 >= len(rows):
            continue
        tcs = tbl.rows[r + 1]._tr.tc_lst
        for (c0, c1) in spans:
            for c in range(c0, c1 + 1):
                if c < len(tcs):
                    _border_tc(tcs[c], 'top', LIGHT)

    npar = doc.add_paragraph(); npar.alignment = WD_ALIGN_PARAGRAPH.LEFT
    inline_to_runs(npar, normalize_note_text(notes_tex), 8.5, numbermap)
    return tbl


def build(frag_path, caption, notes_tex, ncols, out_path, numbermap=None, data_pt=9.5):
    doc = Document()
    _setup_doc(doc)
    add_estout_table(doc, frag_path, caption, notes_tex, ncols, numbermap=numbermap, data_pt=data_pt)
    doc.save(out_path)
    print(f"-> {out_path}")


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description="Build one native DOCX table from an estout fragment.")
    parser.add_argument("fragment")
    parser.add_argument("out")
    parser.add_argument("--caption", default="")
    parser.add_argument("--notes", default="")
    parser.add_argument("--ncols", type=int, required=True)
    parser.add_argument("--font", default=FONT)
    ns = parser.parse_args()
    set_font(ns.font)
    build(ns.fragment, ns.caption, ns.notes, ns.ncols, ns.out)
