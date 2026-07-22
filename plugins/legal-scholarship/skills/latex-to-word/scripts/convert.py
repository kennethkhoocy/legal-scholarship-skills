# -*- coding: utf-8 -*-
"""Parameterized LaTeX to DOCX conversion with native Word tables.

The pipeline prepares a pandoc-friendly body, replaces table and figure floats
with stable placeholders, assembles native table and image objects back into the
DOCX, fixes footnote markers, and optionally renders a PDF/PNG preview.
"""
from __future__ import annotations

import argparse
import base64
import copy
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import traceback
import unicodedata
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Callable, Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

import full_tabular_to_docx as FT
from render_latex_env import render_env_to_png
from siunitx_expand import expand_siunitx_commands
import tex_table_to_docx as T
import toolcheck
from mhchem_unicode import contains_ce_commands, render_ce, replace_ce_commands


PAGEBREAK_TOKEN = "%%PAGEBREAK%%"
TABLE_RE = re.compile(r"%%TABLE:(.+?)%%")
FIG_RE = re.compile(r"%%FIGURE:(.+?)%%")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W = f"{{{W_NS}}}"
M = f"{{{M_NS}}}"

IMAGE_EXTENSIONS = [".png", ".jpg", ".jpeg", ".pdf", ".eps", ".tif", ".tiff"]
DEFAULT_CSL_NAME = "chicago-author-date"
NATBIB_CITATION_COMMANDS = [
    "cite",
    "citep",
    "citet",
    "citeauthor",
    "citeyear",
    "citeyearpar",
    "citealt",
    "citealp",
    "citetext",
    "nocite",
]
BIBLATEX_CITATION_COMMANDS = [
    "autocite",
    "textcite",
    "parencite",
    "smartcite",
    "footcite",
]
CITATION_COMMANDS = NATBIB_CITATION_COMMANDS + BIBLATEX_CITATION_COMMANDS


@dataclass
class Config:
    main_tex: Path
    aux: Path
    out: Path
    workdir: Path
    tables_dirs: list[Path]
    figures_dirs: list[Path]
    font: str
    csl: str | None = None
    render: bool = True
    user_macros: dict[str, tuple[int, str | None, str]] = field(default_factory=dict)


@dataclass
class CitationInfo:
    mode: str
    citation_command_counts: dict[str, int]
    bibliography_commands: list[str]
    bibliography_style: str
    bibliography_paths: list[str]
    missing_bibliography_paths: list[str]
    biblatex_package: bool = False
    printbibliography: bool = False
    addbibresource: bool = False
    manual_bibliography: bool = False
    manual_bibcite_count: int = 0
    warnings: list[str] = field(default_factory=list)
    default_csl: str = DEFAULT_CSL_NAME
    csl: str = ""


@dataclass
class LabelInfo:
    label: str
    number: str
    page: str
    kind: str


@dataclass
class BibciteInfo:
    key: str
    number: str
    year: str
    author: str


@dataclass
class TableRecord:
    kind: str
    env: str
    label: str
    placeholder: str
    start_line: int
    caption: str
    number: str
    engine: str
    notes: str = ""
    source_path: str = ""
    ncols: int | None = None
    align: str | None = None
    tabular_tex: str = ""
    landscape: bool = False
    detection_error: str = ""


@dataclass
class FigureRecord:
    kind: str
    env: str
    label: str
    placeholder: str
    start_line: int
    caption: str
    number: str
    image_path: str
    resolved_path: str = ""
    image_paths: list[str] = field(default_factory=list)
    resolved_paths: list[str] = field(default_factory=list)
    subcaptions: list[str] = field(default_factory=list)
    image_widths: list[float | None] = field(default_factory=list)
    transcript: str = ""


TABLE_FLOAT_ENVS = {"table", "table*", "sidewaystable", "longtable"}
ALGORITHM_FLOAT_ENVS = {"algorithm", "algorithm*", "algorithm2e"}
FIGURE_FLOAT_ENVS = {
    "figure",
    "figure*",
    "wrapfigure",
    "SCfigure",
    "SCfigure*",
    "floatingfigure",
    "sidewaysfigure",
    "sidewaysfigure*",
    *ALGORITHM_FLOAT_ENVS,
}

# Figure captions match the table body font size (the table_data_pt default, 11pt).
# Table captions are sized per-table inside the engines (== that table's data font).
FIGURE_CAPTION_PT = 11.0


@dataclass
class TheoremDefinition:
    env: str
    title: str
    counter: str
    numbered: bool = True


def skip_ws(text: str, pos: int) -> int:
    while pos < len(text) and text[pos].isspace():
        pos += 1
    return pos


def find_matching(text: str, open_pos: int, open_char: str, close_char: str) -> int:
    if open_pos >= len(text) or text[open_pos] != open_char:
        raise ValueError(f"Expected {open_char!r} at position {open_pos}")
    depth = 0
    escaped = False
    comment = False
    for pos in range(open_pos, len(text)):
        ch = text[pos]
        if comment:
            if ch == "\n":
                comment = False
            continue
        if escaped:
            escaped = False
            continue
        if ch == "\\":
            escaped = True
            continue
        if ch == "%":
            comment = True
            continue
        if ch == open_char:
            depth += 1
        elif ch == close_char:
            depth -= 1
            if depth == 0:
                return pos
    raise ValueError(f"No matching {close_char!r} for position {open_pos}")


def read_group(text: str, open_pos: int) -> tuple[str, int]:
    close_pos = find_matching(text, open_pos, "{", "}")
    return text[open_pos + 1 : close_pos], close_pos + 1


def read_bracket(text: str, open_pos: int) -> tuple[str, int]:
    close_pos = find_matching(text, open_pos, "[", "]")
    return text[open_pos + 1 : close_pos], close_pos + 1


def read_bracket_optional(text: str, pos: int) -> tuple[str | None, int]:
    pos = skip_ws(text, pos)
    if pos < len(text) and text[pos] == "[":
        return read_bracket(text, pos)
    return None, pos


def skip_optional_args(text: str, pos: int) -> int:
    pos = skip_ws(text, pos)
    while pos < len(text) and text[pos] == "[":
        pos = read_bracket(text, pos)[1]
        pos = skip_ws(text, pos)
    return pos


def command_regex(command: str) -> re.Pattern[str]:
    return re.compile(rf"\\{re.escape(command)}(?![A-Za-z@])")


def find_command_argument(text: str, command: str, start: int = 0) -> tuple[int, int, int, str] | None:
    pattern = command_regex(command)
    match = pattern.search(text, start)
    while match:
        pos = skip_optional_args(text, match.end())
        if pos < len(text) and text[pos] == "{":
            content, end = read_group(text, pos)
            return match.start(), pos, end, content
        match = pattern.search(text, match.end())
    return None


def iter_command_arguments(text: str, command: str) -> Iterable[tuple[int, int, int, str]]:
    pos = 0
    while True:
        found = find_command_argument(text, command, pos)
        if found is None:
            break
        yield found
        pos = found[2]


def read_command_groups(text: str, command: str, group_count: int, start: int = 0) -> tuple[int, int, list[str]] | None:
    pattern = command_regex(command)
    match = pattern.search(text, start)
    while match:
        pos = skip_optional_args(text, match.end())
        groups: list[str] = []
        ok = True
        for _ in range(group_count):
            pos = skip_ws(text, pos)
            if pos >= len(text) or text[pos] != "{":
                ok = False
                break
            content, pos = read_group(text, pos)
            groups.append(content)
        if ok:
            return match.start(), pos, groups
        match = pattern.search(text, match.end())
    return None


def strip_command_capture(text: str, command: str) -> tuple[str, list[str]]:
    captures: list[str] = []
    out: list[str] = []
    pos = 0
    pattern = command_regex(command)
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        arg_pos = skip_optional_args(text, match.end())
        if arg_pos >= len(text) or text[arg_pos] != "{":
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        content, end = read_group(text, arg_pos)
        captures.append(normalize_ws(content))
        out.append(text[pos : match.start()])
        pos = end
    return "".join(out), captures


def normalize_ws(text: str) -> str:
    text = re.sub(r"%[^\n]*(?:\n|$)", " ", text)
    text = text.replace("~", " ")
    return re.sub(r"\s+", " ", text).strip()


def latex_literal(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "%": r"\%",
        "_": r"\_",
        "&": r"\&",
        "#": r"\#",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def line_number_at(text: str, pos: int) -> int:
    return text.count("\n", 0, pos) + 1


def command_argument_from_block(block: str, command: str) -> str:
    found = find_command_argument(block, command)
    return found[3] if found else ""


def infer_label_kind(label: str) -> str:
    prefix = label.split(":", 1)[0]
    return {
        "tab": "table",
        "tbl": "table",
        "fig": "figure",
        "eq": "equation",
        "alg": "algorithm",
        "algorithm": "algorithm",
        "algocf": "algorithm",
        "sec": "section",
        "subsec": "subsection",
        "subsubsec": "subsubsection",
        "app": "appendix",
    }.get(prefix, "reference")


def default_cref_prefixes() -> dict[str, tuple[str, str]]:
    return {
        "tab": ("Table", "Tables"),
        "tbl": ("Table", "Tables"),
        "table": ("Table", "Tables"),
        "fig": ("Figure", "Figures"),
        "figure": ("Figure", "Figures"),
        "sec": ("Section", "Sections"),
        "section": ("Section", "Sections"),
        "subsec": ("Section", "Sections"),
        "subsection": ("Section", "Sections"),
        "subsubsec": ("Section", "Sections"),
        "subsubsection": ("Section", "Sections"),
        "eq": ("Equation", "Equations"),
        "equation": ("Equation", "Equations"),
        "alg": ("algorithm", "algorithms"),
        "algocf": ("algorithm", "algorithms"),
        "algorithm": ("algorithm", "algorithms"),
        "app": ("Appendix", "Appendices"),
        "appendix": ("Appendix", "Appendices"),
        "reference": ("Reference", "References"),
    }


def parse_cref_prefixes(source: str) -> dict[str, tuple[str, str]]:
    prefixes = default_cref_prefixes()
    begin_doc = re.search(r"\\begin\{document\}", source)
    preamble = source[: begin_doc.start()] if begin_doc else source
    pos = 0
    while True:
        found = read_command_groups(preamble, "crefname", 3, pos)
        if found is None:
            break
        _start, end, groups = found
        kind, singular, plural = [normalize_ws(item) for item in groups]
        if kind and singular and plural:
            prefixes[kind] = (singular, plural)
            short = {
                "table": "tab",
                "figure": "fig",
                "equation": "eq",
                "algorithm": "alg",
                "algocf": "algorithm",
                "section": "sec",
                "subsection": "subsec",
                "subsubsection": "subsubsec",
                "appendix": "app",
            }.get(kind)
            if short:
                prefixes[short] = (singular, plural)
        pos = end
    return prefixes


def parse_aux_labels(aux_path: Path) -> dict[str, LabelInfo]:
    regular: dict[str, LabelInfo] = {}
    cref_kinds: dict[str, str] = {}
    regular_re = re.compile(r"^\\newlabel\{([^}]+)\}\{\{([^{}]*)\}\{([^{}]*)\}")
    cref_re = re.compile(r"^\\newlabel\{([^}]+)@cref\}\{\{\[([^\]]+)\]")

    for line in aux_path.read_text(encoding="utf-8", errors="replace").splitlines():
        cref_match = cref_re.match(line)
        if cref_match:
            cref_kinds[cref_match.group(1)] = cref_match.group(2)
            continue
        match = regular_re.match(line)
        if not match:
            continue
        label = match.group(1)
        if label.endswith("@cref"):
            continue
        regular[label] = LabelInfo(
            label=label,
            number=match.group(2),
            page=match.group(3),
            kind=infer_label_kind(label),
        )

    for label, kind in cref_kinds.items():
        if label in regular:
            regular[label].kind = kind
    return regular


def display_kind(info: LabelInfo | None, label: str) -> str:
    if info is None:
        return infer_label_kind(label)
    if info.kind in {"section", "subsection", "subsubsection"} and re.match(r"^[A-Z]", info.number):
        return "appendix"
    return info.kind


def singular_plural(kind: str, prefixes: dict[str, tuple[str, str]]) -> tuple[str, str]:
    return prefixes.get(kind) or prefixes.get(infer_label_kind(kind), ("Reference", "References"))


def join_english(items: list[str]) -> str:
    if len(items) <= 1:
        return "".join(items)
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return f"{', '.join(items[:-1])}, and {items[-1]}"


def label_number(labels: dict[str, LabelInfo], label: str) -> str:
    info = labels.get(label)
    return info.number if info else f"??{label}??"


def render_cref(labels: dict[str, LabelInfo], raw_labels: str, prefixes: dict[str, tuple[str, str]]) -> str:
    parts = [part.strip() for part in raw_labels.split(",") if part.strip()]
    if not parts:
        return ""
    kinds = [display_kind(labels.get(label), label) for label in parts]
    nums = [label_number(labels, label) for label in parts]
    if len(set(kinds)) == 1:
        singular, plural = singular_plural(kinds[0], prefixes)
        prefix = plural if len(parts) > 1 else singular
        return f"{prefix} {join_english(nums)}"
    rendered = [f"{singular_plural(kind, prefixes)[0]} {num}" for kind, num in zip(kinds, nums)]
    return join_english(rendered)


def render_ref(labels: dict[str, LabelInfo], raw_labels: str) -> str:
    nums = [label_number(labels, part.strip()) for part in raw_labels.split(",") if part.strip()]
    return join_english(nums)


def render_eqref(labels: dict[str, LabelInfo], raw_labels: str) -> str:
    nums = [label_number(labels, part.strip()) for part in raw_labels.split(",") if part.strip()]
    return join_english([f"({num})" for num in nums])


def render_crefrange(
    labels: dict[str, LabelInfo],
    start_label: str,
    end_label: str,
    prefixes: dict[str, tuple[str, str]],
) -> str:
    start = start_label.strip()
    end = end_label.strip()
    start_kind = display_kind(labels.get(start), start)
    end_kind = display_kind(labels.get(end), end)
    start_num = label_number(labels, start)
    end_num = label_number(labels, end)
    if start_kind == end_kind:
        _singular, plural = singular_plural(start_kind, prefixes)
        return f"{plural} {start_num}-{end_num}"
    start_prefix = singular_plural(start_kind, prefixes)[0]
    end_prefix = singular_plural(end_kind, prefixes)[0]
    return f"{start_prefix} {start_num}-{end_prefix} {end_num}"


def resolve_cross_references(
    text: str,
    labels: dict[str, LabelInfo],
    prefixes: dict[str, tuple[str, str]],
) -> tuple[str, dict[str, int], list[str]]:
    pattern = re.compile(r"\\(Crefrange|crefrange|Cref|cref|eqref|autoref|ref)\s*\{")
    counts = {"Crefrange": 0, "crefrange": 0, "Cref": 0, "cref": 0, "eqref": 0, "autoref": 0, "ref": 0}
    missing: list[str] = []
    out: list[str] = []
    pos = 0

    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        open_pos = match.end() - 1
        try:
            raw_labels, end = read_group(text, open_pos)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue

        command = match.group(1)
        if command in {"Crefrange", "crefrange"}:
            second_pos = skip_ws(text, end)
            if second_pos >= len(text) or text[second_pos] != "{":
                out.append(text[pos : match.end()])
                pos = match.end()
                continue
            try:
                end_label, second_end = read_group(text, second_pos)
            except ValueError:
                out.append(text[pos : match.end()])
                pos = match.end()
                continue
            refs = [raw_labels.strip(), end_label.strip()]
            missing.extend([label for label in refs if label and label not in labels])
            replacement = render_crefrange(labels, raw_labels, end_label, prefixes)
            end = second_end
        else:
            refs = [part.strip() for part in raw_labels.split(",") if part.strip()]
            missing.extend([label for label in refs if label not in labels])
            if command == "eqref":
                replacement = render_eqref(labels, raw_labels)
            elif command == "ref":
                replacement = render_ref(labels, raw_labels)
            else:
                replacement = render_cref(labels, raw_labels, prefixes)

        counts[command] += 1
        out.append(text[pos : match.start()])
        out.append(replacement)
        pos = end

    return "".join(out), counts, sorted(set(missing))


def split_citation_keys(raw_keys: str) -> list[str]:
    return [normalize_ws(key) for key in raw_keys.split(",") if normalize_ws(key)]


def normalize_citation_note(note: str | None) -> str:
    if not note:
        return ""
    return latex_plain_text(apply_tex_accents(note)).strip()


def citation_number_for(key: str, bibcites: dict[str, BibciteInfo]) -> str:
    info = bibcites.get(key)
    return info.number if info and info.number else f"??{key}??"


def citation_author_for(key: str, bibcites: dict[str, BibciteInfo]) -> str:
    info = bibcites.get(key)
    return info.author if info and info.author else f"??{key}??"


def citation_year_for(key: str, bibcites: dict[str, BibciteInfo]) -> str:
    info = bibcites.get(key)
    return info.year if info and info.year else f"??{key}??"


def bracketed_numeric_citation(
    keys: list[str],
    bibcites: dict[str, BibciteInfo],
    prenote: str = "",
    postnote: str = "",
) -> str:
    parts = [citation_number_for(key, bibcites) for key in keys]
    inner = ", ".join(parts)
    if prenote:
        inner = f"{prenote} {inner}".strip()
    if postnote:
        inner = f"{inner}, {postnote}".strip(", ")
    return f"[{inner}]"


def textual_numeric_citation(
    keys: list[str],
    bibcites: dict[str, BibciteInfo],
    prenote: str = "",
    postnote: str = "",
) -> str:
    if len(keys) == 1:
        key = keys[0]
        return f"{citation_author_for(key, bibcites)} {bracketed_numeric_citation(keys, bibcites, prenote, postnote)}"
    return "; ".join(
        f"{citation_author_for(key, bibcites)} {bracketed_numeric_citation([key], bibcites)}"
        for key in keys
    )


def resolve_numeric_citations(
    text: str,
    bibcites: dict[str, BibciteInfo],
) -> tuple[str, dict[str, int], list[str]]:
    commands = sorted(NATBIB_CITATION_COMMANDS, key=len, reverse=True)
    pattern = re.compile(r"\\(" + "|".join(re.escape(command) for command in commands) + r")\*?(?![A-Za-z@])")
    counts: dict[str, int] = {}
    missing: list[str] = []
    out: list[str] = []
    pos = 0

    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break

        command = match.group(1)
        arg_pos = skip_ws(text, match.end())
        optional_args: list[str] = []
        malformed = False
        while arg_pos < len(text) and text[arg_pos] == "[" and len(optional_args) < 2:
            try:
                option, arg_pos = read_bracket(text, arg_pos)
            except ValueError:
                malformed = True
                break
            optional_args.append(option)
            arg_pos = skip_ws(text, arg_pos)
        if malformed or arg_pos >= len(text) or text[arg_pos] != "{":
            out.append(text[pos : match.end()])
            pos = match.end()
            continue

        try:
            raw_keys, end = read_group(text, arg_pos)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue

        out.append(text[pos : match.start()])
        counts[command] = counts.get(command, 0) + 1
        prenote = normalize_citation_note(optional_args[0]) if len(optional_args) == 2 else ""
        postnote = normalize_citation_note(optional_args[-1]) if optional_args else ""

        if command == "citetext":
            replacement = apply_tex_accents(raw_keys)
        elif command == "nocite":
            replacement = ""
        else:
            keys = split_citation_keys(raw_keys)
            missing.extend(key for key in keys if key not in bibcites)
            if command in {"citet", "citealt"}:
                replacement = textual_numeric_citation(keys, bibcites, prenote, postnote)
            elif command == "citeauthor":
                replacement = join_english([citation_author_for(key, bibcites) for key in keys])
            elif command == "citeyear":
                replacement = join_english([citation_year_for(key, bibcites) for key in keys])
            elif command == "citeyearpar":
                replacement = f"({join_english([citation_year_for(key, bibcites) for key in keys])})"
            elif command == "citealp":
                replacement = bracketed_numeric_citation(keys, bibcites, prenote, postnote).strip("[]")
            else:
                replacement = bracketed_numeric_citation(keys, bibcites, prenote, postnote)

        out.append(replacement)
        pos = end

    return "".join(out), counts, sorted(set(missing))


def strip_outer_braces(text: str) -> str:
    s = text.strip()
    if s.startswith("{"):
        try:
            inner, end = read_group(s, 0)
        except ValueError:
            return s
        if not s[end:].strip():
            return inner.strip()
    return s


def clean_notes(notes: str | None) -> str:
    if not notes:
        return ""
    s = notes.strip()
    s = re.sub(r"\\begin\{minipage\}\{[^{}]*\}", "", s)
    s = s.replace(r"\end{minipage}", "")
    s = re.sub(r"\\begin\{threeparttable\}", "", s)
    s = s.replace(r"\end{threeparttable}", "")
    s = s.strip()
    if s.startswith(r"\footnotesize"):
        s = s[len(r"\footnotesize") :].strip()
    s = strip_outer_braces(s)
    s = re.sub(r"^\\vspace\{[^{}]*\}", "", s).strip()
    return s


def extract_notes_from_block(block: str) -> str:
    m = re.search(r"\\begin\{tablenotes\}(?:\[[^\]]*\])?", block)
    if m:
        end_m = re.search(r"\\end\{tablenotes\}", block[m.end() :])
        if end_m:
            content = block[m.end() : m.end() + end_m.start()]
            return clean_notes(content)

    m = re.search(r"\\begin\{minipage\}\{[^{}]*\}", block)
    if not m:
        return ""
    end_m = re.search(r"\\end\{minipage\}", block[m.end() :])
    if not end_m:
        return ""
    content = block[m.end() : m.end() + end_m.start()]
    return clean_notes(content)


def contains_tabular(text: str) -> bool:
    return bool(re.search(r"\\begin\s*\{(?:tabular|tabulary|tabularx|longtable)\}", text))


def _matching_env_end(text: str, start: int, env: str) -> int | None:
    open_pat = re.compile(r"\\begin\s*\{" + re.escape(env) + r"\}")
    close_pat = re.compile(r"\\end\s*\{" + re.escape(env) + r"\}")
    depth = 1
    pos = start
    while pos < len(text):
        open_m = open_pat.search(text, pos)
        close_m = close_pat.search(text, pos)
        if close_m is None:
            return None
        if open_m is not None and open_m.start() < close_m.start():
            depth += 1
            pos = open_m.end()
            continue
        depth -= 1
        if depth == 0:
            return close_m.end()
        pos = close_m.end()
    return None


def extract_tabular_from_tex(text: str) -> str:
    m = re.search(r"\\begin\s*\{(tabular|tabulary|tabularx|longtable)\}", text)
    if not m:
        raise ValueError("No supported tabular environment found")
    env = m.group(1)
    end = _matching_env_end(text, m.end(), env)
    if end is None:
        raise ValueError(f"No closing \\end{{{env}}} found")
    return text[m.start() : end]


def normalize_latex_path(path_s: str) -> str:
    return path_s.strip().strip("{}").strip().replace("\\", "/")


def expand_jobname_token(path_s: str, main_tex: Path) -> str:
    stem = main_tex.stem
    return re.sub(r"\{\\jobname\}|(?:\\|/)jobname(?![A-Za-z@])", stem, path_s)


def with_suffix_candidates(path: Path, suffixes: Iterable[str]) -> list[Path]:
    if path.suffix:
        return [path]
    return [path.with_suffix(suffix) for suffix in suffixes]


def dedupe_paths(paths: Iterable[Path]) -> list[Path]:
    out: list[Path] = []
    seen: set[str] = set()
    for path in paths:
        key = str(path.resolve()) if path.exists() else str(path)
        if key not in seen:
            seen.add(key)
            out.append(path)
    return out


def strip_latex_comments(text: str) -> str:
    lines: list[str] = []
    for line in text.splitlines(keepends=True):
        newline = ""
        body = line
        if line.endswith("\r\n"):
            body = line[:-2]
            newline = "\r\n"
        elif line.endswith("\n") or line.endswith("\r"):
            body = line[:-1]
            newline = line[-1]

        cut = len(body)
        for idx, ch in enumerate(body):
            if ch != "%":
                continue
            slash_count = 0
            pos = idx - 1
            while pos >= 0 and body[pos] == "\\":
                slash_count += 1
                pos -= 1
            if slash_count % 2 == 0:
                cut = idx
                break
        lines.append(body[:cut] + newline)
    return "".join(lines)


def split_bibliography_arg(arg: str) -> list[str]:
    paths: list[str] = []
    for part in arg.split(","):
        raw = normalize_latex_path(part)
        if raw:
            paths.append(raw)
    return paths


def resolve_bib_path(path_s: str, main_tex: Path) -> Path:
    raw = normalize_latex_path(expand_jobname_token(path_s, main_tex))
    if not raw:
        raise FileNotFoundError("Empty bibliography path")
    path = Path(raw)
    base = path if path.is_absolute() else main_tex.parent / path
    candidates = [base] if base.suffix else [base.with_suffix(".bib")]
    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve()
    expected = candidates[0]
    raise FileNotFoundError(f"Could not resolve bibliography {path_s!r}; expected {expected}")


def parse_aux_bibliography(aux_path: Path) -> tuple[list[str], str]:
    if not aux_path.exists():
        return [], ""
    bib_args: list[str] = []
    style = ""
    for line in aux_path.read_text(encoding="utf-8", errors="replace").splitlines():
        data_match = re.match(r"\\bibdata\{([^}]*)\}", line)
        if data_match:
            bib_args.extend(split_bibliography_arg(data_match.group(1)))
            continue
        style_match = re.match(r"\\bibstyle\{([^}]*)\}", line)
        if style_match:
            style = normalize_ws(style_match.group(1))
    return bib_args, style


def top_level_brace_groups(text: str) -> list[str]:
    groups: list[str] = []
    pos = 0
    while True:
        pos = skip_ws(text, pos)
        if pos >= len(text) or text[pos] != "{":
            break
        try:
            content, pos = read_group(text, pos)
        except ValueError:
            break
        groups.append(content)
    return groups


ACCENT_COMBINING = {
    '"': "\u0308",
    "'": "\u0301",
    "`": "\u0300",
    "^": "\u0302",
    "~": "\u0303",
    "=": "\u0304",
    ".": "\u0307",
    "c": "\u0327",
    "k": "\u0328",
    "r": "\u030a",
    "u": "\u0306",
    "v": "\u030c",
    "H": "\u030b",
    "b": "\u0331",
    "d": "\u0323",
}


def apply_tex_accents(text: str) -> str:
    def repl_symbol(match: re.Match[str]) -> str:
        accent = match.group(1)
        letter = match.group(2)
        return unicodedata.normalize("NFC", letter + ACCENT_COMBINING.get(accent, ""))

    def repl_letter(match: re.Match[str]) -> str:
        accent = match.group(1)
        letter = match.group(2)
        return unicodedata.normalize("NFC", letter + ACCENT_COMBINING.get(accent, ""))

    s = re.sub(r"\\([\"'`^~=\.])\s*\{?([A-Za-z])\}?", repl_symbol, text)
    s = re.sub(r"\\([ckruvHbd])\s*\{([A-Za-z])\}", repl_letter, s)
    return s


def plain_bibcite_text(fragment: str) -> str:
    return latex_plain_text(apply_tex_accents(strip_outer_braces(fragment))).strip()


def parse_aux_bibcites(aux_path: Path) -> dict[str, BibciteInfo]:
    if not aux_path.exists():
        return {}
    text = aux_path.read_text(encoding="utf-8", errors="replace")
    result: dict[str, BibciteInfo] = {}
    pattern = command_regex("bibcite")
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            break
        key_pos = skip_ws(text, match.end())
        if key_pos >= len(text) or text[key_pos] != "{":
            pos = match.end()
            continue
        try:
            key, after_key = read_group(text, key_pos)
        except ValueError:
            pos = match.end()
            continue
        data_pos = skip_ws(text, after_key)
        if data_pos >= len(text) or text[data_pos] != "{":
            pos = after_key
            continue
        try:
            payload, end = read_group(text, data_pos)
        except ValueError:
            pos = after_key
            continue
        groups = top_level_brace_groups(payload)
        if groups:
            number = plain_bibcite_text(groups[0])
            year = plain_bibcite_text(groups[1]) if len(groups) > 1 else ""
            author = plain_bibcite_text(groups[2]) if len(groups) > 2 else ""
            clean_key = normalize_ws(key)
            result[clean_key] = BibciteInfo(key=clean_key, number=number, year=year, author=author)
        pos = end
    return result


def has_thebibliography(text: str) -> bool:
    return bool(re.search(r"\\begin\{thebibliography\}", text))


def sibling_bbl_path(config: Config) -> Path:
    return config.main_tex.with_suffix(".bbl")


def package_loaded(source: str, package: str) -> bool:
    for _start, _arg_pos, _end, arg in iter_command_arguments(source, "usepackage"):
        packages = [item.strip() for item in arg.split(",")]
        if package in packages:
            return True
    return False


def command_counts(source: str, commands: Iterable[str]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for command in commands:
        count = len(command_regex(command).findall(source))
        if count:
            counts[command] = count
    return counts


def detect_citation_mode(source: str, config: Config) -> CitationInfo:
    clean = strip_latex_comments(source)
    citation_counts = command_counts(clean, CITATION_COMMANDS)
    bibliography_args: list[str] = []
    bibliography_commands: list[str] = []
    warnings: list[str] = []

    for _start, _arg_pos, _end, arg in iter_command_arguments(clean, "bibliography"):
        bibliography_commands.append("bibliography")
        bibliography_args.extend(split_bibliography_arg(arg))
    for _start, _arg_pos, _end, arg in iter_command_arguments(clean, "addbibresource"):
        bibliography_commands.append("addbibresource")
        bibliography_args.extend(split_bibliography_arg(arg))

    styles = [normalize_ws(arg) for _s, _a, _e, arg in iter_command_arguments(clean, "bibliographystyle")]
    aux_bib_args, aux_style = parse_aux_bibliography(config.aux)
    bibliography_args.extend(aux_bib_args)
    bibliography_style = styles[-1] if styles else aux_style

    resolved: list[Path] = []
    missing: list[str] = []
    for raw in bibliography_args:
        try:
            resolved.append(resolve_bib_path(raw, config.main_tex))
        except FileNotFoundError as exc:
            missing.append(str(exc))
    resolved = dedupe_paths(resolved)

    has_biblatex_package = package_loaded(clean, "biblatex")
    has_printbibliography = bool(command_regex("printbibliography").search(clean))
    has_addbibresource = "addbibresource" in bibliography_commands
    has_biblatex_cites = any(command in citation_counts for command in BIBLATEX_CITATION_COMMANDS)
    citations_present = bool(citation_counts)
    manual_bibliography = has_thebibliography(clean) or sibling_bbl_path(config).exists()
    manual_bibcites = parse_aux_bibcites(config.aux)
    biblatex_signal = has_biblatex_package or has_addbibresource or has_printbibliography or has_biblatex_cites
    natbib_signal = bool(bibliography_args or styles or aux_style)

    if biblatex_signal and resolved and (citations_present or has_addbibresource or has_printbibliography):
        mode = "biblatex"
    elif citations_present and resolved and natbib_signal:
        mode = "natbib-bibtex"
    elif citations_present and manual_bibliography:
        mode = "manual-bibliography"
    elif citations_present and manual_bibcites:
        mode = "manual-bibliography"
        manual_bibliography = True
        warnings.append("Citations were resolved from .aux \\bibcite entries without an external .bib file.")
    elif citations_present and resolved:
        mode = "natbib-bibtex"
    elif citations_present:
        mode = "citation-best-effort"
        detail = "; ".join(missing) if missing else "no bibliography source was found"
        warnings.append(f"Citations were detected but no bibliography could be resolved; converting citation commands best-effort ({detail}).")
    else:
        mode = "footnote-only"

    if mode in {"biblatex", "natbib-bibtex"} and missing:
        warnings.append("One or more bibliography paths could not be resolved and were omitted: " + "; ".join(missing))

    return CitationInfo(
        mode=mode,
        citation_command_counts=citation_counts,
        bibliography_commands=sorted(set(bibliography_commands)),
        bibliography_style=bibliography_style,
        bibliography_paths=[str(path) for path in resolved],
        missing_bibliography_paths=missing,
        biblatex_package=has_biblatex_package,
        printbibliography=has_printbibliography,
        addbibresource=has_addbibresource,
        manual_bibliography=manual_bibliography,
        manual_bibcite_count=len(manual_bibcites),
        warnings=warnings,
        csl=config.csl or DEFAULT_CSL_NAME,
    )


def resolve_tex_path(path_s: str, config: Config) -> Path:
    raw = normalize_latex_path(path_s)
    if not raw:
        raise FileNotFoundError("Empty TeX path")
    p = Path(raw)
    bases = [Path()] if p.is_absolute() else [config.main_tex.parent, *config.tables_dirs]
    candidates: list[Path] = []
    for base in bases:
        base_path = p if p.is_absolute() else base / p
        candidates.extend(with_suffix_candidates(base_path, [".tex"]))
    for candidate in dedupe_paths(candidates):
        if candidate.exists():
            return candidate.resolve()
    raise FileNotFoundError(f"Could not resolve TeX input {path_s!r}")


def resolve_graphic_path(path_s: str, config: Config) -> Path:
    raw = normalize_latex_path(path_s)
    if not raw:
        raise FileNotFoundError("Empty graphic path")
    p = Path(raw)
    bases = [Path()] if p.is_absolute() else [config.main_tex.parent, *config.figures_dirs]
    candidates: list[Path] = []
    for base in bases:
        base_path = p if p.is_absolute() else base / p
        candidates.extend(with_suffix_candidates(base_path, IMAGE_EXTENSIONS))
    for candidate in dedupe_paths(candidates):
        if candidate.exists():
            return candidate.resolve()
    raise FileNotFoundError(f"Could not resolve graphic {path_s!r}")


def input_paths_from_block(block: str) -> list[str]:
    paths = [arg for _start, _arg_pos, _end, arg in iter_command_arguments(block, "input")]
    paths.extend(m.group(1).rstrip() for m in re.finditer(r"\\@@input\s+([^\s%]+)", block))
    return paths


def expand_latex_inputs(text: str, config: Config, depth: int = 0) -> str:
    if depth > 5:
        return text

    def repl_input(match: re.Match[str]) -> str:
        raw = match.group(1)
        try:
            path = resolve_tex_path(raw, config)
            return expand_latex_inputs(path.read_text(encoding="utf-8", errors="replace"), config, depth + 1)
        except Exception:
            return match.group(0)

    text = re.sub(r"\\@@input\s+([^\s%]+)", repl_input, text)

    def expand_command(source: str, command: str, wrap_clearpage: bool) -> str:
        out: list[str] = []
        pos = 0
        for start, _arg_pos, end, arg in iter_command_arguments(source, command):
            out.append(source[pos:start])
            try:
                path = resolve_tex_path(arg, config)
                body = expand_latex_inputs(path.read_text(encoding="utf-8", errors="replace"), config, depth + 1)
                # \include semantics: LaTeX issues \clearpage around the included body.
                out.append(f"\\clearpage\n{body}\n\\clearpage" if wrap_clearpage else body)
            except Exception:
                out.append(source[start:end])
            pos = end
        out.append(source[pos:])
        return "".join(out)

    text = expand_command(text, "input", wrap_clearpage=False)
    text = expand_command(text, "include", wrap_clearpage=True)

    def repl_verbatim_input(match: re.Match[str]) -> str:
        # \VerbatimInput (fancyvrb/fvextra): inline the referenced file as a
        # plain verbatim block so pandoc carries it into the docx as a code
        # block; dropping the command silently would lose the content.
        try:
            path = resolve_tex_path(match.group(1), config)
            body = path.read_text(encoding="utf-8", errors="replace")
            return "\\begin{verbatim}\n" + body + "\n\\end{verbatim}"
        except Exception:
            return ""

    text = re.sub(
        r"\\(?:Batch)?VerbatimInput\*?\s*(?:\[[^\]]*\])?\s*\{([^}]+)\}",
        repl_verbatim_input,
        text,
    )
    return text


def detect_table_record(
    env: str,
    label: str,
    caption: str,
    number: str,
    placeholder: str,
    start_line: int,
    block: str,
    config: Config,
) -> TableRecord:
    notes = extract_notes_from_block(block)
    record = TableRecord(
        kind="table",
        env=env,
        label=label,
        placeholder=placeholder,
        start_line=start_line,
        caption=normalize_ws(caption),
        number=number,
        engine="",
        notes=notes,
        landscape=env == "sidewaystable",
    )
    try:
        if contains_tabular(block):
            record.engine = "full_tabular"
            expanded_block = expand_newcommands(expand_latex_inputs(block, config), config.user_macros)
            record.tabular_tex = extract_tabular_from_tex(expanded_block)
            return record

        est = read_command_groups(block, "estauto", 3)
        if est is not None:
            _start, _end, groups = est
            frag_arg, ncols_s, align = groups
            frag_path = resolve_tex_path(frag_arg, config)
            frag = expand_newcommands(
                expand_latex_inputs(frag_path.read_text(encoding="utf-8", errors="replace"), config),
                config.user_macros,
            )
            record.source_path = str(frag_path)
            try:
                record.ncols = int(ncols_s.strip())
            except ValueError as exc:
                raise ValueError(f"Invalid estauto column count {ncols_s!r}") from exc
            record.align = align.strip()
            if contains_tabular(frag):
                record.engine = "full_tabular"
                record.tabular_tex = frag
                record.notes = record.notes or extract_notes_from_block(frag)
            else:
                record.engine = "estout"
            return record

        for input_path_s in input_paths_from_block(block):
            input_path = resolve_tex_path(input_path_s, config)
            input_tex = expand_newcommands(
                expand_latex_inputs(input_path.read_text(encoding="utf-8", errors="replace"), config),
                config.user_macros,
            )
            if contains_tabular(input_tex):
                record.engine = "full_tabular"
                record.source_path = str(input_path)
                record.tabular_tex = input_tex
                record.notes = record.notes or extract_notes_from_block(input_tex)
                return record

        record.detection_error = "No supported table content found"
        return record
    except Exception as exc:
        record.detection_error = f"{type(exc).__name__}: {exc}"
        return record


def environment_spans(text: str, env: str) -> list[tuple[int, int, str]]:
    spans: list[tuple[int, int, str]] = []
    begin_re = re.compile(rf"\\begin\{{{re.escape(env)}\}}(?:\[[^\]]*\])?(?:\{{[^{{}}]*\}})?")
    end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
    pos = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            break
        end = end_re.search(text, begin.end())
        if not end:
            break
        spans.append((begin.start(), end.end(), text[begin.start() : end.end()]))
        pos = end.end()
    return spans


def remove_environment_blocks(text: str, env_names: Iterable[str]) -> str:
    spans: list[tuple[int, int]] = []
    for env in env_names:
        spans.extend((start, end) for start, end, _block in environment_spans(text, env))
    if not spans:
        return text
    spans.sort()
    out: list[str] = []
    pos = 0
    for start, end in spans:
        if start < pos:
            continue
        out.append(text[pos:start])
        pos = end
    out.append(text[pos:])
    return "".join(out)


def figure_outer_caption(block: str) -> str:
    outer = remove_environment_blocks(block, ["subfigure", "subtable"])
    return command_argument_from_block(outer, "caption")


def figure_outer_label(block: str) -> str:
    outer = remove_environment_blocks(block, ["subfigure", "subtable"])
    caption = find_command_argument(outer, "caption")
    if caption is not None:
        label = find_command_argument(outer, "label", caption[2])
        if label is not None:
            return label[3]
    labels = list(iter_command_arguments(outer, "label"))
    return labels[-1][3] if labels else ""


def figure_subcaptions(block: str) -> dict[str, str]:
    captions: dict[str, str] = {}
    for env in ["subfigure", "subtable"]:
        for _start, _end, subblock in environment_spans(block, env):
            caption = command_argument_from_block(subblock, "caption")
            for _g_start, _arg_pos, _g_end, graphic in iter_command_arguments(subblock, "includegraphics"):
                captions[normalize_latex_path(graphic)] = normalize_ws(caption)
    return captions


def figure_graphics(block: str) -> list[str]:
    paths: list[str] = []
    for raw, _width in figure_graphic_specs(block):
        if raw:
            paths.append(raw)
    return paths


def latex_graphic_width_to_inches(width_s: str) -> float | None:
    s = width_s.strip()
    rel = re.match(r"([0-9.]+)?\s*\\(?:line|text)width$", s)
    if rel:
        factor = float(rel.group(1)) if rel.group(1) else 1.0
        return factor * 6.4
    m = re.match(r"([0-9.]+)\s*(in|cm|mm|pt)$", s)
    if not m:
        return None
    value = float(m.group(1))
    unit = m.group(2)
    if unit == "in":
        return value
    if unit == "cm":
        return value / 2.54
    if unit == "mm":
        return value / 25.4
    if unit == "pt":
        return value / 72.27
    return None


def graphic_width_from_options(options: list[str]) -> float | None:
    for opt in options:
        match = re.search(r"(?:^|,)\s*width\s*=\s*([^,\]]+)", opt)
        if match:
            parsed = latex_graphic_width_to_inches(match.group(1))
            if parsed is not None:
                return parsed
    return None


def figure_graphic_specs(block: str) -> list[tuple[str, float | None]]:
    specs: list[tuple[str, float | None]] = []
    pattern = command_regex("includegraphics")
    pos = 0
    while True:
        match = pattern.search(block, pos)
        if not match:
            break
        cursor = skip_ws(block, match.end())
        options: list[str] = []
        while cursor < len(block) and block[cursor] == "[":
            opt, cursor = read_bracket(block, cursor)
            options.append(opt)
            cursor = skip_ws(block, cursor)
        if cursor < len(block) and block[cursor] == "{":
            arg, end = read_group(block, cursor)
            specs.append((normalize_latex_path(arg), graphic_width_from_options(options)))
            pos = end
        else:
            pos = match.end()
    return specs


LATEX_ENV_RENDER_RE = re.compile(r"\\begin\{(?:tikzpicture|pgfpicture|axis|algorithm\*?|algorithm2e)\}")
LATEX_GRAPHIC_RENDER_RE = re.compile(r"\\begin\{(?:tikzpicture|pgfpicture|axis)\}")


def source_preamble(source: str) -> str:
    begin_doc = re.search(r"\\begin\{document\}", source)
    return source[: begin_doc.start()] if begin_doc else source


def command_fragment_from_block(block: str, command: str) -> str:
    found = find_command_argument(block, command)
    if found is None:
        return ""
    start, _arg_pos, end, _content = found
    return block[start:end].strip()


def outer_figure_fragments_from_body(body: str) -> list[str]:
    fragments: list[str] = []
    caption = find_command_argument(body, "caption")
    if caption is not None:
        fragments.append(body[caption[0] : caption[2]].strip())
        label = find_command_argument(body, "label", caption[2])
        if label is not None:
            fragments.append(body[label[0] : label[2]].strip())
        return fragments
    labels = list(iter_command_arguments(body, "label"))
    if labels:
        start, _arg_pos, end, _content = labels[-1]
        fragments.append(body[start:end].strip())
    return fragments


def strip_command_fragments(block: str, commands: Iterable[str]) -> str:
    spans: list[tuple[int, int]] = []
    for command in commands:
        pos = 0
        while True:
            found = find_command_argument(block, command, pos)
            if found is None:
                break
            start, _arg_pos, end, _content = found
            spans.append((start, end))
            pos = end
    if not spans:
        return block
    spans.sort()
    out: list[str] = []
    pos = 0
    for start, end in spans:
        if start < pos:
            continue
        out.append(block[pos:start])
        pos = end
    out.append(block[pos:])
    return "".join(out)


def algorithm_search_transcript(body: str) -> str:
    text = strip_latex_comments(body)
    text = re.sub(r"\\Comment\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r" \1 ", text)
    replacements = {
        "KwData": "Data",
        "KwResult": "Result",
        "ForEach": "foreach",
        "For": "for",
        "uIf": "if then",
        "If": "if then",
        "ElseIf": "else if then",
        "Else": "else",
        "KwRet": "return",
        "Return": "return",
        "KwTo": "to",
        "KwSty": "",
        "DontPrintSemicolon": "",
    }
    for command, word in replacements.items():
        text = re.sub(rf"\\{re.escape(command)}(?![A-Za-z@])", f" {word} ", text)
    text = re.sub(r"\\(?:leftarrow|rightarrow|geq|leq|neq|in|cup|setminus|varnothing|emptyset|top|mathrm|mathcal)\b", " ", text)
    text = re.sub(r"\\[A-Za-z@]+\*?", " ", text)
    text = re.sub(r"[{}$^_\\[\];,()|/+*=<>-]+", " ", text)
    text = text.replace("≥", " ").replace("≤", " ").replace("∅", " ")
    line_numbers = " ".join(str(index) for index in range(1, 41))
    return normalize_ws(f"{text} then {line_numbers}")


def algorithm_transcript_comment(transcript: str) -> str:
    if not transcript:
        return ""
    encoded = base64.b64encode(transcript.encode("utf-8")).decode("ascii")
    return f"% DOCX_ALGORITHM_TRANSCRIPT:{encoded}\n"


def extract_algorithm_transcript(block: str) -> str:
    match = re.search(r"%\s*DOCX_ALGORITHM_TRANSCRIPT:([A-Za-z0-9+/=]+)", block)
    if not match:
        return ""
    try:
        return base64.b64decode(match.group(1)).decode("utf-8")
    except Exception:
        return ""


def rendered_includegraphics(path: Path, width: str | None = None) -> str:
    width_s = normalize_ws(width or r"\linewidth") or r"\linewidth"
    return rf"\includegraphics[width={width_s}]{{{path.resolve().as_posix()}}}"


def rendered_figure_transcript(body: str) -> str:
    body = re.sub(r"\\rule\s*\{[^{}]*\}\s*\{[^{}]*\}", " ", body)
    body = re.sub(r"\\begin\{minipage\}\s*(?:\[[^\]]*\])?\s*\{[^{}]*\}", " ", body)
    body = re.sub(r"\\end\{minipage\}", " ", body)
    body = re.sub(r"\\\\(?:\[[^\]]*\])?", " ", body)
    text = latex_plain_text(body)
    text = re.sub(r"\$+", " ", text)
    return normalize_ws(text)


def resizebox_wrapper_for_span(block: str, span_start: int, span_end: int) -> tuple[int, int, str] | None:
    best: tuple[int, int, str] | None = None
    pattern = command_regex("resizebox")
    for match in pattern.finditer(block):
        if match.start() > span_start:
            break
        try:
            first_pos = skip_ws(block, match.end())
            if first_pos >= len(block) or block[first_pos] != "{":
                continue
            width, after_width = read_group(block, first_pos)
            second_pos = skip_ws(block, after_width)
            if second_pos >= len(block) or block[second_pos] != "{":
                continue
            _height, after_height = read_group(block, second_pos)
            third_pos = skip_ws(block, after_height)
            if third_pos >= len(block) or block[third_pos] != "{":
                continue
            third_close = find_matching(block, third_pos, "{", "}")
        except ValueError:
            continue
        content_start = third_pos + 1
        content_end = third_close
        if content_start <= span_start and span_end <= content_end:
            before = normalize_ws(block[content_start:span_start])
            after = normalize_ws(block[span_end:content_end])
            if before or after:
                continue
            width_s = normalize_ws(width)
            best = (match.start(), third_close + 1, width_s if width_s != "!" else "")
    return best


def renderable_graphic_spans(block: str) -> list[tuple[int, int, str]]:
    spans: list[tuple[int, int, str]] = []
    for env in ("tikzpicture", "pgfpicture", "axis"):
        spans.extend((start, end, env) for start, end, _source in environment_spans(block, env))
    spans.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    selected: list[tuple[int, int, str]] = []
    covered_until = -1
    for start, end, env in spans:
        if start < covered_until:
            continue
        selected.append((start, end, env))
        covered_until = end
    return selected


def render_graphics_in_figure_block(block: str, preamble: str, config: Config) -> str:
    if figure_graphic_specs(block) or not LATEX_GRAPHIC_RENDER_RE.search(block):
        return block
    spans = renderable_graphic_spans(block)
    if not spans:
        return block

    out: list[str] = []
    pos = 0
    for start, end, env in spans:
        wrapper = resizebox_wrapper_for_span(block, start, end)
        replace_start = wrapper[0] if wrapper else start
        replace_end = wrapper[1] if wrapper else end
        if replace_start < pos:
            continue
        env_source = block[start:end]
        if env == "axis":
            env_source = "\\begin{tikzpicture}\n" + env_source + "\n\\end{tikzpicture}\n"
        png_path = render_env_to_png(env_source, preamble, config, run_command=run_command)
        out.append(block[pos:replace_start])
        if png_path is not None:
            width = wrapper[2] if wrapper and wrapper[2] else r"\linewidth"
            out.append(rendered_includegraphics(png_path, width))
        else:
            out.append("% LaTeX graphic rendering failed; omitted.\n")
        pos = replace_end
    out.append(block[pos:])
    return "".join(out)


def render_algorithm_float_block(block: str, preamble: str, config: Config) -> str:
    if figure_graphic_specs(block):
        return block
    begin = re.match(r"(\\begin\{algorithm\*?\}(?:\[[^\]]*\])?)", block)
    if not begin:
        begin = re.match(r"(\\begin\{algorithm2e\}(?:\[[^\]]*\])?)", block)
    if not begin:
        return block
    end_match = re.search(r"\\end\{(?:algorithm\*?|algorithm2e)\}\s*$", block)
    if not end_match:
        return block
    body = block[begin.end() : end_match.start()]
    body_without_caption = strip_command_fragments(body, ["caption", "label"]).strip()
    if not body_without_caption:
        return block

    env_source = "\\begin{algorithm}[H]\n" + body_without_caption + "\n\\end{algorithm}\n"
    png_path = render_env_to_png(env_source, preamble, config, run_command=run_command)
    if png_path is None:
        return block

    fragments = [fragment for fragment in (command_fragment_from_block(body, "caption"), command_fragment_from_block(body, "label")) if fragment]
    replacement = [
        begin.group(1),
        "\n",
        *[fragment + "\n" for fragment in fragments],
        r"\centering",
        "\n",
        rendered_includegraphics(png_path, r"\linewidth"),
        "\n",
        algorithm_transcript_comment(algorithm_search_transcript(body_without_caption)),
        block[end_match.start() : end_match.end()],
    ]
    return "".join(replacement)


def figure_render_body(body: str) -> str:
    stripped = strip_command_fragments(body, ["caption", "label"])
    stripped = re.sub(r"\\(?:centering|ContinuedFloat)\b", "", stripped)
    stripped = re.sub(r"\\caption\*?(?:\[[^\]]*\])?\{[^{}]*\}", "", stripped)
    stripped = strip_latex_comments(stripped)
    return stripped.strip()


def render_imageless_figure_block(block: str, preamble: str, config: Config) -> str:
    if figure_graphic_specs(block):
        return block
    begin = re.match(r"(\\begin\{([^{}]+)\}(?:\[[^\]]*\])?(?:\{[^{}]*\})?)", block)
    if not begin:
        return block
    env = begin.group(2)
    end_match = re.search(rf"\\end\{{{re.escape(env)}\}}\s*$", block)
    if not end_match:
        return block
    body = block[begin.end() : end_match.start()]
    render_body = figure_render_body(body)
    if not render_body:
        return block

    env_source = "\\begin{figure}[H]\n\\centering\n" + render_body + "\n\\end{figure}\n"
    png_path = render_env_to_png(env_source, preamble, config, run_command=run_command)
    if png_path is None:
        return block

    fragments = outer_figure_fragments_from_body(body)
    replacement = [
        begin.group(1),
        "\n",
        *[fragment + "\n" for fragment in fragments],
        r"\centering",
        "\n",
        rendered_includegraphics(png_path, r"\linewidth"),
        "\n",
        algorithm_transcript_comment(rendered_figure_transcript(render_body)),
        block[end_match.start() : end_match.end()],
    ]
    return "".join(replacement)


def render_latex_environment_images(source: str, config: Config) -> str:
    preamble = source_preamble(source)
    float_env_pattern = "|".join(re.escape(env) for env in sorted(FIGURE_FLOAT_ENVS, key=len, reverse=True))
    begin_re = re.compile(rf"\\begin\{{({float_env_pattern})\}}(?:\[[^\]]*\])?")
    out: list[str] = []
    pos = 0

    while True:
        match = begin_re.search(source, pos)
        if not match:
            out.append(source[pos:])
            break
        env = match.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end_match = end_re.search(source, match.end())
        if not end_match:
            out.append(source[pos:])
            break
        block = source[match.start() : end_match.end()]
        out.append(source[pos : match.start()])
        if env in ALGORITHM_FLOAT_ENVS:
            out.append(render_algorithm_float_block(block, preamble, config))
        else:
            rendered = render_graphics_in_figure_block(block, preamble, config)
            if rendered == block and not LATEX_GRAPHIC_RENDER_RE.search(block):
                rendered = render_imageless_figure_block(block, preamble, config)
            out.append(rendered)
        pos = end_match.end()
    return "".join(out)


def graphicspath_dirs(source: str, main_dir: Path) -> list[Path]:
    dirs: list[Path] = []
    pos = 0
    while True:
        found = find_command_argument(source, "graphicspath", pos)
        if found is None:
            break
        _start, _arg_pos, end, content = found
        for m in re.finditer(r"\{([^{}]+)\}", content):
            raw = normalize_latex_path(m.group(1))
            if raw:
                p = Path(raw)
                dirs.append(p if p.is_absolute() else main_dir / p)
        pos = end
    return dirs


def auto_detect_dirs(source: str, main_tex: Path) -> tuple[list[Path], list[Path]]:
    main_dir = main_tex.parent
    table_dirs: list[Path] = []
    figure_dirs: list[Path] = []

    for _start, _arg_pos, _end, arg in iter_command_arguments(source, "input"):
        p = Path(normalize_latex_path(arg))
        if p.parent != Path("."):
            table_dirs.append(main_dir / p.parent)
    for m in re.finditer(r"\\@@input\s+([^\s%]+)", source):
        p = Path(normalize_latex_path(m.group(1)))
        if p.parent != Path("."):
            table_dirs.append(main_dir / p.parent)
    pos = 0
    while True:
        est = read_command_groups(source, "estauto", 3, pos)
        if est is None:
            break
        _start, end, groups = est
        p = Path(normalize_latex_path(groups[0]))
        if p.parent != Path("."):
            table_dirs.append(main_dir / p.parent)
        pos = end

    figure_dirs.extend(graphicspath_dirs(source, main_dir))
    for _start, _arg_pos, _end, arg in iter_command_arguments(source, "includegraphics"):
        p = Path(normalize_latex_path(arg))
        if p.parent != Path("."):
            figure_dirs.append(main_dir / p.parent)

    for name in ("tables", "Tables"):
        p = main_dir / name
        if p.exists():
            table_dirs.append(p)
    for name in ("figures", "Figures"):
        p = main_dir / name
        if p.exists():
            figure_dirs.append(p)

    return dedupe_paths(table_dirs), dedupe_paths(figure_dirs)


def neutralize_author_thanks(text: str) -> tuple[str, list[str]]:
    notes: list[str] = []
    out: list[str] = []
    pos = 0
    pattern = command_regex("author")
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        arg_pos = skip_optional_args(text, match.end())
        if arg_pos >= len(text) or text[arg_pos] != "{":
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        content, end = read_group(text, arg_pos)
        cleaned, captured = strip_command_capture(content, "thanks")
        notes.extend(captured)
        out.append(text[pos:arg_pos])
        out.append("{")
        out.append(cleaned)
        out.append("}")
        pos = end

    text = "".join(out)
    if notes:
        note_lines = "\n".join(rf"\noindent\textit{{Author note:}} {note}\par" for note in notes)
        text = text.replace(r"\maketitle", rf"\maketitle" + "\n\n" + note_lines, 1)
    return text, notes


def preamble_command_declarations(source: str, commands: Iterable[str]) -> list[str]:
    begin_doc = re.search(r"\\begin\{document\}", source)
    preamble = source[: begin_doc.start()] if begin_doc else source
    declarations: list[str] = []
    for command in commands:
        found = find_command_argument(preamble, command)
        if found is None:
            continue
        _start, _arg_pos, _end, content = found
        declarations.append(rf"\{command}{{{content}}}")
    return declarations


def strip_source_preamble(text: str) -> str:
    begin_doc = re.search(r"\\begin\{document\}", text)
    if not begin_doc:
        raise RuntimeError("Source does not contain \\begin{document}.")
    declarations = preamble_command_declarations(text, ["title", "author", "date"])
    minimal_lines = [
        r"\documentclass[12pt]{article}",
        r"\usepackage{amsmath}",
        r"\usepackage{amssymb}",
        r"\usepackage{hyperref}",
        r"\newcommand{\sym}[1]{\textsuperscript{#1}}",
        *declarations,
        "",
    ]
    minimal_preamble = "\n".join(minimal_lines)
    return minimal_preamble + text[begin_doc.start() :]


def convert_sym_macros(text: str) -> str:
    out: list[str] = []
    pos = 0
    while True:
        found = find_command_argument(text, "sym", pos)
        if found is None:
            out.append(text[pos:])
            break
        start, _arg_pos, end, content = found
        out.append(text[pos:start])
        out.append(rf"\textsuperscript{{{content}}}")
        pos = end
    return "".join(out)


def convert_href_commands(text: str) -> str:
    pattern = command_regex("href")
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        first_pos = skip_optional_args(text, match.end())
        if first_pos >= len(text) or text[first_pos] != "{":
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        _target, after_first = read_group(text, first_pos)
        second_pos = skip_ws(text, after_first)
        if second_pos >= len(text) or text[second_pos] != "{":
            out.append(text[pos:after_first])
            pos = after_first
            continue
        visible, end = read_group(text, second_pos)
        out.append(text[pos : match.start()])
        out.append(visible)
        pos = end
    return "".join(out)


def unwrap_rotatebox_commands(text: str) -> str:
    pattern = command_regex("rotatebox")
    out: list[str] = []
    pos = 0
    changed = False
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        try:
            angle_pos = skip_optional_args(text, match.end())
            if angle_pos >= len(text) or text[angle_pos] != "{":
                out.append(text[pos : match.end()])
                pos = match.end()
                continue
            _angle, after_angle = read_group(text, angle_pos)
            body_pos = skip_ws(text, after_angle)
            if body_pos >= len(text) or text[body_pos] != "{":
                out.append(text[pos:after_angle])
                pos = after_angle
                continue
            body, end = read_group(text, body_pos)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        out.append(text[pos : match.start()])
        out.append(body)
        pos = end
        changed = True
    return unwrap_rotatebox_commands("".join(out)) if changed and pattern.search("".join(out)) else "".join(out)


def replace_command_with_single_arg(
    text: str,
    command: str,
    replacement: Callable[[str, str | None], str],
    *,
    consume_trailing_optional: bool = False,
) -> str:
    pattern = command_regex(command)
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        try:
            options, arg_pos = read_bracket_optional(text, match.end())
            while True:
                next_options, next_pos = read_bracket_optional(text, arg_pos)
                if next_options is None:
                    break
                options = next_options if options is None else f"{options},{next_options}"
                arg_pos = next_pos
            arg_pos = skip_ws(text, arg_pos)
            if arg_pos >= len(text) or text[arg_pos] != "{":
                out.append(text[pos : match.end()])
                pos = match.end()
                continue
            body, end = read_group(text, arg_pos)
            if consume_trailing_optional:
                _trailing, end = read_bracket_optional(text, end)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        out.append(text[pos : match.start()])
        out.append(replacement(body, options))
        pos = end
    return "".join(out)


def preserve_marginalia_commands(text: str) -> str:
    def as_footnote(body: str, _options: str | None) -> str:
        return rf"\footnote{{{body}}}"

    def as_todo(body: str, _options: str | None) -> str:
        return "\n\n" + rf"\noindent\textbf{{[TODO:}} {body}\textbf{{]}}" + "\n\n"

    text = replace_command_with_single_arg(text, "marginnote", as_footnote, consume_trailing_optional=True)
    text = replace_command_with_single_arg(text, "marginpar", as_footnote)
    text = replace_command_with_single_arg(text, "todo", as_todo)
    return text


def replace_biblatex_print_commands(text: str) -> str:
    heading = "\n\n" + r"\section*{References}" + "\n\n"
    for command in ("printbibliography", "printbibheading", "bibbysection"):
        pattern = command_regex(command)
        out: list[str] = []
        pos = 0
        while True:
            match = pattern.search(text, pos)
            if not match:
                out.append(text[pos:])
                break
            try:
                end = skip_optional_args(text, match.end())
            except ValueError:
                out.append(text[pos : match.end()])
                pos = match.end()
                continue
            out.append(text[pos : match.start()])
            out.append(heading)
            pos = end
        text = "".join(out)
    return text


def split_latex_keyvals(options: str) -> dict[str, str]:
    values: dict[str, str] = {}
    parts: list[str] = []
    buf: list[str] = []
    brace_depth = 0
    bracket_depth = 0
    for ch in options:
        if ch == "{" and bracket_depth == 0:
            brace_depth += 1
        elif ch == "}" and bracket_depth == 0:
            brace_depth = max(0, brace_depth - 1)
        elif ch == "[":
            bracket_depth += 1
        elif ch == "]":
            bracket_depth = max(0, bracket_depth - 1)
        if ch == "," and brace_depth == 0 and bracket_depth == 0:
            parts.append("".join(buf).strip())
            buf = []
            continue
        buf.append(ch)
    if buf:
        parts.append("".join(buf).strip())

    for part in parts:
        if not part:
            continue
        if "=" not in part:
            values[part.strip()] = ""
            continue
        key, value = part.split("=", 1)
        values[key.strip()] = strip_outer_braces(value.strip())
    return values


def begin_environment_at(text: str, pos: int) -> tuple[str, int] | None:
    match = re.match(r"\\begin\{([^{}]+)\}", text[pos:])
    if not match:
        return None
    return match.group(1), pos + match.end()


def read_environment_optional_args(text: str, pos: int) -> tuple[str | None, int]:
    options: list[str] = []
    pos = skip_ws(text, pos)
    while pos < len(text) and text[pos] == "[":
        option, pos = read_bracket(text, pos)
        options.append(option)
        pos = skip_ws(text, pos)
    return (",".join(options) if options else None), pos


def skip_environment_groups(text: str, pos: int, count: int) -> int:
    for _idx in range(count):
        pos = skip_ws(text, pos)
        if pos >= len(text) or text[pos] != "{":
            break
        _content, pos = read_group(text, pos)
    return pos


def find_matching_environment_end(text: str, env: str, pos: int) -> tuple[int, int] | None:
    token_re = re.compile(rf"\\(begin|end)\{{{re.escape(env)}\}}")
    depth = 1
    while True:
        match = token_re.search(text, pos)
        if not match:
            return None
        if match.group(1) == "begin":
            depth += 1
        else:
            depth -= 1
            if depth == 0:
                return match.start(), match.end()
        pos = match.end()


def bibitem_author_year(label: str) -> tuple[str, str]:
    clean = normalize_ws(apply_tex_accents(label))
    match = re.match(r"(.+?)\(([^()]*)\)\s*$", clean)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return clean, ""


def bibitem_entries(content: str, bibcites: dict[str, BibciteInfo]) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    pattern = command_regex("bibitem")
    pos = 0
    seq = 0
    while True:
        match = pattern.search(content, pos)
        if not match:
            break
        option_pos = skip_ws(content, match.end())
        label = ""
        if option_pos < len(content) and content[option_pos] == "[":
            try:
                label, option_pos = read_bracket(content, option_pos)
            except ValueError:
                pos = match.end()
                continue
        key_pos = skip_ws(content, option_pos)
        if key_pos >= len(content) or content[key_pos] != "{":
            pos = match.end()
            continue
        try:
            key, body_start = read_group(content, key_pos)
        except ValueError:
            pos = match.end()
            continue
        next_match = pattern.search(content, body_start)
        body_end = next_match.start() if next_match else len(content)
        seq += 1
        clean_key = normalize_ws(key)
        info = bibcites.get(clean_key)
        label_author, label_year = bibitem_author_year(label)
        entries.append(
            {
                "key": clean_key,
                "number": (info.number if info and info.number else str(seq)),
                "author": (info.author if info and info.author else label_author),
                "year": (info.year if info and info.year else label_year),
                "body": content[body_start:body_end],
            }
        )
        pos = body_end
    return entries


def manual_bibcite_fallbacks(text: str, existing: dict[str, BibciteInfo]) -> dict[str, BibciteInfo]:
    if not has_thebibliography(text):
        return existing
    merged = dict(existing)
    begin_re = re.compile(r"\\begin\{thebibliography\}")
    pos = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            break
        parsed = begin_environment_at(text, begin.start())
        if parsed is None:
            pos = begin.end()
            continue
        env, after_name = parsed
        content_start = skip_environment_groups(text, after_name, 1)
        end_span = find_matching_environment_end(text, env, content_start)
        if end_span is None:
            pos = begin.end()
            continue
        end_start, end_after = end_span
        for entry in bibitem_entries(text[content_start:end_start], merged):
            key = entry["key"]
            if key not in merged:
                merged[key] = BibciteInfo(
                    key=key,
                    number=entry["number"],
                    year=entry["year"],
                    author=entry["author"],
                )
        pos = end_after
    return merged


def clean_manual_bibliography_body(entry: str) -> str:
    s = strip_latex_comments(entry)
    s = apply_tex_accents(s)
    s = re.sub(r"\\newblock\b", " ", s)
    s = re.sub(r"\\(?:par|smallskip|medskip|bigskip)\b", " ", s)
    return normalize_ws(s)


def render_manual_bibliography(content: str, bibcites: dict[str, BibciteInfo]) -> tuple[str, int]:
    entries = bibitem_entries(content, bibcites)
    if not entries:
        return "", 0
    blocks = [r"\section*{References}", ""]
    for entry in entries:
        body = clean_manual_bibliography_body(entry["body"])
        if body:
            blocks.extend([rf"\noindent [{entry['number']}] {body}", ""])
    return "\n".join(blocks).strip(), len(entries)


def rewrite_thebibliography_environments(text: str, bibcites: dict[str, BibciteInfo]) -> tuple[str, int]:
    if not has_thebibliography(text):
        return text, 0
    begin_re = re.compile(r"\\begin\{thebibliography\}")
    out: list[str] = []
    pos = 0
    replaced = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            out.append(text[pos:])
            break
        parsed = begin_environment_at(text, begin.start())
        if parsed is None:
            out.append(text[pos : begin.end()])
            pos = begin.end()
            continue
        env, after_name = parsed
        content_start = skip_environment_groups(text, after_name, 1)
        end_span = find_matching_environment_end(text, env, content_start)
        if end_span is None:
            out.append(text[pos:])
            break
        end_start, end_after = end_span
        replacement, count = render_manual_bibliography(text[content_start:end_start], bibcites)
        out.append(text[pos : begin.start()])
        out.append("\n\n" + replacement + "\n\n")
        replaced += count
        pos = end_after
    return "".join(out), replaced


def insert_sibling_bbl(text: str, config: Config) -> str:
    if has_thebibliography(text):
        return text
    bbl_path = sibling_bbl_path(config)
    if not bbl_path.exists():
        return text
    bbl = bbl_path.read_text(encoding="utf-8", errors="replace")
    if not has_thebibliography(bbl):
        return text
    text = re.sub(r"\\bibliographystyle\s*\{[^{}]*\}", "", text)
    bibliography = find_command_argument(text, "bibliography")
    if bibliography is not None:
        start, _arg_pos, end, _arg = bibliography
        return text[:start] + "\n\n" + bbl + "\n\n" + text[end:]
    printbib = command_regex("printbibliography").search(text)
    if printbib:
        end = skip_optional_args(text, printbib.end())
        return text[: printbib.start()] + "\n\n" + bbl + "\n\n" + text[end:]
    end_doc = re.search(r"\\end\{document\}", text)
    if end_doc:
        return text[: end_doc.start()] + "\n\n" + bbl + "\n\n" + text[end_doc.start() :]
    return text + "\n\n" + bbl


def int_from_option(value: str) -> int | None:
    match = re.search(r"-?\d+", value)
    return int(match.group(0)) if match else None


def lower_roman(value: int) -> str:
    if value <= 0:
        return str(value)
    numerals = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    out: list[str] = []
    remainder = value
    for amount, symbol in numerals:
        while remainder >= amount:
            out.append(symbol)
            remainder -= amount
    return "".join(out)


def alpha_counter(value: int, uppercase: bool = False) -> str:
    if value <= 0:
        return str(value)
    base = ord("A" if uppercase else "a")
    chars: list[str] = []
    remainder = value
    while remainder:
        remainder -= 1
        chars.append(chr(base + (remainder % 26)))
        remainder //= 26
    return "".join(reversed(chars))


def render_enumitem_label(template: str, value: int) -> str:
    def repl(match: re.Match[str]) -> str:
        kind = match.group(1)
        if kind == "arabic":
            return str(value)
        if kind == "roman":
            return lower_roman(value)
        if kind == "Roman":
            return lower_roman(value).upper()
        if kind == "alph":
            return alpha_counter(value)
        if kind == "Alph":
            return alpha_counter(value, uppercase=True)
        return str(value)

    rendered, count = re.subn(r"\\(arabic|roman|Roman|alph|Alph)\*", repl, template)
    if count:
        return rendered.strip()
    if "*" in template:
        return template.replace("*", str(value)).strip()
    return latex_plain_text(template) or str(value)


def split_top_level_items(content: str) -> list[tuple[str | None, str]]:
    item_positions: list[int] = []
    env_stack: list[str] = []
    brace_depth = 0
    i = 0
    while i < len(content):
        if content.startswith(r"\item", i):
            after = i + len(r"\item")
            if (after >= len(content) or not re.match(r"[A-Za-z@]", content[after])) and brace_depth == 0 and not env_stack:
                item_positions.append(i)
                i = after
                continue
        if content[i] == "\\":
            env_cmd = re.match(r"\\(begin|end)\{([^{}]+)\}", content[i:])
            if env_cmd:
                kind, env = env_cmd.group(1), env_cmd.group(2)
                if kind == "begin":
                    env_stack.append(env)
                elif env_stack and env_stack[-1] == env:
                    env_stack.pop()
                i += env_cmd.end()
                continue
            if i + 1 < len(content) and content[i + 1] in "{}%":
                i += 2
                continue
        elif content[i] == "%" and brace_depth == 0:
            newline = content.find("\n", i)
            i = len(content) if newline < 0 else newline + 1
            continue
        elif not env_stack:
            if content[i] == "{":
                brace_depth += 1
            elif content[i] == "}":
                brace_depth = max(0, brace_depth - 1)
        i += 1

    items: list[tuple[str | None, str]] = []
    for index, start in enumerate(item_positions):
        next_start = item_positions[index + 1] if index + 1 < len(item_positions) else len(content)
        item_start = start + len(r"\item")
        item_start = skip_ws(content, item_start)
        item_label = None
        if item_start < len(content) and content[item_start] == "[":
            item_label, item_start = read_bracket(content, item_start)
        items.append((item_label, content[item_start:next_start]))
    return items


def rewrite_enumitem_environments(text: str) -> str:
    state: dict[str, dict[object, int]] = {
        "series": {},
        "last": {},
        "last_by_label": {},
    }

    def resume_option(options: dict[str, str]) -> tuple[bool, str]:
        for key in ("resume", "resume*"):
            if key in options:
                return True, options.get(key, "").strip()
        return False, ""

    def should_rewrite(options: dict[str, str]) -> bool:
        return any(key in options for key in ("label", "start", "series", "resume", "resume*"))

    def list_start(options: dict[str, str], label_template: str, depth: int) -> int:
        if "start" in options:
            explicit = int_from_option(options["start"])
            if explicit is not None:
                return explicit
        wants_resume, resume_name = resume_option(options)
        series_name = options.get("series", "").strip()
        if wants_resume:
            series_key = resume_name or series_name
            if series_key:
                return int(state["series"].get(series_key, 0)) + 1
            label_key = (depth, label_template)
            return int(state["last_by_label"].get(label_key, state["last"].get(depth, 0))) + 1
        return 1

    def update_list_state(options: dict[str, str], label_template: str, depth: int, last_value: int) -> None:
        wants_resume, resume_name = resume_option(options)
        series_name = options.get("series", "").strip()
        series_key = series_name or (resume_name if wants_resume else "")
        if series_key:
            state["series"][series_key] = last_value
        state["last"][depth] = last_value
        state["last_by_label"][(depth, label_template)] = last_value

    def render_list(content: str, options: dict[str, str], depth: int) -> str:
        items = split_top_level_items(content)
        if not items:
            return process(content, depth)
        label_template = options.get("label", r"\arabic*.")
        start_value = list_start(options, label_template, depth)
        rendered_items: list[str] = []
        for offset, (explicit_label, body) in enumerate(items):
            label = latex_plain_text(explicit_label) if explicit_label is not None else render_enumitem_label(label_template, start_value + offset)
            rendered_body = process(body.strip(), depth + 1).strip()
            rendered_items.append(f"{label} {rendered_body}".rstrip())
        update_list_state(options, label_template, depth, start_value + len(items) - 1)
        return "\n\n" + "\n\n".join(rendered_items) + "\n\n"

    def process(segment: str, depth: int = 0) -> str:
        begin_re = re.compile(r"\\begin\{enumerate\}")
        out: list[str] = []
        pos = 0
        while True:
            begin = begin_re.search(segment, pos)
            if not begin:
                out.append(segment[pos:])
                break
            parsed = begin_environment_at(segment, begin.start())
            if parsed is None:
                out.append(segment[pos : begin.end()])
                pos = begin.end()
                continue
            env, after_name = parsed
            options_s, content_start = read_environment_optional_args(segment, after_name)
            end_span = find_matching_environment_end(segment, env, content_start)
            if end_span is None:
                out.append(segment[pos:])
                break
            end_start, end_after = end_span
            content = segment[content_start:end_start]
            options = split_latex_keyvals(options_s or "")
            out.append(segment[pos : begin.start()])
            if options_s is not None and should_rewrite(options):
                out.append(render_list(content, options, depth))
            else:
                out.append(segment[begin.start() : content_start])
                out.append(process(content, depth + 1))
                out.append(segment[end_start:end_after])
            pos = end_after
        return "".join(out)

    return process(text)


def unwrap_multicol_environments(text: str) -> str:
    begin_re = re.compile(r"\\begin\{multicols\*?\}")
    out: list[str] = []
    pos = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            out.append(text[pos:])
            break
        parsed = begin_environment_at(text, begin.start())
        if parsed is None:
            out.append(text[pos : begin.end()])
            pos = begin.end()
            continue
        env, after_name = parsed
        _options, content_start = read_environment_optional_args(text, after_name)
        content_start = skip_environment_groups(text, content_start, 1)
        _options, content_start = read_environment_optional_args(text, content_start)
        end_span = find_matching_environment_end(text, env, content_start)
        if end_span is None:
            out.append(text[pos:])
            break
        end_start, end_after = end_span
        out.append(text[pos : begin.start()])
        out.append("\n\n")
        out.append(unwrap_multicol_environments(text[content_start:end_start]).strip())
        out.append("\n\n")
        pos = end_after
    return "".join(out)


def strip_label_commands(fragment: str) -> str:
    out: list[str] = []
    pos = 0
    for start, _arg_pos, end, _arg in iter_command_arguments(fragment, "label"):
        out.append(fragment[pos:start])
        pos = end
    out.append(fragment[pos:])
    return "".join(out)


def mhchem_display_replacement(content: str, labels: dict[str, LabelInfo]) -> str | None:
    if not contains_ce_commands(content):
        return None
    clean = strip_latex_comments(content).strip()
    label_items = list(iter_command_arguments(clean, "label"))
    label = label_items[-1][3] if label_items else ""
    body = strip_label_commands(clean).strip()
    body = re.sub(r"\\(?:notag|nonumber)\b", "", body).strip()
    match = command_regex("ce").match(body)
    if not match:
        return None
    arg_pos = skip_ws(body, match.end())
    if arg_pos >= len(body) or body[arg_pos] != "{":
        return None
    try:
        expression, end = read_group(body, arg_pos)
    except ValueError:
        return None
    tail = normalize_ws(body[end:])
    tail = tail.replace(r"\,", "").strip()
    if tail and not re.fullmatch(r"[.,;:!?]+", tail):
        return None
    rendered = render_ce(expression)
    if tail:
        rendered += tail
    number = labels[label].number if label in labels else ""
    suffix = f" ({number})" if number else ""
    return "\n\n\\begin{center}\n" + rendered + suffix + "\n\\end{center}\n\n"


def replace_mhchem_display_equations(text: str, labels: dict[str, LabelInfo]) -> str:
    if not contains_ce_commands(text):
        return text
    pattern = re.compile(r"\\begin\{(equation\*?)\}|\\\[|(?<!\\)\$\$")
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        opener = match.group(0)
        out.append(text[pos : match.start()])
        if opener == r"\[":
            end_start = text.find(r"\]", match.end())
            if end_start < 0:
                out.append(text[match.start() :])
                break
            content = text[match.end() : end_start]
            replacement = mhchem_display_replacement(content, labels)
            out.append(replacement if replacement is not None else text[match.start() : end_start + 2])
            pos = end_start + 2
            continue
        if opener == "$$":
            end_start = text.find("$$", match.end())
            if end_start < 0:
                out.append(text[match.start() :])
                break
            content = text[match.end() : end_start]
            replacement = mhchem_display_replacement(content, labels)
            out.append(replacement if replacement is not None else text[match.start() : end_start + 2])
            pos = end_start + 2
            continue

        env = match.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end_match = end_re.search(text, match.end())
        if not end_match:
            out.append(text[match.start() :])
            break
        content = text[match.end() : end_match.start()]
        replacement = mhchem_display_replacement(content, labels)
        out.append(replacement if replacement is not None else text[match.start() : end_match.end()])
        pos = end_match.end()
    return "".join(out)


def normalize_literal_command_names(text: str) -> str:
    text = re.sub(
        r"\\texttt\{\s*\\textbackslash(?:\{\})?\s*([A-Za-z@]+)\s*\}",
        lambda match: rf"\texttt{{{match.group(1)}}}",
        text,
    )

    def repl_verb(match: re.Match[str]) -> str:
        content = match.group(2)
        if re.fullmatch(r"\\[A-Za-z@]+", content):
            return rf"\texttt{{{content[1:]}}}"
        return match.group(0)

    return re.sub(r"\\verb(.)(.*?)\1", repl_verb, text)


def parse_aux_contents(aux_path: Path, stream: str) -> list[tuple[str, str, str]]:
    if not aux_path.exists():
        return []
    entries: list[tuple[str, str, str]] = []
    prefix = rf"\@writefile{{{stream}}}"
    for line in aux_path.read_text(encoding="utf-8", errors="replace").splitlines():
        start = line.find(prefix)
        if start < 0:
            continue
        pos = skip_ws(line, start + len(prefix))
        if pos >= len(line) or line[pos] != "{":
            continue
        try:
            payload, _end = read_group(line, pos)
        except ValueError:
            continue
        match = re.match(r"\\contentsline\s*\{([^{}]*)\}", payload)
        if not match:
            continue
        pos = skip_ws(payload, match.end())
        if pos >= len(payload) or payload[pos] != "{":
            continue
        try:
            raw_title, pos = read_group(payload, pos)
        except ValueError:
            continue
        pos = skip_ws(payload, pos)
        if pos >= len(payload) or payload[pos] != "{":
            continue
        try:
            page, _pos = read_group(payload, pos)
        except ValueError:
            page = ""
        number = ""
        title = raw_title
        number_match = re.match(r"\\numberline\s*\{", raw_title)
        if number_match:
            number_open = number_match.end() - 1
            try:
                number, number_end = read_group(raw_title, number_open)
                title = raw_title[number_end:]
            except ValueError:
                title = raw_title
        entries.append((latex_plain_text(title), latex_plain_text(number), latex_plain_text(page)))
    return entries


def latex_plain_text(fragment: str) -> str:
    s = unwrap_rotatebox_commands(fragment)
    s = re.sub(r"\\(?:ignorespaces|protected@file@percent|relax)\b", "", s)
    s = re.sub(r"\\textbackslash(?:\{\})?", "", s)
    s = re.sub(r"\\[A-Za-z@]+\*?(?:\[[^\]]*\])?", "", s)
    s = s.replace("{", "").replace("}", "")
    return normalize_ws(s)


def toc_replacement(entries: list[tuple[str, str, str]]) -> str:
    if not entries:
        return ""
    lines = [r"\section*{Contents}", ""]
    for title, number, page in entries:
        prefix = f"{number} " if number else ""
        suffix = f" {page}" if page else ""
        lines.extend([f"{prefix}{title}{suffix}", ""])
    return "\n".join(lines).strip()


def replace_tableofcontents(text: str, aux_path: Path) -> str:
    if r"\tableofcontents" not in text:
        return text
    replacement = toc_replacement(parse_aux_contents(aux_path, "toc"))
    return re.sub(r"\\tableofcontents\b", lambda _match: "\n\n" + replacement + "\n\n", text)


def prefix_appendix_sections(text: str) -> str:
    marker = r"\appendix"
    if marker not in text:
        return text
    head, tail = text.split(marker, 1)
    counter = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        title = match.group(1).strip()
        letter = chr(ord("A") + counter) if counter < 26 else str(counter + 1)
        counter += 1
        if re.match(rf"^{re.escape(letter)}\s+", title):
            return match.group(0)
        return rf"\section{{{letter} {title}}}"

    tail = re.sub(r"\\section\s*\{([^{}]*)\}", repl, tail)
    return head + tail


def replace_lstlisting_environments(text: str, labels: dict[str, LabelInfo]) -> str:
    begin_re = re.compile(r"\\begin\{lstlisting\}(?:\[([^\]]*)\])?")
    out: list[str] = []
    pos = 0
    listing_counter = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            out.append(text[pos:])
            break
        end = re.search(r"\\end\{lstlisting\}", text[begin.end() :])
        if not end:
            out.append(text[pos:])
            break
        end_start = begin.end() + end.start()
        end_pos = begin.end() + end.end()
        options = split_latex_keyvals(begin.group(1) or "")
        label = options.get("label", "")
        caption = normalize_ws(options.get("caption", ""))
        listing_counter += 1
        number = labels[label].number if label in labels else str(listing_counter)
        content = text[begin.end() : end_start].strip("\n")
        numbered_lines = [f"{idx} {line}" for idx, line in enumerate(content.splitlines(), start=1)]
        caption_line = f"\\noindent\\textbf{{Listing {number}: {caption}}}\n\n" if caption else ""
        replacement = "\n\n" + caption_line + "\\begin{verbatim}\n" + "\n".join(numbered_lines) + "\n\\end{verbatim}\n\n"
        out.append(text[pos : begin.start()])
        out.append(replacement)
        pos = end_pos
    return "".join(out)


def parse_acronym_definitions(text: str) -> dict[str, tuple[str, str]]:
    definitions: dict[str, tuple[str, str]] = {}
    pattern = command_regex("acro")
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            break
        key_pos = skip_ws(text, match.end())
        if key_pos >= len(text) or text[key_pos] != "{":
            pos = match.end()
            continue
        try:
            key, after_key = read_group(text, key_pos)
        except ValueError:
            pos = match.end()
            continue
        short = key.strip()
        long_pos = skip_ws(text, after_key)
        if long_pos < len(text) and text[long_pos] == "[":
            try:
                short, long_pos = read_bracket(text, long_pos)
            except ValueError:
                pos = after_key
                continue
            long_pos = skip_ws(text, long_pos)
        if long_pos >= len(text) or text[long_pos] != "{":
            pos = after_key
            continue
        try:
            long, end = read_group(text, long_pos)
        except ValueError:
            pos = after_key
            continue
        definitions[key.strip()] = (latex_plain_text(short) or key.strip(), normalize_ws(long))
        pos = end
    return definitions


def acronym_register_replacement(block: str, definitions: dict[str, tuple[str, str]]) -> str:
    lines = [r"\begin{description}"]
    for key in re.findall(r"\\acro\s*\{([^{}]+)\}", block):
        short, long = definitions.get(key, (key, key))
        lines.append(rf"\item[{short}] {long}")
    lines.append(r"\end{description}")
    return "\n".join(lines)


def replace_acronym_environments(text: str, definitions: dict[str, tuple[str, str]]) -> str:
    spans = environment_spans(text, "acronym")
    if not spans:
        return text
    out: list[str] = []
    pos = 0
    for start, end, block in spans:
        out.append(text[pos:start])
        out.append("\n\n" + acronym_register_replacement(block, definitions) + "\n\n")
        pos = end
    out.append(text[pos:])
    return "".join(out)


def expand_acronym_commands(text: str) -> str:
    definitions = parse_acronym_definitions(text)
    if not definitions:
        return text
    text = replace_acronym_environments(text, definitions)
    used: set[str] = set()
    pattern = re.compile(r"\\(ac|acs|acl|acf|Ac|Acs|Acl|Acf)\s*\{")
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        try:
            key, end = read_group(text, match.end() - 1)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        short, long = definitions.get(key.strip(), (key.strip(), key.strip()))
        command = match.group(1)
        lower = command.lower()
        if lower == "acs":
            rendered = short
        elif lower == "acl":
            rendered = long
        elif lower == "acf":
            rendered = f"{long} ({short})"
            used.add(key.strip())
        else:
            if key.strip() in used:
                rendered = short
            else:
                rendered = f"{long} ({short})"
                used.add(key.strip())
        if command[0].isupper() and rendered:
            rendered = rendered[0].upper() + rendered[1:]
        out.append(text[pos : match.start()])
        out.append(rendered)
        pos = end
    return "".join(out)


NEWCOMMAND_DENYLIST = {
    "arraystretch",
    "tabcolsep",
    "baselinestretch",
    "arrayrulewidth",
    "doublerulesep",
    "extrarowheight",
    "parskip",
    "parindent",
    "textwidth",
    "linewidth",
    # counter-format macros: expanding them rewrites the \renewcommand's own
    # first argument (\renewcommand{\thefigure}{A\arabic{figure}} ->
    # \renewcommand{A\arabic{figure}}{...}), which breaks pandoc; numbering
    # comes from the .aux anyway
    "thefigure",
    "thetable",
    "thesection",
    "thesubsection",
    "thesubsubsection",
    "theequation",
    "thepage",
    "thefootnote",
    "thempfootnote",
    "theparagraph",
    "theenumi",
    "theenumii",
    "theenumiii",
    "theenumiv",
    "today",
    "abstractname",
    "tablename",
    "figurename",
}


def _bare_dimension_or_number(text: str) -> bool:
    return bool(re.fullmatch(r"\s*[+-]?(?:\d+(?:\.\d*)?|\.\d+)\s*(?:pt|pc|in|bp|cm|mm|dd|cc|sp|em|ex)?\s*", text))


def parse_newcommand_definitions(text: str) -> dict[str, tuple[int, str | None, str]]:
    definitions: dict[str, tuple[int, str | None, str]] = {}
    pattern = re.compile(r"\\(?:newcommand|renewcommand|providecommand)\*?\s*")
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            break
        name: str | None = None
        i = skip_ws(text, match.end())
        try:
            if i < len(text) and text[i] == "{":
                group, i = read_group(text, i)
                name_m = re.match(r"\s*\\([A-Za-z]+)\s*$", group)
                if name_m:
                    name = name_m.group(1)
            else:
                name_m = re.match(r"\\([A-Za-z]+)", text[i:])
                if name_m:
                    name = name_m.group(1)
                    i += name_m.end()
            if not name:
                pos = match.end()
                continue

            argc = 0
            default: str | None = None
            opt, i = read_bracket_optional(text, i)
            if opt is not None:
                try:
                    argc = int(opt.strip())
                except ValueError:
                    pos = match.end()
                    continue
                opt2, i = read_bracket_optional(text, i)
                if opt2 is not None:
                    default = opt2
            i = skip_ws(text, i)
            if i >= len(text) or text[i] != "{":
                pos = match.end()
                continue
            body, i = read_group(text, i)
        except ValueError:
            pos = match.end()
            continue

        if name in NEWCOMMAND_DENYLIST:
            pos = i
            continue
        if argc == 0 and _bare_dimension_or_number(body):
            pos = i
            continue
        definitions[name] = (argc, default, body)
        pos = i
    return definitions


def expand_newcommands(
    text: str,
    definitions: dict[str, tuple[int, str | None, str]] | None = None,
    _depth: int = 0,
) -> str:
    if definitions is None:
        definitions = parse_newcommand_definitions(text)
    if not definitions or _depth > 8:
        return text

    name_alt = "|".join(sorted((re.escape(name) for name in definitions), key=len, reverse=True))
    pattern = re.compile(r"\\(" + name_alt + r")(?![A-Za-z])")
    out: list[str] = []
    pos = 0
    changed = False
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        name = match.group(1)
        argc, default, body = definitions[name]
        i = match.end()
        args: list[str] = []
        try:
            if default is not None:
                opt, i = read_bracket_optional(text, i)
                args.append(opt if opt is not None else default)
            while len(args) < argc:
                j = skip_ws(text, i)
                if j >= len(text) or text[j] != "{":
                    break
                arg, i = read_group(text, j)
                args.append(arg)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        if len(args) < argc:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue

        def repl_arg(arg_match: re.Match[str]) -> str:
            idx = int(arg_match.group(1))
            if 1 <= idx <= len(args):
                return args[idx - 1]
            return arg_match.group(0)

        expanded = re.sub(r"#(\d)", repl_arg, body)
        out.append(text[pos : match.start()])
        out.append(expanded)
        pos = i
        changed = True
    result = "".join(out)
    if changed:
        return expand_newcommands(result, definitions, _depth + 1)
    return result


def _glossary_field(body: str, field: str) -> str:
    match = re.search(rf"(?<![A-Za-z]){re.escape(field)}\s*=\s*", body)
    if not match:
        return ""
    pos = match.end()
    try:
        if pos < len(body) and body[pos] == "{":
            value, _end = read_group(body, pos)
        else:
            depth = 0
            start = pos
            while pos < len(body):
                char = body[pos]
                if char == "{":
                    depth += 1
                elif char == "}":
                    depth = max(0, depth - 1)
                elif char == "," and depth == 0:
                    break
                pos += 1
            value = body[start:pos]
    except ValueError:
        return ""
    return latex_plain_text(value).strip()


def parse_glossary_entries(text: str) -> dict[str, tuple[str, str]]:
    entries: dict[str, tuple[str, str]] = {}
    pattern = command_regex("newglossaryentry")
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            break
        key_pos = skip_ws(text, match.end())
        if key_pos >= len(text) or text[key_pos] != "{":
            pos = match.end()
            continue
        try:
            key, after_key = read_group(text, key_pos)
            body_pos = skip_ws(text, after_key)
            if body_pos >= len(text) or text[body_pos] != "{":
                pos = after_key
                continue
            body, end = read_group(text, body_pos)
        except ValueError:
            pos = match.end()
            continue
        clean_key = key.strip()
        name = _glossary_field(body, "name") or clean_key
        description = _glossary_field(body, "description")
        entries[clean_key] = (name, description)
        pos = end
    return entries


def parse_glossary_acronyms(text: str) -> dict[str, tuple[str, str]]:
    acronyms: dict[str, tuple[str, str]] = {}
    pos = 0
    while True:
        result = read_command_groups(text, "newacronym", 3, pos)
        if result is None:
            break
        _start, end, groups = result
        key, short, long = (group.strip() for group in groups)
        acronyms[key] = (latex_plain_text(short) or key, latex_plain_text(long) or normalize_ws(long))
        pos = end
    return acronyms


def expand_glossary_commands(
    text: str,
    glossary_entries: dict[str, tuple[str, str]],
    acronyms: dict[str, tuple[str, str]],
) -> str:
    if not glossary_entries and not acronyms:
        return text
    pattern = re.compile(r"\\(glspl|Glspl|GLSpl|gls|Gls|GLS)(?![A-Za-z@])\s*(?:\[[^\]]*\])?\s*\{")
    used: set[str] = set()
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        try:
            key, end = read_group(text, match.end() - 1)
        except ValueError:
            out.append(text[pos : match.end()])
            pos = match.end()
            continue
        clean_key = key.strip()
        command = match.group(1)
        plural = command.lower().endswith("pl")
        capitalize = command[0].isupper()
        if clean_key in acronyms:
            short, long = acronyms[clean_key]
            if clean_key in used:
                rendered = short + ("s" if plural else "")
            else:
                rendered = f"{long} ({short})"
                used.add(clean_key)
        elif clean_key in glossary_entries:
            name, _description = glossary_entries[clean_key]
            rendered = name + ("s" if plural else "")
            used.add(clean_key)
        else:
            rendered = clean_key
        if capitalize and rendered:
            rendered = rendered[0].upper() + rendered[1:]
        out.append(text[pos : match.start()])
        out.append(rendered)
        pos = end
    return "".join(out)


def render_glossary_sections(
    glossary_entries: dict[str, tuple[str, str]],
    acronyms: dict[str, tuple[str, str]],
) -> str:
    def hyphenation_hints(text: str) -> str:
        return re.sub(r"\b[Cc]omparable\b", lambda match: match.group(0)[:3] + "-" + match.group(0)[3:], text)

    blocks: list[str] = []
    if glossary_entries:
        blocks.append(r"\section*{Glossary}")
        blocks.append(r"\begin{description}")
        for _key, (name, description) in sorted(glossary_entries.items(), key=lambda item: item[1][0].lower()):
            blocks.append(rf"\item[{name}] {hyphenation_hints(description)}")
        blocks.append(r"\end{description}")
    if acronyms:
        blocks.append(r"\section*{Acronyms}")
        blocks.append(r"\begin{description}")
        for _key, (short, long) in sorted(acronyms.items(), key=lambda item: item[1][0].lower()):
            blocks.append(rf"\item[{short}] {long}")
        blocks.append(r"\end{description}")
    return "\n".join(blocks)


def replace_print_glossaries(
    text: str,
    glossary_entries: dict[str, tuple[str, str]],
    acronyms: dict[str, tuple[str, str]],
) -> str:
    print_cmd = re.compile(r"\\(printglossaries|printnoidxglossaries|printindexedglossaries)(?![A-Za-z@])")
    if not print_cmd.search(text):
        return text
    text = re.sub(r"\\(?:makeglossaries|makenoidxglossaries|glsresetall|glsaddall)\b", "", text)
    text = re.sub(r"\\printglossary\s*(?:\[[^\]]*\])?", "", text)
    if not glossary_entries and not acronyms:
        return print_cmd.sub("", text)
    rendered = "\n\n" + render_glossary_sections(glossary_entries, acronyms) + "\n\n"
    text = print_cmd.sub(lambda _m: rendered, text, count=1)
    return print_cmd.sub("", text)


def cleanup_latex_for_pandoc(text: str) -> str:
    text = re.sub(r"\\setlength\s*\\itemsep\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\setcounter\s*\{[^{}]*\}\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\(?:clearpage|newpage)\b", f"\n\n{latex_literal(PAGEBREAK_TOKEN)}\n\n", text)
    text = re.sub(r"(?m)^\s*%[-= ]{5,}\s*$", "", text)
    text = text.replace(r"\\%", r"\%")
    return text


def environment_command_at(tex: str, pos: int) -> tuple[str, str, int] | None:
    match = re.match(r"\\(begin|end)\{([^{}]+)\}", tex[pos:])
    if not match:
        return None
    return match.group(1), match.group(2), pos + match.end()


def split_display_rows(tex: str) -> list[str]:
    rows: list[str] = []
    buf: list[str] = []
    depth = 0
    env_stack: list[str] = []
    i = 0
    while i < len(tex):
        if tex[i] == "\\":
            env_cmd = environment_command_at(tex, i)
            if env_cmd is not None:
                kind, env, end = env_cmd
                if kind == "begin":
                    env_stack.append(env)
                elif env_stack and env_stack[-1] == env:
                    env_stack.pop()
                buf.append(tex[i:end])
                i = end
                continue
            if i + 1 < len(tex) and tex[i + 1] == "\\" and depth == 0 and not env_stack:
                rows.append("".join(buf))
                buf = []
                i += 2
                continue
            if i + 1 < len(tex):
                buf.append(tex[i : i + 2])
                i += 2
                continue
        ch = tex[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        buf.append(ch)
        i += 1
    rows.append("".join(buf))
    return rows


def remove_top_level_alignment_marks(tex: str) -> str:
    out: list[str] = []
    env_stack: list[str] = []
    i = 0
    while i < len(tex):
        if tex[i] == "\\":
            if i + 1 < len(tex) and tex[i + 1] == "&":
                out.append(r"\&")
                i += 2
                continue
            env_cmd = environment_command_at(tex, i)
            if env_cmd is not None:
                kind, env, end = env_cmd
                if kind == "begin":
                    env_stack.append(env)
                elif env_stack and env_stack[-1] == env:
                    env_stack.pop()
                out.append(tex[i:end])
                i = end
                continue
        if tex[i] == "&" and not env_stack:
            i += 1
            continue
        out.append(tex[i])
        i += 1
    return "".join(out)


def strip_row_number_suppression(tex: str) -> str:
    return re.sub(r"\\(?:notag|nonumber)\b", "", tex).strip()


def sanitize_math_content_for_pandoc(tex: str) -> str:
    tex = re.sub(r"\\(?:stepcounter|refstepcounter)\s*\{[^{}]*\}", "", tex)
    tex = re.sub(r"\\(?:addtocounter|setcounter)\s*\{[^{}]*\}\s*\{[^{}]*\}", "", tex)
    return strip_latex_comments(tex)


def sanitize_display_math_for_pandoc(text: str) -> str:
    pattern = re.compile(r"\\begin\{(equation\*?|align\*?|gather\*?|multline\*?)\}|\\\[|(?<!\\)\$\$")
    out: list[str] = []
    pos = 0
    while True:
        match = pattern.search(text, pos)
        if not match:
            out.append(text[pos:])
            break
        out.append(text[pos : match.start()])
        opener = match.group(0)
        if opener == r"\[":
            end_start = text.find(r"\]", match.end())
            if end_start < 0:
                out.append(text[match.start() :])
                break
            content = text[match.end() : end_start]
            out.append(opener + sanitize_math_content_for_pandoc(content) + r"\]")
            pos = end_start + 2
            continue
        if opener == "$$":
            end_start = text.find("$$", match.end())
            if end_start < 0:
                out.append(text[match.start() :])
                break
            content = text[match.end() : end_start]
            out.append(opener + sanitize_math_content_for_pandoc(content) + "$$")
            pos = end_start + 2
            continue

        env = match.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end_match = end_re.search(text, match.end())
        if not end_match:
            out.append(text[match.start() :])
            break
        content = text[match.end() : end_match.start()]
        out.append(opener + sanitize_math_content_for_pandoc(content) + end_match.group(0))
        pos = end_match.end()
    return "".join(out)


def numbered_row_blocks(rows: list[str], numbered: bool, first_label: str = "") -> list[str]:
    blocks: list[str] = []
    for idx, row in enumerate(rows):
        row_had_no_number = bool(re.search(r"\\(?:notag|nonumber)\b", row))
        cleaned = strip_row_number_suppression(remove_top_level_alignment_marks(row))
        if not cleaned:
            continue
        env = "equation" if numbered and not row_had_no_number else "equation*"
        if idx == 0 and first_label and r"\label" not in cleaned:
            cleaned = first_label + "\n" + cleaned
        blocks.append(f"\\begin{{{env}}}\n{cleaned}\n\\end{{{env}}}")
    return blocks


def rewrite_align_environments(text: str) -> str:
    begin_re = re.compile(r"\\begin\{(align\*?|gather\*?)\}")
    out: list[str] = []
    pos = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            out.append(text[pos:])
            break
        env = begin.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end = end_re.search(text, begin.end())
        if not end:
            out.append(text[pos:])
            break
        content = text[begin.end() : end.start()]
        rows = split_display_rows(content)
        blocks = numbered_row_blocks(rows, numbered=not env.endswith("*"))
        out.append(text[pos : begin.start()])
        out.append("\n\n" + "\n\n".join(blocks) + "\n\n")
        pos = end.end()
    return "".join(out)


def find_inner_split(content: str) -> tuple[str, int, int, str] | None:
    begin = re.search(r"\\begin\{(split|aligned)\}", content)
    if not begin:
        return None
    env = begin.group(1)
    end = re.search(rf"\\end\{{{re.escape(env)}\}}", content[begin.end() :])
    if not end:
        return None
    start = begin.start()
    end_start = begin.end() + end.start()
    end_pos = begin.end() + end.end()
    return env, start, end_pos, content[begin.end() : end_start]


def rewrite_equation_split_environments(text: str) -> str:
    begin_re = re.compile(r"\\begin\{(equation\*?)\}")
    out: list[str] = []
    pos = 0
    while True:
        begin = begin_re.search(text, pos)
        if not begin:
            out.append(text[pos:])
            break
        env = begin.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end = end_re.search(text, begin.end())
        if not end:
            out.append(text[pos:])
            break
        content = text[begin.end() : end.start()]
        split = find_inner_split(content)
        if split is None:
            out.append(text[pos : end.end()])
            pos = end.end()
            continue
        _split_env, _split_start, _split_end, split_content = split
        labels = [content[start:end] for start, _arg_pos, end, _arg in iter_command_arguments(content, "label")]
        rows = split_display_rows(split_content)
        blocks = numbered_row_blocks(rows, numbered=not env.endswith("*"), first_label="\n".join(labels))
        out.append(text[pos : begin.start()])
        out.append("\n\n" + "\n\n".join(blocks) + "\n\n")
        pos = end.end()
    return "".join(out)


def rewrite_multiline_math_displays(text: str) -> str:
    text = rewrite_equation_split_environments(text)
    text = rewrite_align_environments(text)
    return text


def parse_newtheorem_definitions(source: str) -> dict[str, TheoremDefinition]:
    defaults = {
        "theorem": "Theorem",
        "lemma": "Lemma",
        "proposition": "Proposition",
        "corollary": "Corollary",
        "definition": "Definition",
        "remark": "Remark",
        "example": "Example",
    }
    definitions = {
        env: TheoremDefinition(env=env, title=title, counter=env, numbered=True)
        for env, title in defaults.items()
    }
    begin_doc = re.search(r"\\begin\{document\}", source)
    preamble = source[: begin_doc.start()] if begin_doc else source
    pattern = re.compile(r"\\newtheorem(\*)?")
    pos = 0
    while True:
        match = pattern.search(preamble, pos)
        if not match:
            break
        numbered = not bool(match.group(1))
        env_pos = skip_ws(preamble, match.end())
        if env_pos >= len(preamble) or preamble[env_pos] != "{":
            pos = match.end()
            continue
        try:
            env, after_env = read_group(preamble, env_pos)
        except ValueError:
            pos = match.end()
            continue
        counter = env.strip()
        title_pos = skip_ws(preamble, after_env)
        if title_pos < len(preamble) and preamble[title_pos] == "[":
            try:
                counter, title_pos = read_bracket(preamble, title_pos)
            except ValueError:
                pos = after_env
                continue
            title_pos = skip_ws(preamble, title_pos)
        if title_pos >= len(preamble) or preamble[title_pos] != "{":
            pos = after_env
            continue
        try:
            title, end = read_group(preamble, title_pos)
        except ValueError:
            pos = after_env
            continue
        env_name = env.strip()
        definitions[env_name] = TheoremDefinition(
            env=env_name,
            title=latex_plain_text(title) or env_name.title(),
            counter=counter.strip() or env_name,
            numbered=numbered,
        )
        pos = end
    return definitions


def theorem_environment_spans(text: str, definitions: dict[str, TheoremDefinition]) -> list[tuple[int, int, str, str]]:
    spans: list[tuple[int, int, str, str]] = []
    for env in definitions:
        begin_re = re.compile(rf"\\begin\{{{re.escape(env)}\}}(?:\[([^\]]*)\])?")
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        pos = 0
        while True:
            begin = begin_re.search(text, pos)
            if not begin:
                break
            end = end_re.search(text, begin.end())
            if not end:
                break
            spans.append((begin.start(), end.end(), env, begin.group(1) or ""))
            pos = end.end()
    spans.sort(key=lambda item: item[0])
    return spans


def transform_theorem_environments(
    text: str,
    definitions: dict[str, TheoremDefinition],
    labels: dict[str, LabelInfo],
) -> str:
    spans = theorem_environment_spans(text, definitions)
    if not spans:
        return text
    out: list[str] = []
    pos = 0
    counters: dict[str, int] = {}
    for start, end, env, note in spans:
        if start < pos:
            continue
        definition = definitions[env]
        begin_match = re.match(rf"\\begin\{{{re.escape(env)}\}}(?:\[[^\]]*\])?", text[start:end])
        if begin_match is None:
            continue
        content_start = start + begin_match.end()
        content_end = end - len(rf"\end{{{env}}}")
        content = text[content_start:content_end].strip()
        label = command_argument_from_block(content, "label")
        number = ""
        if definition.numbered:
            if label in labels:
                number = labels[label].number
                parsed = parse_simple_equation_number(number)
                if parsed is not None:
                    counters[definition.counter] = max(counters.get(definition.counter, 0), parsed)
            else:
                counters[definition.counter] = counters.get(definition.counter, 0) + 1
                number = str(counters[definition.counter])
        body = strip_label_commands(content).strip()
        label_parts = [definition.title]
        if number:
            label_parts.append(number)
        label_text = " ".join(label_parts)
        note_text = f" ({normalize_ws(note)})" if note.strip() else ""
        label_text = f"{label_text}{note_text}."
        replacement = f"\n\n\\noindent\\textbf{{{label_text}}} \\emph{{{body}}}\n\n"
        out.append(text[pos:start])
        out.append(replacement)
        pos = end
    out.append(text[pos:])
    return "".join(out)


def parse_simple_equation_number(number: str) -> int | None:
    if re.fullmatch(r"\d+", number):
        return int(number)
    return None


MATH_TAG_SYMBOLS = {
    "star": "⋆",
    "ast": "*",
    "dagger": "†",
    "ddagger": "‡",
}


def render_manual_math_tag(tag: str) -> str:
    s = tag.strip()
    if s.startswith("$") and s.endswith("$") and len(s) >= 2:
        s = s[1:-1].strip()
    for command, symbol in MATH_TAG_SYMBOLS.items():
        s = re.sub(rf"\\{command}\b", symbol, s)
    s = re.sub(r"\\text\{([^{}]*)\}", r"\1", s)
    s = s.replace("{", "").replace("}", "")
    return normalize_ws(s)


def manual_math_tag(fragment: str) -> str:
    found = find_command_argument(fragment, "tag")
    return render_manual_math_tag(found[3]) if found is not None else ""


def extract_display_math_numbers(source: str, labels: dict[str, LabelInfo]) -> list[list[str]]:
    begin_doc = re.search(r"\\begin\{document\}", source)
    body = source[begin_doc.start() :] if begin_doc else source
    body = strip_latex_comments(body)
    pattern = re.compile(r"\\begin\{(equation\*?|align\*?|gather\*?|multline\*?)\}|\\\[")
    display_numbers: list[list[str]] = []
    counter = 0
    pos = 0

    def assign_number(fragment: str) -> str:
        nonlocal counter
        counter += 1
        label = command_argument_from_block(fragment, "label")
        if label in labels:
            number = labels[label].number
            parsed = parse_simple_equation_number(number)
            if parsed is not None:
                counter = max(counter, parsed)
            return number
        return str(counter)

    while True:
        match = pattern.search(body, pos)
        if not match:
            break
        if match.group(0) == r"\[":
            end = body.find(r"\]", match.end())
            if end < 0:
                break
            display_numbers.append([])
            pos = end + 2
            continue

        env = match.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end_match = end_re.search(body, match.end())
        if not end_match:
            break
        content = body[match.end() : end_match.start()]
        numbers: list[str] = []
        numbered = not env.endswith("*")
        base_env = env.rstrip("*")
        if numbered and base_env in {"align", "gather"}:
            for row in split_display_rows(content):
                if not row.strip() or re.search(r"\\(?:notag|nonumber)\b", row):
                    continue
                tag = manual_math_tag(row)
                if tag:
                    counter += 1
                    numbers.append(tag)
                else:
                    numbers.append(assign_number(row))
        elif numbered and base_env in {"equation", "multline"}:
            if not re.search(r"\\(?:notag|nonumber)\b", content):
                tag = manual_math_tag(content)
                if tag:
                    counter += 1
                    numbers.append(tag)
                else:
                    numbers.append(assign_number(content))
        display_numbers.append(numbers)
        pos = end_match.end()

    return display_numbers


def replace_floats(
    text: str,
    labels: dict[str, LabelInfo],
    config: Config,
) -> tuple[str, list[TableRecord], list[FigureRecord]]:
    float_env_pattern = "|".join(re.escape(env) for env in sorted(TABLE_FLOAT_ENVS | FIGURE_FLOAT_ENVS, key=len, reverse=True))
    begin_re = re.compile(rf"\\begin\{{({float_env_pattern})\}}(?:\[[^\]]*\])?")
    tables: list[TableRecord] = []
    figures: list[FigureRecord] = []
    out: list[str] = []
    pos = 0
    table_idx = 0
    figure_idx = 0

    while True:
        match = begin_re.search(text, pos)
        if not match:
            out.append(text[pos:])
            break

        env = match.group(1)
        end_re = re.compile(rf"\\end\{{{re.escape(env)}\}}")
        end_match = end_re.search(text, match.end())
        if not end_match:
            raise RuntimeError(f"Could not find end of {env} beginning at line {line_number_at(text, match.start())}")

        block = text[match.start() : end_match.end()]
        out.append(text[pos : match.start()])
        start_line = line_number_at(text, match.start())
        is_figure_float = env in FIGURE_FLOAT_ENVS
        caption = figure_outer_caption(block) if is_figure_float else command_argument_from_block(block, "caption")
        label = figure_outer_label(block) if is_figure_float else command_argument_from_block(block, "label")

        if env in TABLE_FLOAT_ENVS:
            table_idx += 1
            label = label or f"idx{table_idx}"
            number = labels[label].number if label in labels else str(table_idx)
            token = f"%%TABLE:{label}%%"
            replacement = f"\n\n{latex_literal(token)}\n\n"
            tables.append(
                detect_table_record(
                    env=env,
                    label=label,
                    caption=caption,
                    number=number,
                    placeholder=token,
                    start_line=start_line,
                    block=block,
                    config=config,
                )
            )
        else:
            figure_idx += 1
            label = label or f"idx{figure_idx}"
            number = labels[label].number if label in labels else str(figure_idx)
            graphic_specs = figure_graphic_specs(block)
            image_paths = [path for path, _width in graphic_specs]
            image_widths = [width for _path, width in graphic_specs]
            image_path = image_paths[0] if image_paths else f"idx{figure_idx}"
            token = f"%%FIGURE:{label}%%"
            replacement = f"\n\n{latex_literal(token)}\n\n"
            subcaption_map = figure_subcaptions(block)
            subcaptions = [subcaption_map.get(path, "") for path in image_paths]
            transcript = extract_algorithm_transcript(block)
            resolved_paths: list[str] = []
            for path_s in image_paths:
                try:
                    resolved_paths.append(str(resolve_graphic_path(path_s, config)))
                except Exception:
                    resolved_paths.append("")
            resolved = resolved_paths[0] if resolved_paths else ""
            figures.append(
                FigureRecord(
                    kind="figure",
                    env=env,
                    label=label,
                    placeholder=token,
                    start_line=start_line,
                    caption=normalize_ws(caption),
                    number=number,
                    image_path=image_path,
                    resolved_path=resolved,
                    image_paths=image_paths,
                    resolved_paths=resolved_paths,
                    subcaptions=subcaptions,
                    image_widths=image_widths,
                    transcript=transcript,
                )
            )

        out.append(replacement)
        pos = end_match.end()

    return "".join(out), tables, figures


def text_from_element(element: ET.Element) -> str:
    return "".join(node.text or "" for node in element.iter(f"{W}t")).strip()


def set_font(style, name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def remove_style_borders(style) -> None:
    ppr = style.element.get_or_add_pPr()
    border = ppr.find(qn("w:pBdr"))
    if border is not None:
        ppr.remove(border)


def build_reference_docx(path: Path, font: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    set_font(normal, font, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(0)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = Pt(15)

    for name, size in [("Heading 1", 14), ("Heading 2", 12.5), ("Heading 3", 12)]:
        style = doc.styles[name]
        set_font(style, font, size, bold=True)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)

    if "Title" in doc.styles:
        title = doc.styles["Title"]
        set_font(title, font, 15, bold=True)
        title.font.underline = False
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(10)
        title.paragraph_format.first_line_indent = Pt(0)
        remove_style_borders(title)

    try:
        author = doc.styles["Author"]
    except KeyError:
        author = doc.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    set_font(author, font, 12)
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(6)
    author.paragraph_format.first_line_indent = Pt(0)

    for name in ["Date", "Subtitle"]:
        if name in doc.styles:
            style = doc.styles[name]
            set_font(style, font, 12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Hyperlink" in doc.styles:
        hyperlink = doc.styles["Hyperlink"]
        set_font(hyperlink, font, 12)
        hyperlink.font.underline = False

    try:
        footnote = doc.styles["Footnote Text"]
    except KeyError:
        footnote = doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    set_font(footnote, font, 10)
    footnote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    footnote.paragraph_format.space_after = Pt(0)
    footnote.paragraph_format.first_line_indent = Pt(0)

    doc.add_paragraph("Reference document for LaTeX body conversion.")
    doc.save(path)


def build_pandoc_tex(config: Config) -> dict[str, object]:
    labels = parse_aux_labels(config.aux)
    source = config.main_tex.read_text(encoding="utf-8", errors="replace")
    # Inline \input/\include bodies once, up front, so every downstream stage
    # (float detection, preamble stripping, citation detection, ...) sees the
    # same fully-assembled source. Floats living in part files then reach the
    # native table/figure engines instead of being resolved by pandoc.
    try:
        source = expand_latex_inputs(source, config)
    except Exception:
        pass
    prefixes = parse_cref_prefixes(source)
    theorem_definitions = parse_newtheorem_definitions(source)
    glossary_entries = parse_glossary_entries(source)
    glossary_acronyms = parse_glossary_acronyms(source)
    config.user_macros = parse_newcommand_definitions(source)
    citations = detect_citation_mode(source, config)
    manual_bibcites = parse_aux_bibcites(config.aux)
    manual_citation_counts: dict[str, int] = {}
    missing_manual_citations: list[str] = []
    manual_bibliography_entries = 0
    numbermap = {label: info.number for label, info in labels.items()}
    if hasattr(T, "set_cref_prefixes"):
        T.set_cref_prefixes(prefixes)
    if hasattr(T, "set_font"):
        T.set_font(config.font)
    else:
        T.FONT = config.font

    source = render_latex_environment_images(source, config)
    text = strip_source_preamble(source)
    text, author_notes = neutralize_author_thanks(text)
    text = convert_href_commands(text)
    text = expand_newcommands(text, config.user_macros)
    text = expand_siunitx_commands(text)
    text = unwrap_rotatebox_commands(text)
    text = preserve_marginalia_commands(text)
    if citations.mode in {"manual-bibliography", "citation-best-effort"}:
        text = insert_sibling_bbl(text, config)
        manual_bibcites = manual_bibcite_fallbacks(text, manual_bibcites)
        text, manual_citation_counts, missing_manual_citations = resolve_numeric_citations(text, manual_bibcites)
    elif citations.mode in {"biblatex", "natbib-bibtex"}:
        text = replace_biblatex_print_commands(text)
    if contains_ce_commands(text):
        text = replace_mhchem_display_equations(text, labels)
        text = replace_ce_commands(text)
    text, tables, figures = replace_floats(text, labels, config)
    text = replace_tableofcontents(text, config.aux)
    text = prefix_appendix_sections(text)
    text = replace_lstlisting_environments(text, labels)
    text = expand_acronym_commands(text)
    text = expand_glossary_commands(text, glossary_entries, glossary_acronyms)
    text = replace_print_glossaries(text, glossary_entries, glossary_acronyms)
    text = transform_theorem_environments(text, theorem_definitions, labels)
    text = normalize_literal_command_names(text)
    text = convert_sym_macros(text)
    text, ref_counts, missing_refs = resolve_cross_references(text, labels, prefixes)
    if citations.mode == "manual-bibliography":
        text, manual_bibliography_entries = rewrite_thebibliography_environments(text, manual_bibcites)
    text = sanitize_display_math_for_pandoc(text)
    text = rewrite_multiline_math_displays(text)
    display_math_numbers = extract_display_math_numbers(text, labels)
    text = rewrite_enumitem_environments(text)
    text = unwrap_multicol_environments(text)
    text = cleanup_latex_for_pandoc(text)

    pandoc_tex = config.workdir / f"{config.main_tex.stem}_pandoc.tex"
    pandoc_tex.write_text(text, encoding="utf-8", newline="\n")
    unresolved_commands = re.findall(r"\\(?:Crefrange|crefrange|Cref|cref|eqref|autoref|ref)\s*\{", text)

    return {
        "labels": len(labels),
        "numbermap": numbermap,
        "cref_prefixes": {key: list(value) for key, value in prefixes.items()},
        "author_notes": author_notes,
        "reference_counts": ref_counts,
        "missing_references": missing_refs,
        "unresolved_reference_commands": len(unresolved_commands),
        "citations": asdict(citations),
        "manual_citation_counts": manual_citation_counts,
        "missing_manual_citations": missing_manual_citations,
        "manual_bibliography_entries": manual_bibliography_entries,
        "tables": [asdict(item) for item in tables],
        "figures": [asdict(item) for item in figures],
        "display_math_numbers": display_math_numbers,
        "table_floats": len(tables),
        "figure_floats": len(figures),
        "pandoc_tex": str(pandoc_tex),
    }


def run_command(command: list[str], cwd: Path, timeout: int | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        cwd=str(cwd),
        text=True,
        encoding="utf-8",
        errors="replace",
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
    )


_PANDOC_EXE: str | None = None


def pandoc_executable() -> str:
    """Resolve the pandoc executable once, exiting cleanly if it is missing."""
    global _PANDOC_EXE
    if _PANDOC_EXE is None:
        try:
            _PANDOC_EXE = toolcheck.find_pandoc()
        except toolcheck.ToolNotFoundError as exc:
            sys.exit(f"Error: {exc}")
    return _PANDOC_EXE


def resolve_user_csl(csl: str, config: Config) -> str:
    if re.match(r"https?://", csl):
        return csl
    raw = Path(csl).expanduser()
    candidates = [raw] if raw.is_absolute() else [Path.cwd() / raw, config.main_tex.parent / raw]
    for candidate in dedupe_paths(candidates):
        if candidate.exists():
            return str(candidate.resolve())
    raise FileNotFoundError(f"CSL file does not exist: {csl}")


def write_default_csl(config: Config) -> Path:
    csl_path = config.workdir / f"{DEFAULT_CSL_NAME}.csl"
    if csl_path.exists():
        return csl_path
    proc = run_command([pandoc_executable(), "--print-default-data-file", "default.csl"], config.workdir, timeout=30)
    if proc.returncode != 0 or not proc.stdout.strip():
        raise RuntimeError(
            "Could not load Pandoc's default Chicago author-date CSL style:\n"
            + (proc.stderr.strip() or proc.stdout.strip())
        )
    csl_path.write_text(proc.stdout, encoding="utf-8", newline="\n")
    return csl_path


def csl_for_pandoc(config: Config, citation_info: dict[str, object]) -> tuple[str, str]:
    if config.csl:
        return resolve_user_csl(config.csl, config), config.csl
    default_path = write_default_csl(config)
    return str(default_path), str(citation_info.get("default_csl") or DEFAULT_CSL_NAME)


def run_pandoc(
    config: Config,
    pandoc_tex: Path,
    reference_docx: Path,
    body_docx: Path,
    citation_info: dict[str, object],
) -> dict[str, object]:
    if body_docx.exists():
        body_docx.unlink()
    resource_path = os.pathsep.join([str(config.main_tex.parent), str(config.workdir)])
    command = [
        pandoc_executable(),
        str(pandoc_tex),
        "--from=latex",
        f"--resource-path={resource_path}",
        f"--reference-doc={reference_docx}",
        "-o",
        str(body_docx),
    ]
    citation_mode = str(citation_info.get("mode") or "footnote-only")
    effective_csl = ""
    if citation_mode in {"biblatex", "natbib-bibtex"}:
        bibliography_paths = [str(path) for path in citation_info.get("bibliography_paths", [])]
        if not bibliography_paths:
            raise RuntimeError("Citation mode is active, but no bibliography paths were provided to Pandoc.")
        command.append("--citeproc")
        for path in bibliography_paths:
            command.append(f"--bibliography={path}")
        csl_path, effective_csl = csl_for_pandoc(config, citation_info)
        command.append(f"--csl={csl_path}")
    proc = run_command(command, config.main_tex.parent)
    (config.workdir / "pandoc_warnings.txt").write_text(proc.stderr, encoding="utf-8")
    if proc.returncode != 0:
        raise RuntimeError(f"pandoc failed with exit code {proc.returncode}:\n{proc.stderr}")
    return {
        "command": " ".join(command),
        "citation_mode": citation_mode,
        "csl": effective_csl,
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "warnings": [line for line in proc.stderr.splitlines() if line.strip()],
    }


def replace_pagebreak_markers(root: ET.Element) -> int:
    converted = 0
    for paragraph in root.iter(f"{W}p"):
        if text_from_element(paragraph) != PAGEBREAK_TOKEN:
            continue
        for child in list(paragraph):
            if child.tag != f"{W}pPr":
                paragraph.remove(child)
        run = ET.Element(f"{W}r")
        br = ET.SubElement(run, f"{W}br")
        br.set(f"{W}type", "page")
        paragraph.append(run)
        converted += 1
    return converted


def ensure_child(parent: ET.Element, tag: str, insert_at_start: bool = False) -> ET.Element:
    child = parent.find(tag)
    if child is None:
        child = ET.Element(tag)
        if insert_at_start:
            parent.insert(0, child)
        else:
            parent.append(child)
    return child


def force_footnote_superscripts_in_root(root: ET.Element) -> tuple[int, int]:
    total = 0
    fixed = 0
    for run in root.iter(f"{W}r"):
        if run.find(f"{W}footnoteReference") is None:
            continue
        total += 1
        rpr = ensure_child(run, f"{W}rPr", insert_at_start=True)
        va = ensure_child(rpr, f"{W}vertAlign")
        if va.get(f"{W}val") != "superscript":
            va.set(f"{W}val", "superscript")
            fixed += 1
    return total, fixed


def force_footnote_reference_style(styles_root: ET.Element) -> None:
    for style in styles_root.iter(f"{W}style"):
        if style.get(f"{W}styleId") != "FootnoteReference":
            continue
        rpr = ensure_child(style, f"{W}rPr")
        va = ensure_child(rpr, f"{W}vertAlign")
        va.set(f"{W}val", "superscript")


def postprocess_docx(docx_path: Path, math_numbers: list[list[str]] | None = None) -> dict[str, int]:
    from lxml import etree

    math_numbers = math_numbers or []
    tmp_docx = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        data = {name: zin.read(name) for name in names}

    ns = {"w": W_NS}

    def lxml_text(element) -> str:
        return "".join(node.text or "" for node in element.xpath(".//w:t", namespaces=ns)).strip()

    def lxml_child(parent, tag: str, insert_at_start: bool = False):
        child = parent.find(tag)
        if child is None:
            child = etree.Element(tag)
            if insert_at_start:
                parent.insert(0, child)
            else:
                parent.append(child)
        return child

    def paragraph_style(paragraph) -> str:
        ppr = paragraph.find(f"{W}pPr")
        if ppr is None:
            return ""
        pstyle = ppr.find(f"{W}pStyle")
        return pstyle.get(f"{W}val") if pstyle is not None else ""

    def append_text_run(paragraph, text: str, bold: bool = False) -> None:
        run = etree.Element(f"{W}r")
        if bold:
            rpr = etree.SubElement(run, f"{W}rPr")
            etree.SubElement(rpr, f"{W}b")
        t = etree.SubElement(run, f"{W}t")
        t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        paragraph.append(run)

    document_root = etree.fromstring(data["word/document.xml"])
    pagebreaks = 0
    for paragraph in document_root.iter(f"{W}p"):
        if lxml_text(paragraph) != PAGEBREAK_TOKEN:
            continue
        for child in list(paragraph):
            if child.tag != f"{W}pPr":
                paragraph.remove(child)
        run = etree.Element(f"{W}r")
        br = etree.SubElement(run, f"{W}br")
        br.set(f"{W}type", "page")
        paragraph.append(run)
        pagebreaks += 1

    math_blocks = 0
    equation_numbers_added = 0
    math_iter = iter(math_numbers)
    for paragraph in document_root.iter(f"{W}p"):
        if paragraph.find(f"{M}oMathPara") is None:
            continue
        math_blocks += 1
        ppr = lxml_child(paragraph, f"{W}pPr", insert_at_start=True)
        spacing = lxml_child(ppr, f"{W}spacing")
        spacing.set(f"{W}before", "160")
        spacing.set(f"{W}after", "160")
        spacing.set(f"{W}line", "360")
        spacing.set(f"{W}lineRule", "auto")
        try:
            nums = next(math_iter)
        except StopIteration:
            nums = []
        nums = [num for num in nums if num]
        if nums:
            tabs = lxml_child(ppr, f"{W}tabs")
            if not tabs.findall(f"{W}tab"):
                tab = etree.SubElement(tabs, f"{W}tab")
                tab.set(f"{W}val", "right")
                tab.set(f"{W}pos", "9360")
            run = etree.Element(f"{W}r")
            etree.SubElement(run, f"{W}tab")
            text = etree.SubElement(run, f"{W}t")
            text.text = ", ".join(f"({num})" for num in nums)
            paragraph.append(run)
            equation_numbers_added += len(nums)

    source_code_paragraphs_left_aligned = 0
    for paragraph in document_root.iter(f"{W}p"):
        if paragraph_style(paragraph) not in {"SourceCode", "CodeBlock"}:
            continue
        ppr = lxml_child(paragraph, f"{W}pPr", insert_at_start=True)
        jc = lxml_child(ppr, f"{W}jc")
        if jc.get(f"{W}val") != "left":
            jc.set(f"{W}val", "left")
            source_code_paragraphs_left_aligned += 1

    description_terms_merged = 0
    body = document_root.find(f"{W}body")
    if body is not None:
        children = list(body)
        idx = 0
        while idx + 1 < len(children):
            paragraph = children[idx]
            following = children[idx + 1]
            if paragraph.tag != f"{W}p" or following.tag != f"{W}p":
                idx += 1
                continue
            if paragraph_style(paragraph) != "DefinitionTerm" or paragraph_style(following) != "Definition":
                idx += 1
                continue
            term = lxml_text(paragraph)
            if not term:
                idx += 1
                continue
            for child in list(paragraph):
                if child.tag != f"{W}pPr":
                    paragraph.remove(child)
            append_text_run(paragraph, term, bold=True)
            append_text_run(paragraph, ": ")
            for child in list(following):
                if child.tag == f"{W}pPr":
                    continue
                following.remove(child)
                paragraph.append(child)
            body.remove(following)
            children.pop(idx + 1)
            description_terms_merged += 1

    footnote_refs = 0
    footnote_fixed = 0
    for run in document_root.iter(f"{W}r"):
        if run.find(f"{W}footnoteReference") is None:
            continue
        footnote_refs += 1
        rpr = lxml_child(run, f"{W}rPr", insert_at_start=True)
        va = lxml_child(rpr, f"{W}vertAlign")
        if va.get(f"{W}val") != "superscript":
            va.set(f"{W}val", "superscript")
            footnote_fixed += 1
    data["word/document.xml"] = etree.tostring(
        document_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    if "word/styles.xml" in data:
        styles_root = etree.fromstring(data["word/styles.xml"])
        for style in styles_root.iter(f"{W}style"):
            if style.get(f"{W}styleId") != "FootnoteReference":
                if style.get(f"{W}styleId") == "DefinitionTerm":
                    rpr = lxml_child(style, f"{W}rPr")
                    lxml_child(rpr, f"{W}b")
                continue
            rpr = lxml_child(style, f"{W}rPr")
            va = lxml_child(rpr, f"{W}vertAlign")
            va.set(f"{W}val", "superscript")
        data["word/styles.xml"] = etree.tostring(
            styles_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    backup = docx_path.with_suffix(docx_path.suffix + ".bak")
    shutil.copyfile(docx_path, backup)
    with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, data[name])
    shutil.move(tmp_docx, docx_path)
    return {
        "pagebreaks_converted": pagebreaks,
        "footnote_reference_runs": footnote_refs,
        "footnote_reference_runs_fixed": footnote_fixed,
        "display_math_blocks": math_blocks,
        "equation_numbers_added": equation_numbers_added,
        "source_code_paragraphs_left_aligned": source_code_paragraphs_left_aligned,
        "description_terms_merged": description_terms_merged,
    }


def insert_doc_before(anchor_p, scratch: Document) -> None:
    body = scratch.element.body
    anchor = anchor_p._p
    for el in list(body):
        if el.tag.endswith("}sectPr"):
            continue
        body.remove(el)
        anchor.addprevious(el)
    anchor.getparent().remove(anchor)


def _sect_set_type(sectPr, val: str) -> None:
    """Ensure a ``w:type`` child equal to ``val`` exists, positioned before
    ``w:pgSz`` (the schema order)."""
    t = sectPr.find(qn("w:type"))
    if t is None:
        t = OxmlElement("w:type")
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is not None:
            pgSz.addprevious(t)
        else:
            sectPr.insert(0, t)
    t.set(qn("w:val"), val)


def _sect_set_landscape(sectPr) -> None:
    """Rotate a section's page to landscape (long edge horizontal), swapping the
    page width/height and setting ``w:orient``. Defaults to US Letter if unset."""
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.append(pgSz)
    try:
        wv = int(pgSz.get(qn("w:w")))
        hv = int(pgSz.get(qn("w:h")))
    except (TypeError, ValueError):
        wv, hv = 12240, 15840
    lo, hi = sorted((wv, hv))
    pgSz.set(qn("w:w"), str(hi))
    pgSz.set(qn("w:h"), str(lo))
    pgSz.set(qn("w:orient"), "landscape")


def insert_doc_landscape_before(anchor_p, scratch: Document, body: Document) -> None:
    """Insert a built table (caption + table + notes) on its OWN landscape page.

    Wraps the inserted block in a next-page landscape Word section, with a portrait
    section boundary before it and the document's portrait section resuming on a new
    page after — the faithful rendering of a LaTeX ``sidewaystable`` (a rotated table
    on its own page). The body's final ``sectPr`` (portrait) is the template for both
    boundary sections, so headers/footers and margins carry through.
    """
    anchor = anchor_p._p
    body_sectPr = body.element.body.find(qn("w:sectPr"))

    # Portrait boundary that closes the preceding section before the table block.
    portrait = copy.deepcopy(body_sectPr)
    _sect_set_type(portrait, "nextPage")
    p_before = OxmlElement("w:p")
    pPr_b = OxmlElement("w:pPr")
    pPr_b.append(portrait)
    p_before.append(pPr_b)
    anchor.addprevious(p_before)

    # The table block (caption, table, notes), in order, before the anchor.
    sbody = scratch.element.body
    for el in list(sbody):
        if el.tag.endswith("}sectPr"):
            continue
        sbody.remove(el)
        anchor.addprevious(el)

    # Landscape boundary that closes the table's own (landscape) section.
    landscape = copy.deepcopy(body_sectPr)
    _sect_set_landscape(landscape)
    _sect_set_type(landscape, "nextPage")
    p_after = OxmlElement("w:p")
    pPr_a = OxmlElement("w:pPr")
    pPr_a.append(landscape)
    p_after.append(pPr_a)
    anchor.addprevious(p_after)

    # The portrait body resumes on a new page after the landscape section.
    _sect_set_type(body_sectPr, "nextPage")
    anchor.getparent().remove(anchor)


def rasterize_graphic_for_docx(image_path: Path, config: Config) -> Path:
    if image_path.suffix.lower() not in {".pdf", ".eps"}:
        return image_path
    import fitz

    cache_dir = config.workdir / "_rasterized_figures"
    cache_dir.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha1(str(image_path.resolve()).encode("utf-8")).hexdigest()[:12]
    out_path = cache_dir / f"{image_path.stem}_{key}.png"
    if out_path.exists():
        return out_path
    with fitz.open(str(image_path)) as pdf:
        if pdf.page_count < 1:
            raise RuntimeError(f"Graphic has no pages: {image_path}")
        pix = pdf.load_page(0).get_pixmap(matrix=fitz.Matrix(220 / 72, 220 / 72), alpha=False)
        pix.save(str(out_path))
    return out_path


def _set_borderless_table(tbl) -> None:
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _caption_paragraph(doc: Document, text: str, font_pt: float = 10.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    T.inline_to_runs(p, text, font_pt)
    return p


def _add_searchable_transcript_after(paragraph: Paragraph, transcript: str) -> Paragraph:
    text = normalize_ws(transcript)
    if not text:
        return paragraph
    p = _paragraph_after(paragraph)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(2)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p


def _add_searchable_transcript_doc(doc: Document, transcript: str) -> None:
    text = normalize_ws(transcript)
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(2)
    run.font.color.rgb = RGBColor(255, 255, 255)


def figure_caption_name(record: dict[str, object]) -> str:
    env = str(record.get("env") or "").rstrip("*")
    return "Algorithm" if env in ALGORITHM_FLOAT_ENVS else "Figure"


def figure_main_caption(record: dict[str, object]) -> str:
    number = str(record.get("number") or "")
    caption = normalize_ws(str(record.get("caption") or ""))
    name = figure_caption_name(record)
    return f"{name} {number}: {caption}" if number and caption else caption


def is_placeholder_image_path(path_s: str) -> bool:
    return bool(re.fullmatch(r"idx\d+", path_s.strip()))


def record_image_paths(record: dict[str, object]) -> list[str]:
    image_paths = [str(item) for item in record.get("image_paths", []) if str(item)]
    fallback = str(record.get("image_path") or "")
    if not image_paths and fallback:
        image_paths = [fallback]
    return [path for path in image_paths if path.strip() and not is_placeholder_image_path(path)]


def build_figure_scratch(record: dict[str, object], config: Config) -> Document:
    scratch = Document()
    T._setup_doc(scratch)
    image_paths = record_image_paths(record)
    subcaptions = [str(item) for item in record.get("subcaptions", [])]
    image_widths = [item if isinstance(item, (int, float)) else None for item in record.get("image_widths", [])]
    main_caption = figure_main_caption(record)
    transcript = str(record.get("transcript") or "")

    resolved: list[Path] = []
    for raw in image_paths:
        try:
            resolved_path = resolve_graphic_path(raw, config)
            resolved.append(rasterize_graphic_for_docx(resolved_path, config))
        except Exception:
            continue
    if not resolved:
        p = scratch.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        T.inline_to_runs(p, "[LaTeX graphic omitted]", 9.0)
        if main_caption:
            _caption_paragraph(scratch, main_caption, FIGURE_CAPTION_PT)
        _add_searchable_transcript_doc(scratch, transcript)
        return scratch

    if len(resolved) <= 2:
        tbl = scratch.add_table(rows=2 if any(subcaptions) else 1, cols=len(resolved))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        _set_borderless_table(tbl)
        cell_width = 6.2 / len(resolved)
        for idx, image_path in enumerate(resolved):
            cell = tbl.cell(0, idx)
            cell.width = Inches(cell_width)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_picture(str(image_path), width=Inches(max(1.0, cell_width - 0.15)))
            if any(subcaptions):
                cap_cell = tbl.cell(1, idx)
                cap_p = cap_cell.paragraphs[0]
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_before = Pt(2)
                cap_p.paragraph_format.space_after = Pt(0)
                label = chr(ord("a") + idx) if idx < 26 else str(idx + 1)
                subcaption = subcaptions[idx] if idx < len(subcaptions) else ""
                if subcaption:
                    T.inline_to_runs(cap_p, f"({label}) {subcaption}", 9.0)
    else:
        for idx, image_path in enumerate(resolved):
            p = scratch.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(str(image_path), width=Inches(5.8))
            if idx < len(subcaptions) and subcaptions[idx]:
                label = chr(ord("a") + idx) if idx < 26 else str(idx + 1)
                _caption_paragraph(scratch, f"({label}) {subcaptions[idx]}", 9.0)

    if main_caption:
        _caption_paragraph(scratch, main_caption, FIGURE_CAPTION_PT)
    _add_searchable_transcript_doc(scratch, transcript)
    return scratch


def _move_before_anchor(anchor_p, elements: list[object]) -> None:
    anchor = anchor_p._p
    for element in elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        anchor.addprevious(element)
    anchor.getparent().remove(anchor)


def _clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _bind_block_keep_with_next(first_p: Paragraph, last_p: Paragraph) -> None:
    """Keep a whole figure block (image, subcaptions, caption) on one page so the
    image and its caption never split across a page break: set keep_with_next on
    every paragraph from first_p up to (not including) last_p, and keep_together on
    all of them. Safe no-op when first_p is last_p (single paragraph)."""
    parent = first_p._parent
    end = last_p._p
    node = first_p._p
    while node is not None:
        if node.tag == qn("w:p"):
            para = Paragraph(node, parent)
            para.paragraph_format.keep_together = True
            if node is not end:
                para.paragraph_format.keep_with_next = True
        if node is end:
            break
        node = node.getnext()


def insert_figure_record(anchor_p, record: dict[str, object], config: Config) -> None:
    image_paths = record_image_paths(record)
    subcaptions = [str(item) for item in record.get("subcaptions", [])]
    image_widths = [item if isinstance(item, (int, float)) else None for item in record.get("image_widths", [])]
    main_caption = figure_main_caption(record)
    transcript = str(record.get("transcript") or "")

    resolved: list[Path] = []
    for raw in image_paths:
        try:
            resolved_path = resolve_graphic_path(raw, config)
            resolved.append(rasterize_graphic_for_docx(resolved_path, config))
        except Exception:
            continue
    if not resolved:
        _clear_paragraph_content(anchor_p)
        anchor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor_p.paragraph_format.space_before = Pt(3)
        anchor_p.paragraph_format.space_after = Pt(3)
        T.inline_to_runs(anchor_p, "[LaTeX graphic omitted]", 9.0)
        last_p = anchor_p
        if main_caption:
            cap_p = _paragraph_after(anchor_p)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(3)
            cap_p.paragraph_format.space_after = Pt(6)
            T.inline_to_runs(cap_p, main_caption, FIGURE_CAPTION_PT)
            last_p = cap_p
        _bind_block_keep_with_next(anchor_p, last_p)
        _add_searchable_transcript_after(last_p, transcript)
        return

    _clear_paragraph_content(anchor_p)
    anchor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor_p.paragraph_format.space_before = Pt(3)
    anchor_p.paragraph_format.space_after = Pt(3)
    anchor_p.paragraph_format.line_spacing = 1.0
    last_p = anchor_p

    if len(resolved) == 1:
        width_in = image_widths[0] if image_widths else None
        anchor_p.add_run().add_picture(str(resolved[0]), width=Inches(min(width_in or 5.8, 6.2)))
    elif len(resolved) == 2:
        anchor_p.paragraph_format.space_after = Pt(2)
        anchor_p.add_run().add_picture(str(resolved[0]), width=Inches(2.85))
        anchor_p.add_run("    ")
        anchor_p.add_run().add_picture(str(resolved[1]), width=Inches(2.85))
        current = anchor_p
        if any(subcaptions):
            cap_p = _paragraph_after(current)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(0)
            cap_p.paragraph_format.space_after = Pt(3)
            left = f"(a) {subcaptions[0]}" if len(subcaptions) > 0 and subcaptions[0] else "(a)"
            right = f"(b) {subcaptions[1]}" if len(subcaptions) > 1 and subcaptions[1] else "(b)"
            T.inline_to_runs(cap_p, f"{left}        {right}", 9.0)
            current = cap_p
        last_p = current
    else:
        anchor_p.add_run().add_picture(str(resolved[0]), width=Inches(5.8))
        current = anchor_p
        if subcaptions and subcaptions[0]:
            cap_p = _paragraph_after(current)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(2)
            cap_p.paragraph_format.space_after = Pt(3)
            T.inline_to_runs(cap_p, f"(a) {subcaptions[0]}", 9.0)
            current = cap_p
        for idx, image_path in enumerate(resolved):
            if idx == 0:
                continue
            p = _paragraph_after(current)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.add_run().add_picture(str(image_path), width=Inches(5.8))
            current = p
            if idx < len(subcaptions) and subcaptions[idx]:
                label = chr(ord("a") + idx) if idx < 26 else str(idx + 1)
                cap_p = _paragraph_after(current)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_before = Pt(2)
                cap_p.paragraph_format.space_after = Pt(0)
                T.inline_to_runs(cap_p, f"({label}) {subcaptions[idx]}", 9.0)
                current = cap_p
        last_p = current

    if main_caption:
        cap_p = _paragraph_after(last_p)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(3)
        cap_p.paragraph_format.space_after = Pt(6)
        T.inline_to_runs(cap_p, main_caption, FIGURE_CAPTION_PT)
        last_p = cap_p
    _bind_block_keep_with_next(anchor_p, last_p)
    _add_searchable_transcript_after(last_p, transcript)


def build_table_scratch(record: dict[str, object], numbermap: dict[str, str], config: Config) -> Document:
    scratch = Document()
    T._setup_doc(scratch)
    land = bool(record.get("landscape"))
    avail = 9.0 if land else 6.4
    number = str(record.get("number") or numbermap.get(str(record.get("label")), ""))
    caption = str(record.get("caption") or "")
    cap = f"Table {number}: {caption}" if caption else f"Table {number}"
    notes = clean_notes(str(record.get("notes") or ""))
    engine = str(record.get("engine") or "")

    if record.get("detection_error"):
        raise RuntimeError(str(record["detection_error"]))
    if engine == "full_tabular":
        tabular_tex = str(record.get("tabular_tex") or "")
        if not tabular_tex and record.get("source_path"):
            tabular_tex = Path(str(record["source_path"])).read_text(encoding="utf-8", errors="replace")
        FT.add_fulltabular_table(
            scratch, tabular_tex, cap, notes,
            numbermap=numbermap, avail_in=avail, landscape=land,
        )
    elif engine == "estout":
        source_path = str(record.get("source_path") or "")
        ncols = record.get("ncols")
        if not source_path or ncols is None:
            raise RuntimeError("Estout table is missing source path or column count")
        # gridcols = data columns + 1 label column; drives the shared font policy.
        T.add_estout_table(
            scratch,
            source_path,
            cap,
            notes,
            int(ncols),
            numbermap=numbermap,
            avail_in=avail,
            data_pt=FT.table_data_pt(int(ncols) + 1, land),
        )
    else:
        raise RuntimeError("No table engine selected")
    return scratch


def assemble_docx(config: Config, body_docx: Path, transform: dict[str, object]) -> dict[str, object]:
    numbermap: dict[str, str] = dict(transform["numbermap"])  # type: ignore[arg-type]
    tables = {str(item["label"]): item for item in transform["tables"]}  # type: ignore[index]
    figures = {str(item["label"]): item for item in transform["figures"]}  # type: ignore[index]
    body = Document(body_docx)
    tables_done: list[str] = []
    figures_done: list[str] = []
    missing: list[dict[str, str]] = []

    for paragraph in list(body.paragraphs):
        text = paragraph.text.strip()
        mt = TABLE_RE.search(text)
        mf = FIG_RE.search(text)
        if mt:
            label = mt.group(1)
            rec = tables.get(label)
            if rec is None:
                missing.append({"kind": "table", "id": label, "error": "No table record"})
                continue
            try:
                scratch = build_table_scratch(rec, numbermap, config)
                if rec.get("landscape"):
                    # sidewaystable: isolate on its own rotated (landscape) page.
                    insert_doc_landscape_before(paragraph, scratch, body)
                else:
                    insert_doc_before(paragraph, scratch)
                tables_done.append(label)
            except Exception as exc:
                missing.append({"kind": "table", "id": label, "error": f"{type(exc).__name__}: {exc}"})
        elif mf:
            label = mf.group(1)
            rec = figures.get(label)
            if rec is None:
                missing.append({"kind": "figure", "id": label, "error": "No figure record"})
                continue
            try:
                insert_figure_record(paragraph, rec, config)
                figures_done.append(label)
            except Exception as exc:
                missing.append({"kind": "figure", "id": label, "error": f"{type(exc).__name__}: {exc}"})

    config.out.parent.mkdir(parents=True, exist_ok=True)
    body.save(config.out)
    return {
        "tables_inserted": tables_done,
        "figures_inserted": figures_done,
        "missing": missing,
        "docx": str(config.out),
        "figure_records": list(figures.values()),
    }


def count_footnotes(archive: zipfile.ZipFile) -> tuple[int, str]:
    if "word/footnotes.xml" not in archive.namelist():
        return 0, ""
    root = ET.fromstring(archive.read("word/footnotes.xml"))
    footnotes: list[tuple[int, str]] = []
    for footnote in root.findall(f"{W}footnote"):
        ftype = footnote.get(f"{W}type")
        fid_raw = footnote.get(f"{W}id")
        if ftype or fid_raw is None:
            continue
        try:
            fid = int(fid_raw)
        except ValueError:
            continue
        if fid < 1:
            continue
        footnotes.append((fid, text_from_element(footnote)))
    footnotes.sort(key=lambda item: item[0])
    return len(footnotes), (footnotes[0][1] if footnotes else "")


def verify_docx(docx_path: Path) -> dict[str, object]:
    doc = Document(docx_path)
    with zipfile.ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        document_root = ET.fromstring(document_xml)
        document_text = "".join(node.text or "" for node in document_root.iter(f"{W}t"))
        footnote_count, first_footnote = count_footnotes(archive)
        media = [name for name in archive.namelist() if name.startswith("word/media/")]

    unresolved_placeholders = document_text.count("%%TABLE:") + document_text.count("%%FIGURE:")
    unresolved_refs = len(re.findall(r"\\(?:Crefrange|crefrange|Cref|cref|eqref|autoref|ref)\s*\{", document_text))
    unresolved_citations = len(
        re.findall(
            r"\\(?:cite|citep|citet|citeauthor|citeyear|autocite|textcite|parencite|smartcite|footcite)\s*(?:\[[^\]]*\]\s*)*\{",
            document_text,
        )
    )
    return {
        "native_word_tables": len(doc.tables),
        "embedded_media": len(media),
        "footnotes": footnote_count,
        "first_footnote": first_footnote,
        "omml_equations": document_xml.count("<m:oMath"),
        "unresolved_placeholders": unresolved_placeholders,
        "unresolved_reference_commands": unresolved_refs,
        "unresolved_citation_commands": unresolved_citations,
        "document_text_chars": len(document_text),
    }


def export_docx_to_pdf(docx_path: Path, pdf_path: Path, workdir: Path) -> tuple[str | None, list[str]]:
    messages: list[str] = []
    if pdf_path.exists():
        pdf_path.unlink()

    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        doc.Close(False)
        word.Quit()
        if pdf_path.exists():
            return "word-com", messages
        messages.append("Word COM completed without producing the expected PDF")
    except Exception as exc:
        messages.append(f"Word COM failed: {type(exc).__name__}: {exc}")
        try:
            word.Quit()  # type: ignore[name-defined]
        except Exception:
            pass

    soffice = shutil.which("soffice.com") or shutil.which("soffice.exe") or shutil.which("soffice")
    if not soffice:
        messages.append("LibreOffice soffice was not found on PATH")
        return None, messages

    profile = workdir / ".lo_profile"
    profile.mkdir(parents=True, exist_ok=True)
    command = [
        soffice,
        f"-env:UserInstallation={profile.resolve().as_uri()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(pdf_path.parent),
        str(docx_path),
    ]
    proc = run_command(command, workdir, timeout=180)
    if proc.returncode != 0:
        messages.append(f"LibreOffice failed ({proc.returncode}): {proc.stderr.strip()}")
        return None, messages
    produced = pdf_path.parent / f"{docx_path.stem}.pdf"
    if produced.exists() and produced != pdf_path:
        produced.replace(pdf_path)
    if not pdf_path.exists():
        messages.append("LibreOffice completed without producing the expected PDF")
        return None, messages
    return "libreoffice", messages


def rasterize_first_pages(pdf_path: Path, out_dir: Path, limit: int = 3, dpi: int = 160) -> tuple[int, list[str]]:
    import fitz

    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("page_*.png"):
        old.unlink()
    written: list[str] = []
    with fitz.open(str(pdf_path)) as pdf:
        page_count = pdf.page_count
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        for idx in range(min(limit, page_count)):
            pix = pdf.load_page(idx).get_pixmap(matrix=matrix, alpha=False)
            out = out_dir / f"page_{idx + 1:02d}.png"
            pix.save(str(out))
            written.append(str(out))
    return page_count, written


def render_outputs(config: Config) -> dict[str, object]:
    # QA render lives in the workdir: writing <out-stem>.pdf beside the docx
    # would clobber a same-stem PDF built from the LaTeX source of truth.
    pdf_path = config.workdir / f"{config.out.stem}.docx-render.pdf"
    preview_dir = config.workdir / "preview"
    method, messages = export_docx_to_pdf(config.out, pdf_path, config.workdir)
    pages = None
    pngs: list[str] = []
    if pdf_path.exists():
        pages, pngs = rasterize_first_pages(pdf_path, preview_dir)
    return {
        "pdf": str(pdf_path) if pdf_path.exists() else None,
        "method": method,
        "messages": messages,
        "pages": pages,
        "preview_pngs": pngs,
    }


def compact_warnings(warnings: Iterable[str]) -> str:
    lines = [line.strip() for line in warnings if line.strip()]
    return "none" if not lines else "\n".join(f"  {line}" for line in lines)


def write_summary(config: Config, summary: dict[str, object]) -> None:
    summary_path = config.workdir / "conversion_summary.json"
    summary_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")


def print_summary(config: Config, summary: dict[str, object]) -> None:
    transform = summary["transform"]
    assembly = summary["assembly"]
    verify = summary["verify"]
    render = summary["render"]
    print("CONVERSION SUMMARY")
    print(f"Source: {config.main_tex}")
    print(f"Aux: {config.aux}")
    print(f"Output DOCX: {config.out}")
    print(f"Workdir: {config.workdir}")
    print(f"Font: {config.font}")
    print(f"Citation mode: {transform['citations']['mode']}")
    if transform["citations"].get("bibliography_paths"):
        print(f"Bibliographies: {len(transform['citations']['bibliography_paths'])}")
        print(f"CSL: {summary['pandoc'].get('csl') or transform['citations'].get('csl')}")
    if transform["citations"].get("warnings"):
        print("Citation warnings:")
        for warning in transform["citations"]["warnings"]:
            print(f"  {warning}")
    print(f"Tables detected: {transform['table_floats']}")
    print(f"Figures detected: {transform['figure_floats']}")
    print(f"Tables inserted: {len(assembly['tables_inserted'])}")
    print(f"Figures inserted: {len(assembly['figures_inserted'])}")
    print(f"Native Word tables: {verify['native_word_tables']}")
    print(f"Embedded media: {verify['embedded_media']}")
    print(f"Footnotes: {verify['footnotes']}")
    print(f"OMML equations: {verify['omml_equations']}")
    print(f"Unresolved placeholders: {verify['unresolved_placeholders']}")
    print(f"Unresolved reference commands: {verify['unresolved_reference_commands']}")
    print(f"Unresolved citation commands: {verify['unresolved_citation_commands']}")
    print(f"Rendered pages: {render.get('pages')}")
    print(f"PDF renderer: {render.get('method')}")
    if assembly["missing"]:
        print("Missing inserts:")
        for item in assembly["missing"]:
            print(f"  {item['kind']} {item['id']}: {item['error']}")
    if summary["pandoc"]["warnings"]:
        print("Pandoc warnings:")
        print(compact_warnings(summary["pandoc"]["warnings"]))
    if render.get("messages"):
        print("Render messages:")
        for msg in render["messages"]:
            print(f"  {msg}")
    print(f"Summary JSON: {config.workdir / 'conversion_summary.json'}")


def build_config(args: argparse.Namespace) -> Config:
    main_tex = Path(args.main_tex).expanduser().resolve()
    if not main_tex.exists():
        raise FileNotFoundError(f"Main TeX file does not exist: {main_tex}")
    aux = Path(args.aux).expanduser().resolve() if args.aux else main_tex.with_suffix(".aux")
    if not aux.exists():
        raise FileNotFoundError(f"Aux file is required and was not found: {aux}")
    out = Path(args.out).expanduser().resolve() if args.out else main_tex.with_suffix(".docx")
    workdir = Path(args.workdir).expanduser().resolve() if args.workdir else out.parent / f"_{out.stem}_work"

    source = main_tex.read_text(encoding="utf-8", errors="replace")
    auto_tables, auto_figures = auto_detect_dirs(source, main_tex)
    tables_dirs = [Path(args.tables_dir).expanduser().resolve(), *auto_tables] if args.tables_dir else auto_tables
    figures_dirs = [Path(args.figures_dir).expanduser().resolve(), *auto_figures] if args.figures_dir else auto_figures
    return Config(
        main_tex=main_tex,
        aux=aux,
        out=out,
        workdir=workdir,
        tables_dirs=dedupe_paths(tables_dirs),
        figures_dirs=dedupe_paths(figures_dirs),
        font=args.font,
        csl=getattr(args, "csl", None),
        render=not args.no_render,
    )


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert an academic LaTeX manuscript to DOCX with native tables.")
    parser.add_argument("main_tex", help="Path to the main .tex file")
    parser.add_argument("--out", help="Output .docx path")
    parser.add_argument("--aux", help="Path to the compiled .aux file; defaults to <main>.aux")
    parser.add_argument("--tables-dir", help="Directory to search for table inputs")
    parser.add_argument("--figures-dir", help="Directory to search for figure files")
    parser.add_argument("--font", default="Linux Libertine G", help="Body/table font for the reference DOCX")
    parser.add_argument(
        "--csl",
        help=(
            "CSL style path or URL for Pandoc citeproc when citation mode is detected; "
            "defaults to Pandoc's chicago-author-date style"
        ),
    )
    parser.add_argument("--workdir", help="Working directory; defaults to a temp directory beside the output")
    parser.add_argument("--no-render", action="store_true", help="Skip PDF/PNG rendering")
    return parser.parse_args(argv)


def convert(config: Config) -> dict[str, object]:
    config.workdir.mkdir(parents=True, exist_ok=True)
    reference_docx = config.workdir / "reference.docx"
    body_docx = config.workdir / "body.docx"

    transform = build_pandoc_tex(config)
    build_reference_docx(reference_docx, config.font)
    pandoc = run_pandoc(
        config,
        Path(str(transform["pandoc_tex"])),
        reference_docx,
        body_docx,
        dict(transform["citations"]),  # type: ignore[arg-type]
    )
    body_postprocess = postprocess_docx(body_docx, transform.get("display_math_numbers"))  # type: ignore[arg-type]
    assembly = assemble_docx(config, body_docx, transform)
    final_postprocess = postprocess_docx(config.out)
    verify = verify_docx(config.out)
    render = render_outputs(config) if config.render else {"pdf": None, "method": None, "messages": [], "pages": None, "preview_pngs": []}

    summary = {
        "config": {
            "main_tex": str(config.main_tex),
            "aux": str(config.aux),
            "out": str(config.out),
            "workdir": str(config.workdir),
            "tables_dirs": [str(path) for path in config.tables_dirs],
            "figures_dirs": [str(path) for path in config.figures_dirs],
            "font": config.font,
            "csl": config.csl,
            "render": config.render,
        },
        "transform": transform,
        "pandoc": pandoc,
        "body_postprocess": body_postprocess,
        "assembly": assembly,
        "final_postprocess": final_postprocess,
        "verify": verify,
        "render": render,
    }
    write_summary(config, summary)
    return summary


def main(argv: list[str] | None = None) -> int:
    try:
        config = build_config(parse_args(argv or sys.argv[1:]))
        summary = convert(config)
        print_summary(config, summary)
        if summary["verify"]["unresolved_placeholders"] != 0:
            return 2
        if summary["assembly"]["missing"]:
            return 3
        return 0
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}", file=sys.stderr)
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
