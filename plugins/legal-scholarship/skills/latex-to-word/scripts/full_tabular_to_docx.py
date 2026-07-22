# -*- coding: utf-8 -*-
"""Render full LaTeX tabular/booktabs tables as native Word tables.

This module complements tex_table_to_docx.py, which handles estout fragments.
It imports the established inline renderer and low-level Word helpers from that
engine, while parsing tables that already contain their own tabular environment.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Iterable

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt

from siunitx_expand import expand_siunitx_commands, format_bare_s_number, format_separate_uncertainty
from tex_table_to_docx import FONT, HEAVY, LIGHT, _border_tc, _run, inline_to_runs, normalize_note_text


@dataclass
class ColumnSpec:
    align: str
    width_in: float | None = None
    raw: str = ""


@dataclass
class CellSpec:
    text: str
    span: int = 1
    align: str | None = None
    rowspan: int = 1
    vcont: bool = False


@dataclass
class ParsedTabular:
    rows: list[list[CellSpec]]
    columns: list[ColumnSpec]
    top_heavy: set[int]
    bottom_heavy: set[int]
    top_light: set[int]
    partial_top: dict[int, list[tuple[int, int]]]
    space_before: set[int]


def _brace(s: str, open_idx: int) -> tuple[str, int]:
    """Return the braced contents starting at s[open_idx] and the next index."""
    depth = 0
    i = open_idx
    while i < len(s):
        c = s[i]
        if c == "\\" and i + 1 < len(s):
            i += 2
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return s[open_idx + 1:i], i + 1
        i += 1
    return s[open_idx + 1:], len(s)


def _bracket(s: str, open_idx: int) -> tuple[str, int]:
    depth = 0
    i = open_idx
    while i < len(s):
        c = s[i]
        if c == "\\" and i + 1 < len(s):
            i += 2
            continue
        if c == "[":
            depth += 1
        elif c == "]":
            depth -= 1
            if depth == 0:
                return s[open_idx + 1:i], i + 1
        i += 1
    return s[open_idx + 1:], len(s)


def _skip_ws(s: str, i: int) -> int:
    while i < len(s) and s[i].isspace():
        i += 1
    return i


def _read_group(s: str, i: int) -> tuple[str, int]:
    i = _skip_ws(s, i)
    if i >= len(s) or s[i] != "{":
        raise ValueError(f"expected braced group near: {s[i:i + 40]!r}")
    return _brace(s, i)


def _read_optional_bracket(s: str, i: int) -> tuple[str | None, int]:
    i = _skip_ws(s, i)
    if i < len(s) and s[i] == "[":
        return _bracket(s, i)
    return None, i


def _strip_comments(tex: str) -> str:
    out = []
    for line in tex.splitlines():
        cut = len(line)
        i = 0
        while i < len(line):
            if line[i] == "%" and (i == 0 or line[i - 1] != "\\"):
                cut = i
                break
            i += 1
        out.append(line[:cut])
    return "\n".join(out)


def _drop_leading_longtable_caption(body: str) -> str:
    stripped = body.lstrip()
    prefix = body[: len(body) - len(stripped)]
    if not stripped.startswith(r"\caption"):
        return body
    i = len(r"\caption")
    _, i = _read_optional_bracket(stripped, i)
    try:
        _, i = _read_group(stripped, i)
    except ValueError:
        return body
    i = _skip_ws(stripped, i)
    if stripped.startswith(r"\label", i):
        try:
            _, i = _read_group(stripped, i + len(r"\label"))
        except ValueError:
            pass
    i = _skip_ws(stripped, i)
    if stripped.startswith(r"\\", i):
        i += 2
    return prefix + stripped[i:]


def _clean_longtable_body(body: str) -> str:
    body = _drop_leading_longtable_caption(body)
    first_head = re.search(r"\\endfirsthead\b", body)
    last_foot = re.search(r"\\endlastfoot\b", body)
    if first_head and last_foot and first_head.start() < last_foot.end():
        body = body[: first_head.start()] + "\n" + body[last_foot.end() :]
    body = re.sub(r"\\end(?:firsthead|head|foot|lastfoot)\b", "", body)
    return body


def _extract_tabular(tabular_tex: str) -> tuple[str, str]:
    m = re.search(r"\\begin\{(tabular|tabulary|tabularx|longtable)\}", tabular_tex)
    if not m:
        raise ValueError("tabular_tex does not contain a supported tabular environment")
    env = m.group(1)
    i = m.end()
    _, i = _read_optional_bracket(tabular_tex, i)
    if env in {"tabulary", "tabularx"}:
        _, i = _read_group(tabular_tex, i)
    colspec, i = _read_group(tabular_tex, i)
    end_span = _matching_env_end(tabular_tex, i, env)
    if end_span is None:
        raise ValueError(f"could not find \\end{{{env}}}")
    body = tabular_tex[i:end_span[0]]
    if env == "longtable":
        body = _clean_longtable_body(body)
    return colspec, body


def _matching_env_end(text: str, start: int, env: str) -> tuple[int, int] | None:
    open_pat = re.compile(r"\\begin\s*\{" + re.escape(env) + r"\}")
    close_pat = re.compile(r"\\end\s*\{" + re.escape(env) + r"\}")
    depth = 1
    pos = start
    while pos < len(text):
        open_m = open_pat.search(text, pos)
        close_m = close_pat.search(text, pos)
        if close_m is None:
            return None
        if open_m is not None and open_m.start() < close_m.start():
            depth += 1
            pos = open_m.end()
            continue
        depth -= 1
        if depth == 0:
            return close_m.start(), close_m.end()
        pos = close_m.end()
    return None


def _expand_star_specs(spec: str) -> str:
    out = []
    i = 0
    while i < len(spec):
        if spec[i] == "*":
            j = _skip_ws(spec, i + 1)
            if j < len(spec) and spec[j] == "{":
                count_s, j = _brace(spec, j)
                j = _skip_ws(spec, j)
                if j < len(spec) and spec[j] == "{":
                    inner, j = _brace(spec, j)
                    try:
                        out.append(inner * int(count_s.strip()))
                        i = j
                        continue
                    except ValueError:
                        pass
        out.append(spec[i])
        i += 1
    expanded = "".join(out)
    return expanded if expanded == spec else _expand_star_specs(expanded)


def _latex_dim_to_inches(dim: str) -> float | None:
    rel = re.match(r"\s*([0-9.]+)\s*\\(?:line|text)width\s*$", dim)
    if rel:
        return float(rel.group(1)) * 6.4
    m = re.match(r"\s*([0-9.]+)\s*(cm|mm|in|pt)\s*$", dim)
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2)
    if unit == "in":
        return val
    if unit == "cm":
        return val / 2.54
    if unit == "mm":
        return val / 25.4
    if unit == "pt":
        return val / 72.27
    return None


def _alignment_from_spec(spec: str, default: str = "c") -> str:
    cleaned = re.sub(r"@\{[^{}]*\}", "", spec).replace("|", "")
    for c in cleaned:
        if c in "lcr":
            return c
        if c in "pmbX":
            return "l"
    return default


def parse_colspec_details(colspec: str) -> tuple[list[ColumnSpec], list[tuple[int, int]]]:
    spec = _expand_star_specs(colspec)
    cols: list[ColumnSpec] = []
    decimal_pairs: list[tuple[int, int]] = []
    pending_decimal_left: int | None = None

    def append_col(col: ColumnSpec) -> None:
        nonlocal pending_decimal_left
        cols.append(col)
        if pending_decimal_left is not None:
            decimal_pairs.append((pending_decimal_left, len(cols) - 1))
            pending_decimal_left = None

    i = 0
    while i < len(spec):
        c = spec[i]
        if c.isspace() or c == "|":
            i += 1
            continue
        if c in "@!":
            j = _skip_ws(spec, i + 1)
            if j < len(spec) and spec[j] == "{":
                raw_arg, i = _brace(spec, j)
                if c == "@" and raw_arg.strip() == "." and cols:
                    pending_decimal_left = len(cols) - 1
            else:
                i += 1
            continue
        if c in "><":
            j = _skip_ws(spec, i + 1)
            if j < len(spec) and spec[j] == "{":
                _, i = _brace(spec, j)
            else:
                i += 1
            continue
        if c in "lcr":
            append_col(ColumnSpec(c, raw=c))
            i += 1
            continue
        if c in "pmbX":
            raw = c
            width = None
            j = _skip_ws(spec, i + 1)
            if j < len(spec) and spec[j] == "{":
                raw_arg, j = _brace(spec, j)
                raw += "{" + raw_arg + "}"
                width = _latex_dim_to_inches(raw_arg)
            else:
                j = i + 1
            append_col(ColumnSpec("l", width, raw=raw))
            i = j
            continue
        if c == "S":
            j = _skip_ws(spec, i + 1)
            if j < len(spec) and spec[j] == "[":
                opt, j = _bracket(spec, j)
                append_col(ColumnSpec("r", raw=f"S[{opt}]"))
                i = j
            else:
                append_col(ColumnSpec("r", raw=c))
                i += 1
            continue
        if c == "D":
            j = i + 1
            for _ in range(3):
                j = _skip_ws(spec, j)
                if j < len(spec) and spec[j] == "{":
                    _, j = _brace(spec, j)
            append_col(ColumnSpec("r", raw="D"))
            i = j
            continue
        i += 1
    if not cols:
        raise ValueError(f"no columns parsed from colspec {colspec!r}")
    return cols, decimal_pairs


def parse_colspec(colspec: str) -> list[ColumnSpec]:
    cols, _decimal_pairs = parse_colspec_details(colspec)
    return cols


def _split_top_level(tex: str, sep: str) -> list[str]:
    parts: list[str] = []
    buf: list[str] = []
    depth = 0
    env_depth = 0
    begin_re = re.compile(r"\\begin\s*\{(tabular|tabularx|tabulary|array|longtable)\}")
    end_re = re.compile(r"\\end\s*\{(tabular|tabularx|tabulary|array|longtable)\}")
    i = 0
    while i < len(tex):
        begin_m = begin_re.match(tex, i)
        if begin_m:
            env_depth += 1
            buf.append(tex[i:begin_m.end()])
            i = begin_m.end()
            continue
        end_m = end_re.match(tex, i)
        if end_m:
            env_depth = max(0, env_depth - 1)
            buf.append(tex[i:end_m.end()])
            i = end_m.end()
            continue
        c = tex[i]
        if c == "\\":
            if sep == "\\\\" and i + 1 < len(tex) and tex[i + 1] == "\\" and depth == 0 and env_depth == 0:
                parts.append("".join(buf))
                buf = []
                i += 2
                j = _skip_ws(tex, i)
                if j < len(tex) and tex[j] == "[":
                    _, i = _bracket(tex, j)
                else:
                    i = j
                continue
            if sep == "&" and i + 1 < len(tex) and tex[i + 1] == "&":
                buf.append(tex[i:i + 2])
                i += 2
                continue
            if i + 1 < len(tex) and tex[i + 1] in "{}&%$#_":
                buf.append(tex[i:i + 2])
                i += 2
                continue
            buf.append(c)
            i += 1
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth = max(0, depth - 1)
        if sep == "&" and c == "&" and depth == 0 and env_depth == 0:
            parts.append("".join(buf))
            buf = []
        else:
            buf.append(c)
        i += 1
    parts.append("".join(buf))
    return parts


def _split_rows(body: str) -> list[str]:
    return _split_top_level(body, "\\\\")


def _split_cells(row: str) -> list[str]:
    return _split_top_level(row, "&")


def _parse_multicolumn(cell: str) -> CellSpec | None:
    s = cell.strip()
    if not s.startswith(r"\multicolumn"):
        return None
    i = len(r"\multicolumn")
    try:
        span_s, i = _read_group(s, i)
        align_s, i = _read_group(s, i)
        text, i = _read_group(s, i)
    except ValueError:
        return None
    try:
        span = int(span_s.strip())
    except ValueError:
        return None
    rest = s[i:].strip()
    if rest:
        text = (text + " " + rest).strip()
    return CellSpec(text=text.strip(), span=span, align=_alignment_from_spec(align_s))


def _parse_multirow(cell: str) -> CellSpec | None:
    s = cell.strip()
    if not s.startswith(r"\multirow"):
        return None
    i = len(r"\multirow")
    _, i = _read_optional_bracket(s, i)
    try:
        nrows_s, i = _read_group(s, i)
        _width, i = _read_group(s, i)
        text, i = _read_group(s, i)
    except ValueError:
        return None
    try:
        nrows = int(nrows_s.strip())
    except ValueError:
        nrows = 1
    rest = s[i:].strip()
    if rest:
        text = (text + " " + rest).strip()
    return CellSpec(text=text.strip(), rowspan=max(1, nrows))


def _consume_leading_commands(
    seg: str,
    rows: list[list[CellSpec]],
    top_heavy_next: bool,
    top_light_next: bool,
    addspace_next: bool,
    pending_cmid: list[tuple[int, int]],
    bottom_heavy: set[int],
) -> tuple[str, bool, bool, bool, list[tuple[int, int]]]:
    s = seg.strip()
    while s:
        cm = re.match(r"\\cmidrule(?:\([^)]*\))?\s*\{(\d+)-(\d+)\}\s*", s)
        if cm:
            pending_cmid.append((int(cm.group(1)) - 1, int(cm.group(2)) - 1))
            s = s[cm.end():].lstrip()
            continue
        rule = re.match(r"\\(toprule|midrule|bottomrule|addlinespace|hline)(?:\[[^\]]*\])?\s*", s)
        if not rule:
            break
        cmd = rule.group(1)
        if cmd == "toprule":
            top_heavy_next = True
        elif cmd in ("midrule", "hline"):
            top_light_next = True
        elif cmd == "addlinespace":
            addspace_next = True
        elif cmd == "bottomrule" and rows:
            bottom_heavy.add(len(rows) - 1)
        s = s[rule.end():].lstrip()
    return s, top_heavy_next, top_light_next, addspace_next, pending_cmid


def _parse_cells(seg: str, ncols: int) -> list[CellSpec]:
    parsed: list[CellSpec] = []
    for raw_cell in _split_cells(seg):
        mc = _parse_multicolumn(raw_cell)
        if mc is not None:
            parsed.append(mc)
        else:
            mr = _parse_multirow(raw_cell)
            if mr is not None:
                parsed.append(mr)
            else:
                parsed.append(CellSpec(raw_cell.strip()))
    span_total = sum(c.span for c in parsed)
    while span_total < ncols:
        parsed.append(CellSpec(""))
        span_total += 1
    return parsed


def _column_groups(ncols: int, decimal_pairs: list[tuple[int, int]]) -> list[list[int]]:
    pair_starts = {left: right for left, right in decimal_pairs}
    groups: list[list[int]] = []
    idx = 0
    while idx < ncols:
        right = pair_starts.get(idx)
        if right == idx + 1:
            groups.append([idx, right])
            idx = right + 1
        else:
            groups.append([idx])
            idx += 1
    return groups


def _group_index_for_raw(raw_idx: int, groups: list[list[int]]) -> int:
    for idx, group in enumerate(groups):
        if group[0] <= raw_idx <= group[-1]:
            return idx
    return max(0, len(groups) - 1)


def _collapse_decimal_text(left: str, right: str) -> str:
    left = left.strip()
    right = right.strip()
    if left and right:
        return f"{left}.{right}"
    return left or right


def _collapse_decimal_columns(
    rows: list[list[CellSpec]],
    columns: list[ColumnSpec],
    decimal_pairs: list[tuple[int, int]],
    partial_top: dict[int, list[tuple[int, int]]],
) -> tuple[list[list[CellSpec]], list[ColumnSpec], dict[int, list[tuple[int, int]]]]:
    if not decimal_pairs:
        return rows, columns, partial_top

    groups = _column_groups(len(columns), decimal_pairs)
    decimal_groups = {tuple(pair) for pair in decimal_pairs}
    new_columns: list[ColumnSpec] = []
    for group in groups:
        if len(group) == 2 and tuple(group) in decimal_groups:
            width = None
            widths = [columns[idx].width_in for idx in group]
            if all(width is not None for width in widths):
                width = sum(widths)  # type: ignore[arg-type]
            new_columns.append(ColumnSpec("r", width, raw=columns[group[0]].raw + "@{.}" + columns[group[1]].raw))
        else:
            new_columns.append(columns[group[0]])

    collapsed_rows: list[list[CellSpec]] = []
    for row in rows:
        positioned: list[tuple[int, int, CellSpec]] = []
        raw_idx = 0
        for cell in row:
            start = raw_idx
            end = raw_idx + max(1, cell.span) - 1
            positioned.append((start, end, cell))
            raw_idx = end + 1

        new_row: list[CellSpec] = []
        group_idx = 0
        while group_idx < len(groups):
            group = groups[group_idx]
            if len(group) == 2 and tuple(group) in decimal_groups:
                left_cell = next((cell for start, end, cell in positioned if start == group[0] and end == group[0]), None)
                right_cell = next((cell for start, end, cell in positioned if start == group[1] and end == group[1]), None)
                if left_cell is not None and right_cell is not None:
                    new_row.append(CellSpec(_collapse_decimal_text(left_cell.text, right_cell.text), align="r"))
                    group_idx += 1
                    continue

            covering = next((item for item in positioned if item[0] <= group[0] <= item[1]), None)
            if covering is None:
                new_row.append(CellSpec(""))
                group_idx += 1
                continue
            start, end, cell = covering
            if start < group[0]:
                group_idx += 1
                continue
            first_group = _group_index_for_raw(start, groups)
            last_group = _group_index_for_raw(end, groups)
            new_row.append(
                CellSpec(
                    cell.text,
                    span=last_group - first_group + 1,
                    align=cell.align,
                    rowspan=cell.rowspan,
                    vcont=cell.vcont,
                )
            )
            group_idx = last_group + 1
        collapsed_rows.append(new_row)

    new_partial: dict[int, list[tuple[int, int]]] = {}
    for row_idx, spans in partial_top.items():
        for c0, c1 in spans:
            new_partial.setdefault(row_idx, []).append(
                (_group_index_for_raw(c0, groups), _group_index_for_raw(c1, groups))
            )
    return collapsed_rows, new_columns, new_partial


def parse_fulltabular(tabular_tex: str) -> ParsedTabular:
    colspec, body = _extract_tabular(tabular_tex)
    columns, decimal_pairs = parse_colspec_details(colspec)
    ncols = len(columns)

    rows: list[list[CellSpec]] = []
    top_heavy: set[int] = set()
    bottom_heavy: set[int] = set()
    top_light: set[int] = set()
    partial_top: dict[int, list[tuple[int, int]]] = {}
    space_before: set[int] = set()

    pending_cmid: list[tuple[int, int]] = []
    top_heavy_next = False
    top_light_next = False
    addspace_next = False
    active_first_col_vmerge = 0

    for seg in _split_rows(_strip_comments(body)):
        s, top_heavy_next, top_light_next, addspace_next, pending_cmid = _consume_leading_commands(
            seg,
            rows,
            top_heavy_next,
            top_light_next,
            addspace_next,
            pending_cmid,
            bottom_heavy,
        )
        if not s.strip():
            continue
        ridx = len(rows)
        parsed = _parse_cells(s, ncols)
        if active_first_col_vmerge > 0:
            if parsed and parsed[0].span == 1 and not parsed[0].text.strip():
                parsed[0].vcont = True
            elif s.lstrip().startswith("&"):
                parsed.insert(0, CellSpec("", vcont=True))
                parsed = parsed[:ncols]
            active_first_col_vmerge -= 1
        if parsed and parsed[0].rowspan > 1 and not parsed[0].vcont:
            active_first_col_vmerge = parsed[0].rowspan - 1
        rows.append(parsed)
        if top_heavy_next:
            top_heavy.add(ridx)
            top_heavy_next = False
        if top_light_next:
            top_light.add(ridx)
            top_light_next = False
        if addspace_next:
            space_before.add(ridx)
            addspace_next = False
        if pending_cmid:
            partial_top.setdefault(ridx, []).extend(pending_cmid)
            pending_cmid = []

    if rows and not bottom_heavy:
        bottom_heavy.add(len(rows) - 1)
    rows, columns, partial_top = _collapse_decimal_columns(rows, columns, decimal_pairs, partial_top)
    return ParsedTabular(rows, columns, top_heavy, bottom_heavy, top_light, partial_top, space_before)


def _table_widths(columns: list[ColumnSpec], avail_in: float) -> list[float]:
    ncols = len(columns)
    if any(c.width_in for c in columns):
        base: list[float] = []
        for col in columns:
            if col.width_in is not None:
                base.append(col.width_in)
            elif col.align == "l":
                base.append(0.9)
            else:
                base.append(0.55)
        total = sum(base)
        if total <= 0:
            return [avail_in / ncols] * ncols
        scale = avail_in / total if total > avail_in or total < avail_in * 0.88 else 1.0
        return [max(0.35, w * scale) for w in base]

    if ncols == 1:
        return [avail_in]
    if ncols >= 8 and columns[0].align != "l":
        lead = [0.55, 0.95, 1.35]
        if ncols <= len(lead):
            return [avail_in / ncols] * ncols
        rem_count = ncols - len(lead)
        rem_width = max(0.42, (avail_in - sum(lead)) / rem_count)
        return lead + [rem_width] * rem_count
    label_in = min(1.8, avail_in * 0.34)
    data_count = ncols - 1
    data_in = min(0.70, max(0.42, (avail_in - label_in) / data_count))
    return [label_in] + [data_in] * data_count


def _set_tbl_grid(tbl, widths: Iterable[float]) -> None:
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, grid)
    existing = list(grid.findall(qn("w:gridCol")))
    widths = list(widths)
    while len(existing) < len(widths):
        gc = OxmlElement("w:gridCol")
        grid.append(gc)
        existing.append(gc)
    for gc, width in zip(existing, widths):
        gc.set(qn("w:w"), str(int(width * 1440)))


def _set_table_fixed(tbl) -> None:
    tbl_pr = tbl._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_cell_margins(tbl, top=0, start=36, bottom=0, end=36) -> None:
    tbl_pr = tbl._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for edge, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _makecell_parts(text: str) -> tuple[list[str], str | None] | None:
    s = text.strip()
    if not s.startswith(r"\makecell"):
        return None
    i = len(r"\makecell")
    align = None
    opt, i2 = _read_optional_bracket(s, i)
    if opt is not None:
        align = _alignment_from_spec(opt, default="c")
        i = i2
    try:
        inner, end = _read_group(s, i)
    except ValueError:
        return None
    if s[end:].strip():
        return None
    parts = [p.strip() for p in _split_rows(inner)]
    return parts, align


def _nested_tabular_parts(text: str) -> list[list[str]] | None:
    s = text.strip()
    if not re.match(r"\\begin\s*\{(?:tabular|tabularx|tabulary|array)\}", s):
        return None
    try:
        _colspec, body = _extract_tabular(s)
    except ValueError:
        return None
    rows: list[list[str]] = []
    for seg in _split_rows(_strip_comments(body)):
        seg = seg.strip()
        if not seg:
            continue
        rows.append([cell.strip() for cell in _split_cells(seg)])
    return rows or None


def _clean_inline(text: str) -> str:
    s = text.strip()
    s = expand_siunitx_commands(s)
    s = re.sub(r"\\vspace\{[^{}]*\}", "", s)
    s = s.replace(r"\footnotesize", "")
    s = re.sub(r"\\tnote\{([^{}]*)\}", r"^{\1}", s)
    s = s.replace("``", '"').replace("''", '"')
    def sqrt_repl(m):
        arg = m.group(1)
        return "\u221a" + arg if re.fullmatch(r"[A-Za-z0-9]+", arg) else "\u221a(" + arg + ")"

    s = re.sub(r"\\sqrt\{([^{}]*)\}", sqrt_repl, s)
    s = re.sub(r"\\textsubscript\{([^{}]*)\}", r"_{\1}", s)
    s = re.sub(r"\\textsuperscript\{([^{}]*)\}", r"^{\1}", s)
    return s.strip()


def _is_s_column(column: ColumnSpec | None) -> bool:
    return bool(column and column.raw.strip().startswith("S"))


def _render_inline(
    para,
    text: str,
    pt: float,
    numbermap=None,
    s_column: ColumnSpec | None = None,
) -> str | None:
    if _is_s_column(s_column):
        text = format_bare_s_number(text)
        if "(" in s_column.raw:
            text = format_separate_uncertainty(text)
    nested = _nested_tabular_parts(text)
    if nested is not None:
        for ridx, cols in enumerate(nested):
            if ridx:
                para.add_run().add_break()
            inline_to_runs(para, _clean_inline("  ".join(cols)), pt, numbermap)
        return "l"
    makecell = _makecell_parts(text)
    align_override = None
    if makecell is not None:
        parts, align_override = makecell
        for idx, part in enumerate(parts):
            if idx:
                para.add_run().add_break()
            inline_to_runs(para, _clean_inline(part), pt, numbermap)
        return align_override
    inline_to_runs(para, _clean_inline(text), pt, numbermap)
    return None


def _note_lines(notes_tex: str) -> list[str]:
    s = notes_tex.strip()
    if not s:
        return []
    s = re.sub(r"%[^\n]*(?:\n|$)", " ", s)
    s = s.replace(r"\footnotesize", "")
    item_re = re.compile(r"\\item(?:\[([^\]]*)\])?")
    matches = list(item_re.finditer(s))
    if not matches:
        return [s.strip()]
    lines: list[str] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(s)
        label = (match.group(1) or "").strip()
        text = re.sub(r"\s+", " ", s[start:end]).strip()
        if label:
            lines.append(f"^{label} {text}".strip())
        elif text:
            lines.append(text)
    return [line for line in lines if line]


def _docx_align(align: str | None, cidx: int, span: int, columns: list[ColumnSpec]):
    if align is None and cidx < len(columns):
        align = columns[cidx].align
    if align is None:
        align = "l" if cidx == 0 and span == 1 else "c"
    if align == "r":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if align == "l":
        return WD_ALIGN_PARAGRAPH.LEFT
    return WD_ALIGN_PARAGRAPH.CENTER


def _set_cell_width(cell, width_in: float) -> None:
    cell.width = Emu(int(Inches(width_in)))
    tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _row_border(tbl, r: int, edge: str, sz: int) -> None:
    if r < 0 or r >= len(tbl.rows):
        return
    for tc in tbl.rows[r]._tr.tc_lst:
        _border_tc(tc, edge, sz)


def _partial_top_border(tbl, r: int, c0: int, c1: int, sz: int) -> None:
    if r < 0 or r >= len(tbl.rows):
        return
    tcs = tbl.rows[r]._tr.tc_lst
    for c in range(c0, c1 + 1):
        if 0 <= c < len(tcs):
            _border_tc(tcs[c], "top", sz)


def _vmerge_tc(tc, restart: bool) -> None:
    tc_pr = tc.get_or_add_tcPr()
    vmerge = tc_pr.find(qn("w:vMerge"))
    if vmerge is None:
        vmerge = OxmlElement("w:vMerge")
        tc_pr.append(vmerge)
    vmerge.set(qn("w:val"), "restart" if restart else "continue")


def table_data_pt(total_cols: int, landscape: bool = False) -> float:
    """Body font (pt) for a converted table.

    Tables are sized UP from the LaTeX ``\\small``/``\\footnotesize`` so the Word
    output reads comfortably (publication house style), stepping down only for very
    wide tables so they still fit the portrait text column. A landscape page
    (sidewaystable) has ~9in of width, so it always takes the full size. ``total_cols``
    is the grid column count (estout: data columns + 1 label column; full tabular:
    the parsed column count). Shared by both table engines so the policy is defined
    once."""
    if landscape:
        return 11.0
    if total_cols >= 9:
        return 9.0
    if total_cols == 8:
        return 10.0
    return 11.0


def add_fulltabular_table(doc, tabular_tex, caption, notes_tex, numbermap=None,
                          avail_in=6.4, landscape=False):
    """Add a parsed full-tabular LaTeX table to an existing python-docx Document."""
    parsed = parse_fulltabular(tabular_tex)
    is_longtable = bool(re.search(r"\\begin\{longtable\}", tabular_tex))
    nrows = len(parsed.rows)
    ncols = len(parsed.columns)
    if nrows == 0:
        raise ValueError("tabular contains no data rows")
    data_pt = table_data_pt(ncols, landscape)

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption matches the table body font size.
    inline_to_runs(cap_p, _clean_inline(caption), data_pt, numbermap)

    widths = _table_widths(parsed.columns, avail_in)
    # Fill the full available width (scale up or down) so every table spans
    # margin-to-margin (landscape: the wider sidewaystable page).
    _wsum = sum(widths)
    if _wsum > 0:
        widths = [w * avail_in / _wsum for w in widths]
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _set_table_fixed(tbl)
    _set_cell_margins(tbl)
    _set_tbl_grid(tbl, widths)

    for r, row in enumerate(parsed.rows):
        cidx = 0
        for cell_spec in row:
            if cidx >= ncols:
                break
            span = max(1, min(cell_spec.span, ncols - cidx))
            cell = tbl.cell(r, cidx)
            if span > 1:
                cell = cell.merge(tbl.cell(r, cidx + span - 1))
            if cell_spec.rowspan > 1:
                _vmerge_tc(cell._tc, True)
            elif cell_spec.vcont:
                _vmerge_tc(cell._tc, False)
            width = sum(widths[cidx:cidx + span])
            _set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cell.paragraphs[0]
            # Tight, uniformly single-spaced rows (drop the \addlinespace 4pt gap
            # and the longtable 1.08 leading) for a compact publication look.
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            column = parsed.columns[cidx] if span == 1 and cidx < len(parsed.columns) else None
            align_override = _render_inline(para, cell_spec.text, data_pt, numbermap, s_column=column)
            effective_align = align_override or cell_spec.align
            para.alignment = _docx_align(effective_align, cidx, span, parsed.columns)
            cidx += span

    for r in parsed.top_heavy:
        _row_border(tbl, r, "top", HEAVY)
    for r in parsed.top_light:
        _row_border(tbl, r, "top", LIGHT)
    for r in parsed.bottom_heavy:
        _row_border(tbl, r, "bottom", HEAVY)
    for r, spans in parsed.partial_top.items():
        for c0, c1 in spans:
            _partial_top_border(tbl, r, c0, c1, LIGHT)

    for idx, note_line in enumerate(_note_lines(notes_tex)):
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_p.paragraph_format.space_before = Pt(3 if idx == 0 else 0)
        note_p.paragraph_format.space_after = Pt(0)
        inline_to_runs(note_p, normalize_note_text(_clean_inline(note_line)), 8.2, numbermap)

    return tbl
